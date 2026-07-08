"""Generate Voice-to-Command fund transfer (UPI limit → user choice → IMPS/NEFT) demo script."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Voice-to-Command-Fund-Transfer-Demo-Script.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
    return row


def build():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Indian Bank — Voice-to-Command Demo Script")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 61, 124)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("UPI limit → fallback options → customer choice → Fund Transfer · ~2 minutes")
    r.font.size = Pt(12)
    r.italic = True

    doc.add_paragraph()

    add_heading(doc, "Storyline at a glance", 2)
    doc.add_paragraph(
        "A customer asks to send ₹5 lakh to Rahul Sharma via UPI. The AI concierge flags that "
        "UPI allows only ₹1 lakh per transaction, then presents clear fallback options — IMPS, "
        "NEFT, RTGS, or splitting into smaller UPI payments. The customer picks how they want "
        "to proceed (e.g. IMPS for instant other-bank transfer), and the app opens Fund Transfer "
        "with that choice pre-set. The voice assistant then guides them through the form to completion."
    )

    add_heading(doc, "Before you start", 2)
    add_bullet(doc, "App at http://localhost:5173 · npm run dev running · mic enabled.")
    add_bullet(doc, "Turn Voice-to-Command nav-only mode OFF (grey “Mode inactive” in Demo panel).")
    add_bullet(doc, "Use the AI Banking Assistant on home — not direct screen-jump mode.")
    add_bullet(doc, "Demo MPIN: 1234")

    note = doc.add_paragraph()
    r = note.add_run("Channel limits: ")
    r.bold = True
    note.add_run(
        "UPI ₹1,00,000 · IMPS ₹5,00,000 · NEFT/RTGS no per-txn cap. "
        "When UPI is blocked, the assistant should offer alternatives — not force a single rail."
    )

    doc.add_paragraph()

    add_heading(doc, "Fallback options the assistant offers", 2)
    doc.add_paragraph(
        "After detecting a UPI limit breach, the AI presents these choices (read aloud or show in chat):"
    )

    options = [
        (
            "IMPS",
            "Instant · 24×7 · up to ₹5 lakh",
            "Best when payee is at another bank and money must move today.",
            "Say: “IMPS” or “instant transfer”",
        ),
        (
            "NEFT",
            "Batch · typically same/next working day · no cap",
            "Best for large within-bank or other-bank transfers when instant is not required.",
            "Say: “NEFT”",
        ),
        (
            "RTGS",
            "Real-time · high value · usually ₹2 lakh minimum",
            "Best for very large same-day transfers (mention if amount were ₹10 lakh+).",
            "Say: “RTGS”",
        ),
        (
            "Split UPI",
            "Multiple UPI payments · ₹1 lakh each",
            "Best if customer insists on UPI — e.g. five payments of ₹1 lakh (not recommended for demo).",
            "Say: “Split into UPI” (optional — skip in main demo)",
        ),
    ]

    t_opts = doc.add_table(rows=1, cols=4)
    t_opts.style = "Table Grid"
    hdr = t_opts.rows[0].cells
    hdr[0].text = "Option"
    hdr[1].text = "Speed / limit"
    hdr[2].text = "When to suggest"
    hdr[3].text = "Customer says"
    for row in options:
        add_table_row(t_opts, row)

    doc.add_paragraph()

    add_heading(doc, "Act 1 — UPI request & limit advisory (0:00 – 0:25)", 2)
    doc.add_paragraph(
        "Goal: Customer states intent; AI catches the UPI breach before any wrong screen opens."
    )

    steps1 = [
        (
            "0:00",
            "Tap 🧑‍💼 AI Banking Assistant on home.",
            "Greeting: “Tell me what you'd like to do…”",
        ),
        (
            "0:05",
            "Say:\n“Send 5 lakh rupees to Rahul Sharma via UPI”",
            "Assistant responds — does NOT open UPI yet.",
        ),
        (
            "0:12",
            "Listen for limit advisory.",
            "💭 “₹5 lakh exceeds UPI's ₹1 lakh per-transaction limit.” "
            "Spoken reply: “UPI can't handle ₹5 lakh in one go. Here are your options…”",
        ),
        (
            "0:18",
            "Assistant lists fallbacks (see table above).",
            "Chat shows numbered options: 1 IMPS · 2 NEFT · 3 RTGS · 4 Split UPI.",
        ),
    ]

    t1 = doc.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    hdr = t1.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps1:
        add_table_row(t1, row)

    doc.add_paragraph()

    add_heading(doc, "Act 2 — Customer chooses a fallback (0:25 – 0:45)", 2)
    doc.add_paragraph(
        "Goal: Customer stays in control — they pick the payment rail, not the system."
    )

    steps2 = [
        (
            "0:25",
            "Say (recommended for demo):\n“Use IMPS — other bank, HDFC”",
            "Assistant confirms: “Got it — IMPS to HDFC for ₹5 lakh. Opening fund transfer.” "
            "TTS: “Redirecting to fund transfer.”",
        ),
        (
            "0:32",
            "Alternative — say “NEFT” instead.",
            "Fund Transfer opens with Within Bank / NEFT path pre-selected if payee is Indian Bank.",
        ),
        (
            "0:38",
            "Wait for navigation.",
            "IMPS Fund Transfer screen slides in · AI overlay open · context carries payee + amount.",
        ),
        (
            "0:42",
            "Point out the chosen rail on screen.",
            "Form header shows Other Bank (IMPS) or Within Bank (NEFT) based on customer choice.",
        ),
    ]

    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    hdr = t2.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps2:
        add_table_row(t2, row)

    branch = doc.add_paragraph()
    r = branch.add_run("Branch — if customer picks NEFT (within Indian Bank): ")
    r.bold = True
    branch.add_run(
        "Say “NEFT to Indian Bank account” → when form opens, confirm Within Bank tab · "
        "provide account number + payee name · amount ₹5 lakh."
    )

    doc.add_paragraph()

    add_heading(doc, "Act 3 — Complete transfer via voice (0:45 – 1:45)", 2)
    doc.add_paragraph(
        "Goal: Finish the IMPS path (other bank). Fields fill live as customer speaks."
    )

    steps3 = [
        (
            "0:45",
            "If bank not yet set — AI asks Indian Bank or other?\nSay: “Other bank, HDFC”",
            "Other Bank tab active.",
        ),
        (
            "0:52",
            "AI: Account or mobile?\nSay: “Account number”",
            "Account path selected.",
        ),
        (
            "0:58",
            "Say IFSC: “HDFC0001234”",
            "IFSC field fills.",
        ),
        (
            "1:04",
            "Say account: “9876543210123456”",
            "Payee account fills.",
        ),
        (
            "1:10",
            "Say amount: “5 lakh” (may already be pre-filled from Act 1).",
            "Amount shows ₹5,00,000.",
        ),
        (
            "1:16",
            "Say: “Skip” for remarks.",
            "Review screen ready.",
        ),
        (
            "1:22",
            "Tap Proceed on review.",
            "MPIN sheet opens.",
        ),
    ]

    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    hdr = t3.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps3:
        add_table_row(t3, row)

    doc.add_paragraph()

    add_heading(doc, "Act 4 — Authorise & success (1:45 – 2:00)", 2)

    steps4 = [
        ("1:45", "Enter MPIN: 1234", "Processing · IMPS RRN generated."),
        ("1:52", "Success → Home", "Demo complete."),
    ]

    t4 = doc.add_table(rows=1, cols=3)
    t4.style = "Table Grid"
    hdr = t4.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps4:
        add_table_row(t4, row)

    doc.add_paragraph()

    add_heading(doc, "Sample dialogue — limit + choice moment", 2)
    dialogue = [
        ("Customer", "Send 5 lakh rupees to Rahul Sharma via UPI."),
        (
            "Assistant",
            "💭 ₹5 lakh exceeds UPI's ₹1 lakh limit.\n"
            "UPI allows up to ₹1 lakh per transaction, so ₹5 lakh can't go through UPI in one payment. "
            "You can choose:\n"
            "1 — IMPS (instant, up to ₹5 lakh)\n"
            "2 — NEFT (same/next day, no cap)\n"
            "3 — RTGS (real-time, for high value)\n"
            "4 — Split into multiple UPI payments of ₹1 lakh each\n"
            "Which would you prefer?",
        ),
        ("Customer", "IMPS to HDFC, please."),
        (
            "Assistant",
            "Perfect — IMPS to HDFC for ₹5 lakh. Opening fund transfer now.",
        ),
    ]
    for speaker, line in dialogue:
        p = doc.add_paragraph()
        r = p.add_run(f"{speaker}: ")
        r.bold = True
        p.add_run(line)

    doc.add_paragraph()

    add_heading(doc, "Quick reference — voice lines", 2)
    lines = [
        ("Trigger limit check", "Send 5 lakh rupees to Rahul Sharma via UPI"),
        ("Choose IMPS", "Use IMPS / IMPS please / instant transfer"),
        ("Choose NEFT", "Use NEFT / NEFT to Indian Bank"),
        ("Choose RTGS", "Use RTGS"),
        ("Decline UPI split", "No, use IMPS instead"),
        ("Other bank + IFSC", "HDFC0001234"),
        ("Account number", "9876543210123456"),
        ("Confirm amount", "5 lakh"),
    ]
    t5 = doc.add_table(rows=1, cols=2)
    t5.style = "Table Grid"
    t5.rows[0].cells[0].text = "Intent"
    t5.rows[0].cells[1].text = "Say this"
    for intent, phrase in lines:
        add_table_row(t5, [intent, f'"{phrase}"'])

    doc.add_paragraph()

    add_heading(doc, "Presenter tips", 2)
    tips = [
        "Pause after the options list — the choice moment is the hero of this demo.",
        "Recommend IMPS for the main run (instant, other bank, ₹5L fits the cap).",
        "Use NEFT branch if audience asks about within-bank or next-day settlement.",
        "Mention RTGS verbally even if you don't run it — shows the assistant knows all rails.",
        "Skip “Split UPI” unless asked — it highlights why fallback exists.",
        "If the live AI auto-routes without asking, still say your choice (“IMPS please”) — the transfer agent will follow.",
        "Type lines in chat if mic fails.",
    ]
    for tip in tips:
        add_bullet(doc, tip)

    add_heading(doc, "What this demo proves", 2)
    proofs = [
        "UPI limit enforced with a clear explanation — no silent failure or wrong screen.",
        "Customer-led fallback: IMPS / NEFT / RTGS / split UPI offered before proceeding.",
        "Context preserved: payee, amount, and chosen rail carry into Fund Transfer.",
        "Voice co-pilot completes the form after the customer’s channel choice.",
        "Full authorisation loop with MPIN and success confirmation.",
    ]
    for p in proofs:
        add_bullet(doc, p)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
