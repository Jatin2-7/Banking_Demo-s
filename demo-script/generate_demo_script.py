"""Generate Voice-to-Navigation demo script Word document."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Voice-to-Navigation-Demo-Script.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
    return row


def build():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Indian Bank — Voice-to-Navigation Demo Script")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 61, 124)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Complete workflow · ~90 seconds")
    r.font.size = Pt(12)
    r.italic = True

    doc.add_paragraph()

    add_heading(doc, "Storyline at a glance", 2)
    doc.add_paragraph(
        "Prateek opens the Indian Bank app and uses voice alone to review his savings "
        "account statement for a chosen date window, spot a few key debits, then "
        "navigate straight to card services and change his card PIN — without hunting "
        "through menus."
    )

    add_heading(doc, "Before you start (30 sec setup)", 2)
    add_bullet(doc, " — App running at http://localhost:5173 (npm run dev from repo root)")
    add_bullet(doc, " — In the Demo panel (right), turn on Voice-to-Command mode (green “Mode active” banner)")
    add_bullet(doc, " — Mic enabled (ElevenLabs STT or browser speech in Chrome/Edge)")
    add_bullet(doc, " — Home screen visible inside the phone frame; balance card shows Savings account")

    add_heading(doc, "Act 1 — Voice navigate to Transaction History (0:00 – 0:35)", 2)
    doc.add_paragraph(
        "Goal: Prove hands-free screen routing from the home dashboard."
    )

    steps1 = [
        (
            "0:00",
            "Tap the 🧑‍💼 AI RM button on the home screen (bottom-right of phone).",
            "AI RM · Voice Navigation panel slides up with greeting: “Tell me the screen you want to open…”",
        ),
        (
            "0:05",
            "Say (or type):\n“Show me my account transaction history”",
            "App speaks: “Opening your account statement.” Account Statement screen slides in.",
        ),
        (
            "0:12",
            "Optional — tap Change Period on the statement header, or stay on Recent Transactions.",
            "Savings account XXXXXX1762, balance ₹2,51,000.00, scrollable txn list loads.",
        ),
        (
            "0:18",
            "With the Account Assistant open, say:\n“Show me my transactions from 1 April to 15 May 2026”",
            "Assistant reads the on-screen list and highlights salary credits (₹1,25,000 on 1 Apr & 1 May), "
            "car-loan EMIs (₹45,000), and recent UPI spends — only from data visible on the statement.",
        ),
        (
            "0:30",
            "Follow-up (optional):\n“What was my biggest debit in that period?”",
            "Assistant answers from the statement (e.g. salary is credit; largest debit is EMI ₹45,000 or IMPS ₹12,000).",
        ),
    ]

    t1 = doc.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    hdr = t1.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps1:
        add_table_row(t1, row)

    doc.add_paragraph()

    add_heading(doc, "Act 2 — Review & pivot (0:35 – 0:55)", 2)
    doc.add_paragraph(
        "Goal: Brief pause on the statement so the audience registers real transaction data, "
        "then pivot to a new intent without going back to menus."
    )

    steps2 = [
        (
            "0:35",
            "Scroll the statement — point out salary credit and one UPI bill payment.",
            "Audience sees live mock data (BESCOM, Swiggy, HDFC Car Loan EMI, etc.).",
        ),
        (
            "0:42",
            "Close the Account Assistant (×) or leave it open — either works.",
            "Full-width statement remains; 🧑‍💼 FAB available again.",
        ),
        (
            "0:45",
            "Tap 🧑‍💼 again OR use Demo panel mic → say:\n“I want to change my card PIN”",
            "Voice router matches card PIN intent → Debit Card dashboard opens with Set/Reset Card PIN flow.",
        ),
    ]

    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    hdr = t2.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps2:
        add_table_row(t2, row)

    note = doc.add_paragraph()
    r = note.add_run("Note — Credit card vs debit card: ")
    r.bold = True
    note.add_run(
        "Saying “open my credit card” opens the Credit Card dashboard (limits, statement). "
        "PIN change is implemented on the Debit Card screen. For this demo, use "
        "“change my card PIN” or “reset my debit card PIN” to land directly on the PIN reset form."
    )

    doc.add_paragraph()

    add_heading(doc, "Act 3 — Change card PIN on screen (0:55 – 1:30)", 2)
    doc.add_paragraph(
        "Goal: Complete a self-service card action after voice navigation — no branch visit."
    )

    steps3 = [
        (
            "0:55",
            "If not auto-opened: on Debit Card → Card Actions → tap Set/Reset Card PIN.",
            "Modal: Set/Reset Card PIN with New PIN and Re-enter PIN fields.",
        ),
        (
            "1:00",
            "Enter New PIN: 5678 · Re-enter PIN: 5678 · tap Confirm.",
            "Authentication screen — Transaction PIN keypad (6 dots).",
        ),
        (
            "1:08",
            "Enter any 6-digit auth PIN on the keypad (e.g. 123456).",
            "Success dialog: “Card PIN changed successfully” with green checkmark.",
        ),
        (
            "1:15",
            "Tap Ok → Home.",
            "Returns to Indian Bank home dashboard. Demo complete.",
        ),
    ]

    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    hdr = t3.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps3:
        add_table_row(t3, row)

    doc.add_paragraph()

    add_heading(doc, "Quick reference — voice lines", 2)
    lines = [
        ("Open statement", "Show me my account transaction history"),
        ("Filter by dates (in-statement AI)", "Show me my transactions from 1 April to 15 May 2026"),
        ("Ask about txns", "What was my salary credit last month?"),
        ("Navigate to PIN change", "Change my card PIN  /  Reset my debit card PIN"),
        ("Open credit card (alternate)", "Open my credit card"),
        ("Return home", "Go back home"),
    ]
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text = "Intent"
    t4.rows[0].cells[1].text = "Say this"
    for intent, phrase in lines:
        add_table_row(t4, [intent, f'"{phrase}"'])

    doc.add_paragraph()

    add_heading(doc, "Presenter tips", 2)
    tips = [
        "Enable Voice-to-Command mode first — the home AI RM then navigates only (no long chat).",
        "Demo panel → Voice-to-Command tab has one-click example phrases if the mic fails.",
        "Transaction History opens with the Account Assistant already visible — use it for date-range questions.",
        "PIN demo uses Debit Card screen; mention Credit Card dashboard is one voice command away if asked.",
        "Total runtime: ~90 sec with one follow-up question; can trim to ~60 sec by skipping Act 2 scroll.",
    ]
    for tip in tips:
        add_bullet(doc, tip)

    add_heading(doc, "What this demo proves", 2)
    proofs = [
        "Voice-to-Command: utterance → screen, no menu drilling.",
        "Contextual AI on Transaction History reads the same data shown on screen.",
        "Seamless pivot from “review money” to “secure my card” in one session.",
        "End-to-end self-service PIN change with on-device confirmation flow.",
    ]
    for p in proofs:
        add_bullet(doc, p)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
