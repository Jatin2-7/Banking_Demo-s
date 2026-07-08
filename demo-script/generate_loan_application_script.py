"""Generate Voice Assistance — Loan Application demo script."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Voice-Assistance-Loan-Application-Demo-Script.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
    return row


def build():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Indian Bank — Voice Assistance Demo Script")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 61, 124)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Live form-fill with AI Loan RM · ~2 minutes")
    r.font.size = Pt(12)
    r.italic = True

    doc.add_paragraph()

    add_heading(doc, "Storyline at a glance", 2)
    doc.add_paragraph(
        "Prateek wants a personal loan. Instead of tapping through every dropdown, he opens the "
        "Loan Application (LOS) screen and speaks naturally to the AI relationship manager. "
        "Fields fill live on the form as he talks — occupation, product, purpose, amount, tenure, "
        "and branch pincode — then he reviews and submits. The demo ends with a captured application "
        "reference number."
    )

    add_heading(doc, "Before you start", 2)
    add_bullet(doc, "App at http://localhost:5173 · npm run dev · OpenAI key in .env (required for AI form-fill).")
    add_bullet(doc, "Mic enabled (ElevenLabs STT or browser speech).")
    add_bullet(doc, "Home screen visible · customer name on loan form will show “Prateek”.")

    add_heading(doc, "Form fields the voice assistant fills", 2)
    fields = [
        ("Occupation type", "Salaried / Self employed / Professional"),
        ("Sub product", "IB Clean Loan to Salary · Personal loan · Overdraft against salary"),
        ("Purpose of loan", "Education · Medical · Marriage/edu/medical · Family/household · Other"),
        ("Variant", "Standard · Flexi · Lite"),
        ("Facility type", "Term loan · Term + OD · Overdraft"),
        ("Proposal type", "Fresh proposal · Top-up · Takeover"),
        ("Interest type", "Floating (pre-set, read-only)"),
        ("Requested amount", "e.g. ₹5,00,000"),
        ("Tenure (months)", "e.g. 36"),
        ("Processing branch", "6-digit pincode, e.g. 560034"),
    ]
    t_fields = doc.add_table(rows=1, cols=2)
    t_fields.style = "Table Grid"
    t_fields.rows[0].cells[0].text = "Field"
    t_fields.rows[0].cells[1].text = "Options / example"
    for row in fields:
        add_table_row(t_fields, row)

    doc.add_paragraph()

    add_heading(doc, "Act 1 — Open loan application (0:00 – 0:20)", 2)
    doc.add_paragraph(
        "Goal: Reach the Loan details screen (Step 1 of 3) with the AI assistant ready."
    )

    steps1 = [
        (
            "0:00",
            "Option A — tap Services → “Apply New Loan” on home.\n"
            "Option B — tap 🧑‍💼 AI Banking Assistant → say “I want to apply for a personal loan”.\n"
            "Option C — Voice-to-Command mode ON → say “apply for a loan”.",
            "Loan Application screen opens · Step 1 “Loan details” · progress dots at top.",
        ),
        (
            "0:10",
            "Tap the AI button (loan assistant FAB) if not already open.",
            "Loan form assist panel slides up: “Let's fill this form together…”",
        ),
        (
            "0:15",
            "Point at the empty form behind the assistant.",
            "Audience sees dropdowns and amount/tenure fields — all blank.",
        ),
    ]

    t1 = doc.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    hdr = t1.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps1:
        add_table_row(t1, row)

    doc.add_paragraph()

    add_heading(doc, "Act 2 — Voice-fill the form (0:20 – 1:30)", 2)
    doc.add_paragraph(
        "Goal: Show live co-pilot behaviour — the RM can fill multiple fields from one natural sentence. "
        "Highlighted rows pulse gold as each field updates."
    )

    steps2 = [
        (
            "0:20",
            "Say (compound — fills several fields at once):\n"
            "“I'm salaried. I need IB Clean Loan to Salary for medical expenses — "
            "standard variant, term loan, fresh proposal.”",
            "Occupation → Salaried · Sub product → IB Clean Loan · Purpose → Medical · "
            "Variant → Standard · Facility → Term loan · Proposal → Fresh proposal.",
        ),
        (
            "0:40",
            "Assistant confirms briefly and asks for amount + tenure.\n"
            "Say: “5 lakh for 36 months”",
            "Amount → ₹5,00,000 · Tenure → 36 months.",
        ),
        (
            "0:55",
            "Say: “Processing branch pincode 560034” (Bengaluru example).",
            "Processing branch field fills · all required rows complete.",
        ),
        (
            "1:05",
            "Optional — if a dropdown is unclear, assistant uses request_field to highlight it.",
            "That row rings gold · assistant explains choices in plain language.",
        ),
        (
            "1:15",
            "Say: “Looks good — proceed” or tap Next on the form.",
            "Step advances to “2. Review application”.",
        ),
    ]

    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    hdr = t2.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps2:
        add_table_row(t2, row)

    doc.add_paragraph()

    add_heading(doc, "Act 3 — Review & submit (1:30 – 1:50)", 2)

    steps3 = [
        (
            "1:30",
            "Scroll review screen — confirm occupation, product, ₹5,00,000, 36 months, pincode.",
            "Summary table matches what voice filled on Step 1.",
        ),
        (
            "1:38",
            "Say: “Submit the application” or tap Submit Application.",
            "Step 3 “Status” · success card with reference number (LOS…).",
        ),
        (
            "1:45",
            "Read reference aloud · tap Back to home.",
            "Returns to Indian Bank dashboard · demo complete.",
        ),
    ]

    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    hdr = t3.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Presenter action / Voice line"
    hdr[2].text = "What the audience sees"
    for row in steps3:
        add_table_row(t3, row)

    doc.add_paragraph()

    add_heading(doc, "Sample dialogue — voice form-fill", 2)
    dialogue = [
        ("Customer", "I'm salaried and want IB Clean Loan for medical expenses."),
        (
            "AI RM",
            "Got it — I've set salaried, IB Clean Loan to Salary, and medical purpose on your screen. "
            "Would you like standard or flexi variant, and term loan or overdraft?",
        ),
        ("Customer", "Standard term loan, fresh proposal. Five lakh for 36 months."),
        (
            "AI RM",
            "Done — ₹5,00,000 for 36 months is filled. Which pincode should we use for processing branch?",
        ),
        ("Customer", "560034."),
        (
            "AI RM",
            "All required fields look complete. Tap Next when you're ready to review.",
        ),
    ]
    for speaker, line in dialogue:
        p = doc.add_paragraph()
        r = p.add_run(f"{speaker}: ")
        r.bold = True
        p.add_run(line)

    doc.add_paragraph()

    add_heading(doc, "Alternate entry paths", 2)
    paths = [
        ("Idle assist", "Wait ~8 seconds on the loan form without tapping — AI opens automatically with idle primer."),
        ("Stuck user", "Rapid taps on empty fields trigger “Need help?” — tap Help to open the RM assistant."),
        ("Hindi / Hinglish", "Same flow in Hindi — assistant replies in Hindi; form values stay in English ids."),
        ("Step-by-step", "Instead of one compound sentence, answer each question the RM asks one at a time."),
    ]
    for name, desc in paths:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}: ")
        r.bold = True
        p.add_run(desc)

    doc.add_paragraph()

    add_heading(doc, "Quick reference — voice lines", 2)
    lines = [
        ("Open loan (AI home)", "I want to apply for a personal loan"),
        ("Open loan (nav mode)", "Apply for a loan"),
        ("Occupation + product", "I'm salaried, need IB Clean Loan for medical expenses"),
        ("Variant + facility", "Standard term loan, fresh proposal"),
        ("Amount + tenure", "5 lakh for 36 months / five lakh for thirty six months"),
        ("Branch pincode", "Processing branch pincode 560034"),
        ("Proceed", "Proceed / next / looks good"),
        ("Submit", "Submit the application"),
    ]
    t4 = doc.add_table(rows=1, cols=2)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text = "Intent"
    t4.rows[0].cells[1].text = "Say this"
    for intent, phrase in lines:
        add_table_row(t4, [intent, f'"{phrase}"'])

    doc.add_paragraph()

    add_heading(doc, "Presenter tips", 2)
    tips = [
        "Keep the form visible behind the assistant — the live fill is the hero moment.",
        "Use one compound sentence in Act 2 to show multi-field inference; use step-by-step if the AI misses a field.",
        "Gold ring highlight = field just updated by the agent — point at it when it appears.",
        "If OpenAI key is missing, assistant will error — verify /api/health hasOpenAIKey: true.",
        "Demo disclaimer on review screen: “Demo only — not submitted to any loan system.”",
        "60-second cut: pre-open loan screen, say only amount + tenure + pincode after salaried/medical is set.",
    ]
    for tip in tips:
        add_bullet(doc, tip)

    add_heading(doc, "What this demo proves", 2)
    proofs = [
        "Voice-first loan origination — no manual dropdown hunting for every field.",
        "RM-style co-pilot: natural language mapped to exact form option ids live on screen.",
        "Multi-field extraction from a single utterance (occupation, product, purpose, amount, tenure).",
        "Guided recovery when user is stuck (idle prompt, rage-click help, request_field highlight).",
        "Full LOS journey: details → review → captured reference — end-to-end in ~2 minutes.",
    ]
    for p in proofs:
        add_bullet(doc, p)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
