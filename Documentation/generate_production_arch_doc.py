"""
Production Architecture Document — Enterprise edition
Silver Suits AI — AI Voice RM Platform
"""
import sys
sys.path.insert(0, '/tmp/docxlib')

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Enterprise palette ────────────────────────────────────────────────
BG       = '#F8FAFC'
LANE_HDR = '#0F172A'
TEXT     = '#0F172A'
SUB      = '#64748B'
ARROW    = '#334155'

LANES = {
    'user':   ('#E0F2FE', '#0369A1', 'CHANNEL LAYER'),
    'bank':   ('#DCFCE7', '#15803D', 'BANK INTEGRATION LAYER'),
    'core':   ('#EEF2FF', '#4338CA', 'SILVER SUITS AI PLATFORM'),
    'ai':     ('#F3E8FF', '#7C3AED', 'AI SERVICES (MANAGED)'),
    'rule':   ('#FFEDD5', '#C2410C', 'RULE ENGINE'),
    'data':   ('#FEE2E2', '#B91C1C', 'DATA & ANALYTICS'),
}


def lane(ax, y, h, W, key):
    fc, ec, label = LANES[key]
    ax.add_patch(Rectangle((0, y), W, h, fc=fc, ec='none', zorder=0))
    ax.add_patch(Rectangle((0, y), 0.55, h, fc=LANE_HDR, ec='none', zorder=1))
    ax.text(0.28, y + h / 2, label, ha='center', va='center',
            fontsize=6.5, fontweight='bold', color='white', rotation=90, zorder=2)


def comp(ax, x, y, w, h, title, sub=None, ec='#4338CA', fc='white', ts=8.5, ss=7):
    p = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0,rounding_size=0.12',
                        fc=fc, ec=ec, lw=1.8, zorder=3)
    ax.add_patch(p)
    cx = x + w / 2
    if sub:
        ax.text(cx, y + h * 0.65, title, ha='center', va='center',
                fontsize=ts, fontweight='bold', color=TEXT, zorder=4)
        ax.text(cx, y + h * 0.30, sub, ha='center', va='center',
                fontsize=ss, color=SUB, zorder=4)
    else:
        ax.text(cx, y + h / 2, title, ha='center', va='center',
                fontsize=ts, fontweight='bold', color=TEXT, zorder=4)


def arr(ax, x1, y1, x2, y2, label=None, color=ARROW, rad=0, dashed=False):
    ls = (0, (4, 3)) if dashed else '-'
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.6,
                                linestyle=ls,
                                connectionstyle=f'arc3,rad={rad}',
                                shrinkA=5, shrinkB=5), zorder=5)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my - 0.12, label, ha='center', va='top',
                fontsize=6.5, color=SUB, zorder=6,
                bbox=dict(fc=BG, ec='none', pad=1, alpha=0.95))


def fig_title(ax, W, title, subtitle, y=0.12):
    ax.text(W / 2, y, title, ha='center', fontsize=14, fontweight='bold', color=TEXT)
    ax.text(W / 2, y + 0.42, subtitle, ha='center', fontsize=9, color=SUB)


def legend(ax, W, y, items):
    for i, (col, lbl) in enumerate(items):
        xo = 0.7 + i * (W - 1.4) / len(items)
        ax.plot([xo, xo + 0.5], [y, y], color=col, lw=2.5)
        ax.text(xo + 0.65, y, lbl, fontsize=7, color=SUB, va='center')


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 1 — LAYERED SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════
def fig_layered_arch():
    W, H = 18, 13
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 1 — Production System Architecture (Layered View)',
              'Indian Bank VPC · AWS ap-south-1 (Mumbai) · DPDP / RBI compliant deployment')

    # Lane layout: (y_start, height, key)
    lanes = [
        (0.9,  1.35, 'user'),
        (2.35, 1.55, 'bank'),
        (4.0,  2.35, 'core'),
        (6.45, 1.75, 'ai'),
        (8.3,  1.55, 'rule'),
        (9.95, 1.55, 'data'),
    ]
    for y, h, key in lanes:
        lane(ax, y, h, W, key)

    cw, ch = 2.55, 0.72
    ox = 0.75

    # Channel
    comps_user = [
        ('Web Banking App', 'React · VoiceModal · AG-UI screens'),
        ('Mobile App (V3)', 'iOS / Android wrapper'),
        ('Agent Desktop (V2)', 'CRM overlay · RM assist'),
    ]
    for i, (t, s) in enumerate(comps_user):
        comp(ax, ox + i * 2.85, 1.05, cw, ch, t, s, ec=LANES['user'][1])

    # Bank integration
    comps_bank = [
        ('Event Gateway', 'POST /api/trigger'),
        ('Identity (SSO)', 'OAuth 2.0 · JWT'),
        ('Core Banking APIs', 'Accounts · Pay · KYC'),
        ('Webhook / Callback', 'Session outcomes'),
        ('Network Security', 'mTLS · IP allowlist'),
    ]
    for i, (t, s) in enumerate(comps_bank):
        comp(ax, ox + i * 2.85, 2.55, 2.35 if i == 4 else cw, ch, t, s, ec=LANES['bank'][1])

    # Silver Suits core — 2 rows
    row1 = [
        ('API Gateway', 'Auth · rate limit · WAF'),
        ('Session Manager', 'Redis · 30min TTL'),
        ('Dialogue Engine', 'State machine · saga'),
        ('AG-UI Agents', 'SSE · 5 screen agents'),
        ('Event Bus', 'Turn + trigger events'),
        ('Logger', 'Structured JSON logs'),
    ]
    row2 = [
        ('STT Proxy', '/api/stt'),
        ('TTS Proxy', '/api/tts'),
        ('Tool Registry', 'Bank tool bindings'),
        ('Bank API Adapter', 'Retry · circuit breaker'),
        ('MPIN Gate', 'Payment confirmation'),
        ('Load Balancer', 'ALB · auto-scale'),
    ]
    for i, (t, s) in enumerate(row1):
        comp(ax, ox + i * 2.85, 4.25, cw, ch, t, s, ec=LANES['core'][1])
    for i, (t, s) in enumerate(row2):
        comp(ax, ox + i * 2.85, 5.25, cw, ch, t, s, ec=LANES['core'][1])

    # AI services
    comps_ai = [
        ('ElevenLabs STT', 'Scribe v2 · India'),
        ('Azure OpenAI', 'GPT-4o · intent + reply'),
        ('Cartesia TTS', 'sonic-3 · MP3'),
        ('Azure OpenAI', 'AG-UI function calls'),
        ('Analysis LLM (V3)', 'Post-session insights'),
    ]
    for i, (t, s) in enumerate(comps_ai):
        ec = SUB if 'V3' in t else LANES['ai'][1]
        comp(ax, ox + i * 2.85, 6.75, cw, ch, t, s, ec=ec)

    # Rule engine
    comps_rule = [
        ('Trigger Ingestor', 'Schema validate'),
        ('Fact Composer', 'User + event + form'),
        ('Rule Evaluator', 'In-memory <100ms'),
        ('Action Dispatcher', '5 actuation types'),
        ('Rule Config Store', 'Bank-managed JSON'),
        ('Emergency Override', 'Compliance inject'),
    ]
    for i, (t, s) in enumerate(comps_rule):
        comp(ax, ox + i * 2.85, 8.55, cw, ch, t, s, ec=LANES['rule'][1])

    # Data
    comps_data = [
        ('DocumentDB', 'Sessions · turns · audit'),
        ('Redis Cache', 'Active session state'),
        ('S3', 'Rules · backups · audio opt.'),
        ('Analytics Dashboard', 'Read-only · bank ops'),
        ('CloudWatch', 'Metrics · alerts'),
    ]
    for i, (t, s) in enumerate(comps_data):
        comp(ax, ox + i * 2.85, 10.2, cw, ch, t, s, ec=LANES['data'][1])

    # Key vertical flows (center)
    cx = W / 2
    arr(ax, cx, 1.77, cx, 2.35, 'HTTPS')
    arr(ax, cx, 3.1, cx, 4.0, 'REST / JWT')
    arr(ax, cx, 6.35, cx, 6.45, 'API calls', color=LANES['ai'][1])
    arr(ax, cx, 8.2, cx, 8.3, 'events', color=LANES['rule'][1])
    arr(ax, cx, 9.85, cx, 9.95, 'persist', color=LANES['data'][1])

    # Bank → Rule (horizontal)
    arr(ax, 2.0, 3.1, 2.0, 8.55, 'trigger', color=LANES['bank'][1], rad=0.35)
    # Core → Bank APIs
    arr(ax, 12.5, 5.6, 8.0, 3.1, 'bank APIs', color=LANES['bank'][1], rad=-0.2)

    legend(ax, W, 11.75, [
        (LANES['bank'][1], 'Bank-owned'),
        (LANES['core'][1], 'Silver Suits platform'),
        (LANES['ai'][1], 'External AI (zero retention)'),
        (LANES['rule'][1], 'Rule engine'),
        (LANES['data'][1], 'Bank data plane'),
    ])

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 2 — RUNTIME SEQUENCE (Voice + Rule parallel)
# ═══════════════════════════════════════════════════════════════════════
def fig_runtime_flows():
    W, H = 18, 8.5
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 2 — Runtime Flows',
              'Left: Voice conversation path  |  Right: Rule engine path (can run in parallel)')

    # Swimlane headers
    actors_l = ['Customer', 'Bank App', 'AI Platform', 'AI Services', 'Bank APIs', 'Database']
    actors_r = ['Bank CBS/CRM', 'Rule Engine', 'AI Platform', 'Bank App', 'Database']
    aw = 2.5
    for i, a in enumerate(actors_l):
        comp(ax, 0.6 + i * aw, 1.0, aw - 0.15, 0.55, a, ec=LANES['user'][1] if i == 0 else LANES['core'][1], ts=8)
    ax.text(7.75, 0.75, 'VOICE CONVERSATION FLOW', ha='center', fontsize=10,
            fontweight='bold', color=LANES['core'][1])
    for i, a in enumerate(actors_r):
        comp(ax, 9.5 + i * aw, 1.0, aw - 0.15, 0.55, a, ec=LANES['rule'][1] if i == 1 else LANES['bank'][1], ts=8)
    ax.text(16.25, 0.75, 'RULE ENGINE FLOW', ha='center', fontsize=10,
            fontweight='bold', color=LANES['rule'][1])

    # Divider
    ax.plot([9.0, 9.0], [1.0, 7.8], color='#CBD5E1', lw=1.5, linestyle='--')

    # Voice flow steps (y positions)
    steps_l = [
        (0, 1, 2, '1. Speak'),
        (1, 2, 3, '2. Audio → STT'),
        (2, 3, 4, '3. LLM intent'),
        (3, 4, 5, '4. Bank action'),
        (4, 3, 2, '5. TTS reply'),
        (5, 5, 5, '6. Log turn'),
    ]
    y0 = 2.0
    for j, (from_i, to_i, via_i, lbl) in enumerate(steps_l):
        y = y0 + j * 0.85
        x1 = 0.6 + from_i * aw + aw / 2
        x2 = 0.6 + to_i * aw + aw / 2
        xv = 0.6 + via_i * aw + aw / 2 if via_i != to_i else None
        if xv and via_i != from_i:
            arr(ax, x1, y, xv, y, color=LANES['core'][1])
            arr(ax, xv, y, x2, y, lbl, color=LANES['ai'][1] if via_i == 3 else LANES['core'][1])
        else:
            arr(ax, x1, y, x2, y, lbl, color=LANES['core'][1])

    # Rule flow
    steps_r = [
        (0, 1, '1. Event POST'),
        (1, 2, '2. Build facts'),
        (2, 3, '3. Evaluate rules'),
        (3, 4, '4. Dispatch action'),
        (4, 5, '5. Log decision'),
    ]
    for j, (from_i, to_i, lbl) in enumerate(steps_r):
        y = y0 + j * 0.85
        x1 = 9.5 + from_i * aw + aw / 2
        x2 = 9.5 + to_i * aw + aw / 2
        arr(ax, x1, y, x2, y, lbl, color=LANES['rule'][1])

    # Timing note
    ax.text(7.75, 7.5, 'Typical voice turn: 1.5–2.5s (STT 400ms + LLM 600ms + TTS 400ms + bank API 200ms)',
            ha='center', fontsize=8, color=SUB,
            bbox=dict(boxstyle='round', facecolor='white', edgecolor='#CBD5E1'))
    ax.text(16.25, 7.5, 'Rule engine: <100ms trigger-to-action (in-memory evaluation)',
            ha='center', fontsize=8, color=SUB,
            bbox=dict(boxstyle='round', facecolor='white', edgecolor='#CBD5E1'))

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 3 — RULE ENGINE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════
def fig_rule_engine():
    W, H = 16, 9
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 3 — Rule Engine Architecture',
              'Event-driven · Bank-configured JSON rules · Sub-100ms evaluation · Full audit trail')

    # Pipeline boxes (horizontal, large)
    py = 2.2
    pw, ph = 2.6, 1.35
    pipeline = [
        ('① Trigger\nIngestor', 'POST /api/trigger\nSchema validation', LANES['bank'][1]),
        ('② Fact\nComposer', 'user + event + form\n+ session context', LANES['rule'][1]),
        ('③ Rule\nEvaluator', 'AND/OR conditions\nPriority resolution', LANES['rule'][1]),
        ('④ Action\nDispatcher', 'Route to actuation\nhandler', LANES['rule'][1]),
        ('⑤ Audit &\nCallback', 'Log + bank webhook', LANES['data'][1]),
    ]
    px0 = 0.8
    for i, (title, sub, ec) in enumerate(pipeline):
        comp(ax, px0 + i * 3.05, py, pw, ph, title, sub, ec=ec, ts=9, ss=7)
        if i < 4:
            arr(ax, px0 + i * 3.05 + pw, py + ph / 2,
                px0 + (i + 1) * 3.05, py + ph / 2, color=LANES['rule'][1])

    # Trigger sources (top)
    ax.text(W / 2, 1.35, 'TRIGGER SOURCES', ha='center', fontsize=9,
            fontweight='bold', color=SUB)
    triggers = [
        ('App Analytics', 'rage_click · invalid_field'),
        ('Bank CBS / CRM', 'churn_flag · large_txn'),
        ('Scheduler (V2)', 'EMI due · FD maturity'),
        ('Voice Turn Bus', 'Post-turn events'),
        ('Analytics Bus (V2)', 'Session outcome events'),
    ]
    for i, (t, s) in enumerate(triggers):
        comp(ax, 0.6 + i * 3.05, 1.55, 2.7, 0.65, t, s, ec=LANES['bank'][1], ts=7.5, ss=6.5)
        arr(ax, 0.6 + i * 3.05 + 1.35, 2.2, px0 + 1.3, py, color=LANES['bank'][1], rad=0.15)

    # Actions (bottom)
    ax.text(W / 2, 4.0, 'ACTUATION HANDLERS', ha='center', fontsize=9,
            fontweight='bold', color=SUB)
    actions = [
        ('popup', 'Open VoiceModal\n+ context prefill'),
        ('assign_rm', 'CRM ticket +\nRM notification'),
        ('push_notif', 'FCM / APNs (V2)'),
        ('form_prefill', 'AG-UI warm start'),
        ('escalate_queue', 'Human RM queue (V2)'),
    ]
    for i, (t, s) in enumerate(actions):
        ec = SUB if 'V2' in s else LANES['rule'][1]
        comp(ax, 0.6 + i * 3.05, 4.25, 2.7, 0.75, t, s, ec=ec, ts=8, ss=6.5)
        arr(ax, px0 + 3 * 3.05 + pw / 2, py + ph, 0.6 + i * 3.05 + 1.35, 4.25,
            color=LANES['rule'][1], rad=0.2)

    # Rule config side panel
    rb = FancyBboxPatch((12.0, 5.3), 3.5, 2.8, boxstyle='round,pad=0.1',
                         fc='white', ec=LANES['rule'][1], lw=2, zorder=3)
    ax.add_patch(rb)
    ax.text(13.75, 5.55, 'RULE CONFIG (Bank-Managed)', ha='center',
            fontsize=8.5, fontweight='bold', color=LANES['rule'][1], zorder=4)
    sample = (
        'rule_id · priority · trigger[]\n'
        'conditions: AND/OR tree\n'
        'action · action_config\n'
        'enabled · version\n'
        '─────────────\n'
        'Hot-reload < 5 sec\n'
        'CI schema validation\n'
        'Git-versioned deploy'
    )
    ax.text(12.2, 5.95, sample, ha='left', va='top', fontsize=7,
            color=TEXT, fontfamily='monospace', zorder=4, linespacing=1.4)
    arr(ax, 12.0, 6.7, 10.5, 2.9, 'load rules', color=LANES['rule'][1], dashed=True)

    # Emergency override
    comp(ax, 0.6, 5.3, 4.5, 0.9, 'Emergency Override Layer (V2)',
         'RBI directives · global AI pause · mandatory disclosures',
         ec='#D97706', fc='#FFFBEB', ts=8.5, ss=7)
    arr(ax, 2.85, 5.3, 4.5, 2.9, 'priority inject', color='#D97706', rad=-0.15)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 4 — DATA & DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════
def fig_data_deploy():
    W, H = 16, 8
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 4 — Data Architecture & AWS Deployment',
              'Bank-owned data plane · Silver Suits compute in bank VPC')

    # Left: Data model
    ax.text(4, 0.95, 'DATA MODEL (DocumentDB)', ha='center', fontsize=10,
            fontweight='bold', color=LANES['data'][1])
    collections = [
        ('sessions', 'id · user · outcome · lang · duration'),
        ('turns', 'utterance · intent · slots · latency'),
        ('rule_events', 'rule_id · facts · action'),
        ('users', 'segment · churn · preferences'),
        ('audit_log', 'immutable · 7yr retention'),
    ]
    for i, (name, fields) in enumerate(collections):
        comp(ax, 0.5, 1.3 + i * 1.05, 7, 0.85, name, fields, ec=LANES['data'][1], ts=8.5, ss=7)

    # Right: AWS deployment
    ax.text(12, 0.95, 'AWS DEPLOYMENT (ap-south-1)', ha='center', fontsize=10,
            fontweight='bold', color=LANES['core'][1])
    aws = [
        ('Internet / Bank WAN', 'HTTPS 443 only'),
        ('ALB + WAF', 'TLS 1.3 termination'),
        ('ECS Fargate Cluster', '3+ AI platform containers'),
        ('ElastiCache Redis', 'Session + override flags'),
        ('DocumentDB Cluster', 'Multi-AZ · encrypted'),
        ('S3 + KMS', 'Rules · logs · backups'),
        ('CloudWatch', 'Metrics · alarms · logs'),
    ]
    for i, (name, desc) in enumerate(aws):
        comp(ax, 8.5, 1.3 + i * 0.88, 7, 0.72, name, desc, ec=LANES['core'][1], ts=8.5, ss=7)

    # Analytics bar
    comp(ax, 0.5, 6.8, 15, 0.75, 'Analytics Dashboard (Read-Only)',
         'Queries DocumentDB only · No write path · Role-based access for bank operations team',
         ec=LANES['data'][1], fc='#FEF2F2', ts=9, ss=7.5)
    arr(ax, 4, 6.3, 4, 6.8, 'read', color=LANES['data'][1])
    arr(ax, 12, 6.3, 12, 6.8, 'read', color=LANES['data'][1])

    # Security strip
    sec = 'AES-256 at rest  ·  TLS 1.3 in transit  ·  mTLS bank integration  ·  Zero AI vendor retention  ·  India data residency'
    ax.text(W / 2, 7.75, sec, ha='center', fontsize=8, color='white', fontweight='bold',
            bbox=dict(boxstyle='round', facecolor=LANE_HDR, pad=0.5))

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


print('Generating diagrams...')
buf1 = fig_layered_arch()
buf2 = fig_runtime_flows()
buf3 = fig_rule_engine()
buf4 = fig_data_deploy()


# ═══════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════
NAVY  = RGBColor(0x0F, 0x17, 0x2A)
ACCENT = RGBColor(0x1E, 0x40, 0xAF)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x05, 0x76, 0x48)
ORANGE = RGBColor(0xC2, 0x41, 0x0C)


def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_colour)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def tbl(doc, headers, rows, widths=None, hdr='0F172A'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hr.cells[i], hdr)
        p = hr.cells[i].paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
    for idx, row in enumerate(rows):
        tr = t.add_row()
        for i, v in enumerate(row):
            if idx % 2 == 1:
                set_cell_bg(tr.cells[i], 'F1F5F9')
            p = tr.cells[i].paragraphs[0]
            p.clear()
            p.add_run(str(v)).font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            set_col_width(t, i, w)
    doc.add_paragraph()
    return t


def h1(doc, t):
    p = doc.add_heading(t, 1)
    if p.runs:
        p.runs[0].font.color.rgb = NAVY
        p.runs[0].font.size = Pt(18)


def h2(doc, t, c=None):
    p = doc.add_heading(t, 2)
    if p.runs:
        p.runs[0].font.color.rgb = c or ACCENT
        p.runs[0].font.size = Pt(13)


def h3(doc, t):
    p = doc.add_heading(t, 3)
    if p.runs:
        p.runs[0].font.color.rgb = NAVY
        p.runs[0].font.size = Pt(11)


def para(doc, t):
    p = doc.add_paragraph(t)
    if p.runs:
        p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def callout(doc, t):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'EEF2FF')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    r = p.add_run(t)
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(8)


def img(doc, buf, cap, w=6.8):
    doc.add_picture(buf, width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        c = doc.add_paragraph(cap)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.runs:
            c.runs[0].font.size = Pt(9)
            c.runs[0].italic = True
            c.runs[0].font.color.rgb = GREY
    doc.add_paragraph()


def pb(doc):
    doc.add_page_break()


# ── Setup ───────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.top_margin = sec.bottom_margin = Inches(0.7)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ═══ COVER ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SILVER SUITS AI')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = GREY

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('AI Voice RM Platform')
r2.bold = True
r2.font.size = Pt(24)
r2.font.color.rgb = NAVY

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Production Architecture\n& Rule Engine Specification')
r3.font.size = Pt(16)
r3.font.color.rgb = ACCENT

doc.add_paragraph()
tbl(doc, ['Document', 'Value'], [
    ['Client', 'Indian Bank'],
    ['Classification', 'Confidential — Internal / Partner Use'],
    ['Version', '3.0'],
    ['Date', 'May 2026'],
    ['Deployment', 'Bank VPC · AWS ap-south-1 (Mumbai)'],
    ['Compliance', 'DPDP Act 2023 · RBI Digital Lending · Zero AI retention'],
], widths=[4, 12])

pb(doc)

# ═══ 1. EXECUTIVE SUMMARY ════════════════════════════════════════════════
h1(doc, '1. Executive Summary')
para(doc,
     'The AI Voice RM Platform embeds an intelligent voice and screen assistant inside the bank\'s '
     'digital channels. It helps customers complete banking journeys by voice, detects frustration '
     'in real time, and executes bank-defined business rules to intervene with the right action — '
     'AI popup, human RM assignment, or notification.')

callout(doc,
        'Scope: In-app voice assistant (V1–V5). Push notifications and human RM escalation (V2). '
        'All customer data remains in bank-owned infrastructure in India.')

h2(doc, 'Solution at a Glance')
tbl(doc,
    ['Capability', 'Description', 'Owner'],
    [
        ['Voice banking', '5 flows: send money, balance, transfer, bill pay, flight (demo)', 'Silver Suits AI + Bank APIs'],
        ['Screen assistants', '5 AG-UI agents fill loan, IMPS, deposit, txn screens via voice', 'Silver Suits AI'],
        ['Behavior detection', 'Rage taps (5 in 900ms) and invalid field (2×) trigger help', 'Bank App + Platform'],
        ['Rule engine', 'Real-time IF/THEN on bank events; <100ms; JSON-configured rules', 'Silver Suits builds · Bank configures'],
        ['Persistence', 'Sessions, turns, rule decisions, audit — DocumentDB', 'Bank'],
        ['Analytics', 'Volume, outcomes, timelines, rule effectiveness', 'Bank'],
    ],
    widths=[3.5, 9, 4.5],
)

h2(doc, 'Architecture Principles')
for p in [
    'Bank owns customer data — Silver Suits is a data processor under DPDP.',
    'Deploy inside bank VPC — no public exposure of application tier.',
    'AI vendors operate with zero retention — ephemeral requests only, India regions.',
    'Rules are bank-managed — product team edits JSON; no code deploy for rule changes.',
    'Every decision is auditable — rule matches, actions, and API calls logged immutably.',
]:
    doc.add_paragraph(p, style='List Bullet')

pb(doc)

# ═══ 2. SYSTEM ARCHITECTURE ════════════════════════════════════════════
h1(doc, '2. System Architecture')
para(doc,
     'Figure 1 presents the layered production architecture. Each layer has a defined ownership '
     'boundary and communicates only through documented APIs. The platform supports three '
     'concurrent capabilities: voice dialogue, AG-UI screen agents, and the rule engine.')

img(doc, buf1, 'Figure 1 — Layered production architecture (6 layers, bank VPC deployment)', w=6.9)

h2(doc, '2.1 Layer Responsibilities')
tbl(doc,
    ['Layer', 'Components', 'Responsibility', 'Owner'],
    [
        ['Channel', 'Web app, mobile (V3), agent desktop (V2)', 'User interface, mic, rage detection, MPIN gate', 'Bank'],
        ['Bank Integration', 'SSO, CBS events, core APIs, webhooks, mTLS', 'Identity, banking operations, outcome callbacks', 'Bank'],
        ['AI Platform', 'Gateway, sessions, dialogue engine, AG-UI, proxies, event bus', 'Orchestrate AI flows, isolate API keys, scale compute', 'Silver Suits'],
        ['AI Services', 'ElevenLabs STT, Azure OpenAI, Cartesia TTS', 'Speech and language intelligence — zero retention', 'Silver Suits (contracted)'],
        ['Rule Engine', 'Ingestor, fact composer, evaluator, dispatcher, config', 'Evaluate bank rules and actuate interventions', 'Silver Suits · Bank configures'],
        ['Data & Analytics', 'DocumentDB, Redis, S3, dashboard, CloudWatch', 'Persist, query, monitor — read-only analytics', 'Bank'],
    ],
    widths=[2.2, 3.8, 5.5, 2.5],
)

h2(doc, '2.2 Runtime Flows')
para(doc,
     'Figure 2 separates the voice conversation path from the rule engine path. Both can execute '
     'in parallel — e.g. a rage-click trigger fires the rule engine while an active voice session continues.')

img(doc, buf2, 'Figure 2 — Runtime sequence: voice turn path and rule engine path', w=6.9)

h3(doc, 'Voice Turn Latency Budget')
tbl(doc,
    ['Step', 'Component', 'Target latency'],
    [
        ['Speech-to-text', 'ElevenLabs Scribe v2 (via proxy)', '300–500 ms'],
        ['Intent extraction', 'Azure OpenAI GPT-4o', '400–700 ms'],
        ['Bank tool execution', 'Bank API Adapter', '100–300 ms'],
        ['Reply generation', 'Azure OpenAI', 'included in LLM call'],
        ['Text-to-speech', 'Cartesia sonic-3 (via proxy)', '300–500 ms'],
        ['End-to-end per turn', 'Full pipeline', '1.5–2.5 s (p50)'],
    ],
    widths=[3.5, 5.5, 4],
)

pb(doc)

# ═══ 3. AI PLATFORM COMPONENTS ═══════════════════════════════════════════
h1(doc, '3. AI Platform — Component Specification')

h2(doc, '3.1 Core Services')
tbl(doc,
    ['Component', 'Technology', 'Interface', 'Function'],
    [
        ['API Gateway', 'Nginx + Express', 'HTTPS / JWT', 'Authentication, rate limiting (100 req/min/user), routing'],
        ['Session Manager', 'Redis ElastiCache', 'Internal', 'Active session state, 30-min TTL, multi-instance sync'],
        ['Dialogue Engine', 'Node.js engine.js + saga.js', 'POST /api/engine/turn', 'Manifest-driven state machine; 9 states; tool orchestration'],
        ['AG-UI Agents', 'Node.js + SSE', 'POST /api/agui/{id}', 'Streaming form-fill agents for 5 banking screens'],
        ['STT Proxy', 'Node.js → ElevenLabs', 'POST /api/stt', 'Audio in → transcript + language code'],
        ['TTS Proxy', 'Node.js → Cartesia', 'POST /api/tts', 'Text in → MP3 bytes'],
        ['Tool Registry', 'Node.js', 'Internal', 'Maps dialogue steps to bank API calls'],
        ['Bank API Adapter', 'Axios + Opossum', 'Bank REST', 'Retry, circuit breaker, timeout on bank integrations'],
        ['Event Bus', 'Redis Pub/Sub', 'Internal', 'Publishes turn events to rule engine subscribers'],
        ['Logger', 'Winston → CloudWatch + DocumentDB', 'Internal', 'Structured per-turn and per-session logs'],
    ],
    widths=[3, 2.8, 3.2, 8],
)

h2(doc, '3.2 Voice Flow Manifests')
tbl(doc,
    ['Manifest', 'Slots collected', 'Bank tools invoked', 'Confirmation'],
    [
        ['send_money', 'recipient, amount, VPA', 'contacts.search → execute_upi_payment', 'MPIN + yes/no'],
        ['check_balance', 'account type', 'balance.check', 'Verbal confirm'],
        ['internal_transfer', 'amount, from, to account', 'accounts.list → execute_transfer', 'MPIN'],
        ['pay_bill', 'biller, account, amount', 'billers.search → execute_bill_payment', 'MPIN'],
        ['book_flight', 'route, date, passenger', 'flights.search → execute_flight', 'Verbal confirm'],
    ],
    widths=[3, 4.5, 5, 4.5],
)

h2(doc, '3.3 AG-UI Screen Agents')
tbl(doc,
    ['Screen', 'Agent ID', 'Tools', 'Primary use'],
    [
        ['Home', 'indian_bank_home_assistant', 'navigate_to', 'Route by amount: UPI ≤₹1L, IMPS >₹1L'],
        ['Loan LOS', 'indian_bank_loan_los', 'set_field, validate_form', '10-field loan application by voice'],
        ['IMPS Transfer', 'indian_bank_imps_transfer', 'set_field', '9-field transfer form'],
        ['Create Deposit', 'indian_bank_create_deposit', 'set_field', 'FD / RD / MMD selection and fill'],
        ['Transaction History', 'indian_bank_txn_history', 'set_field, navigate_to', 'Query txns; navigate on fraud signal'],
    ],
    widths=[3, 4.5, 3.5, 6],
)

h2(doc, '3.4 Bank Integration APIs')
tbl(doc,
    ['API', 'Direction', 'Auth', 'Payload summary'],
    [
        ['POST /api/trigger', 'Bank → Platform', 'mTLS + API key', 'user_id, event_type, user_profile, form_profile, event_data'],
        ['POST /api/engine/turn', 'App → Platform', 'JWT', 'session_id, input{type, text|selection|confirm}'],
        ['POST /api/stt', 'App → Platform', 'JWT', 'audio blob (WebM/MP3/WAV)'],
        ['POST /api/tts', 'App → Platform', 'JWT', 'text, language code'],
        ['POST /api/agui/{agentId}', 'App → Platform', 'JWT', 'messages[], threadId — SSE response stream'],
        ['Webhook callback', 'Platform → Bank', 'mTLS', 'session_id, outcome, transcript_summary, slots, metrics'],
    ],
    widths=[3.5, 2.5, 2.5, 9.5],
)

pb(doc)

# ═══ 4. RULE ENGINE ═══════════════════════════════════════════════════════
h1(doc, '4. Rule Engine Architecture')
para(doc,
     'The rule engine is an event-driven decision service. Banks define rules as versioned JSON '
     'documents. On each trigger, the engine composes a fact object, evaluates conditions in '
     'memory, dispatches the highest-priority matching action, and writes an immutable audit record.')

img(doc, buf3, 'Figure 3 — Rule engine: ingest → compose → evaluate → dispatch → audit', w=6.9)

h2(doc, '4.1 Processing Pipeline', ORANGE)
tbl(doc,
    ['Stage', 'Input', 'Output', 'SLA'],
    [
        ['Trigger Ingestor', 'POST /api/trigger JSON', 'Validated event envelope', '< 10 ms'],
        ['Fact Composer', 'Event + user_profile + form_profile + Redis session', 'Flat fact object', '< 15 ms'],
        ['Rule Evaluator', 'Facts + rule config store', 'Matched rule_id + action', '< 50 ms'],
        ['Action Dispatcher', 'Matched rule + action_config', 'Side effect (popup, CRM, etc.)', '< 25 ms'],
        ['Audit & Callback', 'Decision record', 'DocumentDB write + bank webhook', '< 20 ms'],
        ['Total', '—', '—', '< 100 ms p99'],
    ],
    widths=[3, 4.5, 4, 2.5],
    hdr='C2410C',
)

h2(doc, '4.2 Trigger Catalog')
tbl(doc,
    ['event_type', 'Source', 'Typical use case', 'Release'],
    [
        ['rage_click', 'Bank app', 'User frustration — offer AI help', 'V1'],
        ['invalid_field', 'Bank app', 'Repeated form errors', 'V1'],
        ['churn_risk_flag', 'Bank CRM', 'Proactive retention outreach', 'V1'],
        ['large_txn_attempt', 'Bank CBS', 'High-value transaction guardrail', 'V1'],
        ['form_abandoned', 'Bank app', 'Loan/deposit form drop-off', 'V1'],
        ['loan_emi_due', 'Scheduler', 'Payment reminder + assist', 'V2'],
        ['high_churn_risk', 'Post-session analysis', 'Follow-up after negative session', 'V3'],
    ],
    widths=[3.2, 2.5, 5.8, 1.5],
)

h2(doc, '4.3 Actuation Matrix')
tbl(doc,
    ['Action', 'Mechanism', 'Customer experience', 'Release'],
    [
        ['popup', 'SSE push to app → VoiceModal', 'In-app AI assistant opens with context', 'V1'],
        ['assign_rm', 'POST bank CRM API', 'Human RM receives ticket + transcript summary', 'V1'],
        ['form_prefill', 'SET_CONTEXT on AG-UI agent', 'Target screen opens with fields pre-populated', 'V1'],
        ['push_notif', 'Bank notification service → FCM/APNs', 'Mobile alert with deep link', 'V2'],
    ],
    widths=[2.5, 4.5, 5.5, 1.5],
)

h2(doc, '4.4 Rule Configuration Schema')
p = doc.add_paragraph()
code = (
    '{\n'
    '  "rule_id": "R042",\n'
    '  "name": "Premium customer — rage intervention",\n'
    '  "priority": 10,\n'
    '  "enabled": true,\n'
    '  "trigger": ["rage_click", "invalid_field"],\n'
    '  "conditions": {\n'
    '    "AND": [\n'
    '      { "field": "user.segment", "op": "==", "value": "premium" },\n'
    '      { "field": "event.rage_count", "op": ">=", "value": 3 }\n'
    '    ]\n'
    '  },\n'
    '  "action": "popup",\n'
    '  "action_config": {\n'
    '    "message": "Need help with your application?",\n'
    '    "manifest": "loan_assist",\n'
    '    "prefill": { "screen": "loan_form" }\n'
    '  }\n'
    '}'
)
r = p.add_run(code)
r.font.name = 'Consolas'
r.font.size = Pt(8.5)
doc.add_paragraph()

h2(doc, '4.5 Condition Operators')
tbl(doc,
    ['Operator', 'Type', 'Example'],
    [
        ['==  !=', 'Equality', 'user.segment == "premium"'],
        ['>=  <=  >  <', 'Numeric / timestamp', 'event.rage_count >= 3'],
        ['CONTAINS', 'String / array', 'user.products CONTAINS "home_loan"'],
        ['IN  NOT_IN', 'Set membership', 'user.lang_pref IN ["hi","ta","te"]'],
        ['AND  OR  NOT', 'Logic combinators', '{"AND": [cond1, cond2]}'],
        ['COUNT (V2)', 'Event frequency', 'COUNT("rage_click", 3600000) >= 3'],
        ['WITHIN_HOURS (V2)', 'Recency', 'user.last_session_ts WITHIN_HOURS 24'],
    ],
    widths=[3.5, 3, 10.5],
)

h2(doc, '4.6 Emergency Override (V2)')
tbl(doc,
    ['Control', 'Effect', 'Authority'],
    [
        ['Global AI pause', 'All dialogue endpoints return 503', 'Bank CISO'],
        ['System prompt injection', 'Compliance text prepended to every LLM call', 'Bank compliance + Silver Suits SRE'],
        ['Rule blocklist', 'Named rules skipped at evaluation', 'Bank product / compliance'],
        ['Action blocklist', 'Action type disabled globally', 'Bank compliance'],
        ['Mandatory disclosure prefix', 'Prepended to every bot reply', 'Bank legal'],
    ],
    widths=[3.5, 5.5, 4],
)

pb(doc)

# ═══ 5. DATA & DEPLOYMENT ════════════════════════════════════════════════
h1(doc, '5. Data Architecture & Deployment')
img(doc, buf4, 'Figure 4 — DocumentDB data model and AWS deployment topology', w=6.9)

h2(doc, '5.1 Data Collections')
tbl(doc,
    ['Collection', 'Key fields', 'Written by', 'Read by'],
    [
        ['sessions', 'session_id, user_id, outcome, manifest, lang, turn_count, timestamps', 'Logger', 'Dashboard, replay (V2)'],
        ['turns', 'utterance, intent, slots, bot_reply, tool_calls, latency_ms{stt,llm,tts}', 'Logger', 'Dashboard, analysis (V3)'],
        ['rule_events', 'rule_id, fact_snapshot, action, action_config, timestamp', 'Rule engine', 'Dashboard, audit'],
        ['users', 'segment, lang_pref, churn_score, consent_ts, session_count', 'Platform', 'Rule enricher (V2)'],
        ['audit_log', 'actor, action, entity, hash — append-only', 'All services', 'Compliance, regulators'],
    ],
    widths=[2.5, 6.5, 2.5, 2.5],
)

h2(doc, '5.2 Analytics Dashboard Views')
tbl(doc,
    ['View', 'Metrics', 'Release'],
    [
        ['Operations overview', 'Session volume, success/fail/cancel rates, daily trend', 'V1'],
        ['Conversation explorer', 'Per-user timeline, turn-by-turn detail', 'V1'],
        ['Language & region', 'Distribution across en/hi/ta/te (+6 in V2)', 'V1'],
        ['Journey funnel', 'Step dropout by manifest and segment', 'V2'],
        ['Latency & reliability', 'p50/p95/p99 per service (STT, LLM, TTS)', 'V2'],
        ['Rule performance', 'Fire rate, action conversion, A/B results', 'V2'],
        ['Sentiment & churn', 'Post-session scores, cohort trends', 'V3'],
    ],
    widths=[4, 9, 1.5],
)

h2(doc, '5.3 AWS Infrastructure')
tbl(doc,
    ['Service', 'Purpose', 'Configuration'],
    [
        ['ALB + WAF', 'HTTPS entry, DDoS protection', 'TLS 1.3, access logs → S3'],
        ['ECS Fargate', 'AI platform containers', 'Min 3 tasks, CPU auto-scale, rolling deploy'],
        ['ElastiCache Redis', 'Sessions, pub/sub, override flags', 'Multi-AZ, encryption in transit + at rest'],
        ['DocumentDB', 'Persistent data', 'Multi-AZ, 35-day backup, TLS required'],
        ['S3 + KMS', 'Rule configs, artifacts, optional audio', 'SSE-KMS, versioning, cross-region DR (V3)'],
        ['CloudWatch', 'Monitoring', 'Dashboards: latency, error rate, rule eval count'],
        ['PrivateLink (V2)', 'Private Azure OpenAI connectivity', 'No public internet for LLM traffic'],
    ],
    widths=[3.5, 4.5, 9],
)

h2(doc, '5.4 Security & Compliance Controls')
tbl(doc,
    ['Control', 'Implementation', 'Regulation'],
    [
        ['Data residency', 'AWS ap-south-1 only; AI vendors India region', 'RBI IT Framework'],
        ['Encryption at rest', 'AES-256, bank-managed KMS', 'DPDP 2023'],
        ['Encryption in transit', 'TLS 1.3, mTLS for bank integration', 'DPDP 2023'],
        ['Zero AI retention', 'Contractual + technical ephemeral context', 'RBI Digital Lending'],
        ['Consent management', 'Logged before first AI session', 'DPDP Section 6'],
        ['Audit trail', '7-year immutable append-only log', 'RBI audit requirements'],
        ['Access control', 'RBAC on dashboard; JWT scoped API access', 'Internal security policy'],
    ],
    widths=[3.5, 5.5, 4],
)

pb(doc)

# ═══ 6. ROADMAP ═══════════════════════════════════════════════════════════
h1(doc, '6. Delivery Roadmap & Team')

h2(doc, '6.1 Release Phases')
tbl(doc,
    ['Phase', 'Timeline', 'Deliverables', 'Exit criteria'],
    [
        ['V1 — Production', 'Months 1–3',
         'VPC deploy, bank API integration, 4 languages, 5 voice flows, rule engine (popup + assign_rm), DocumentDB logging, basic dashboard, DPDP DPA',
         '500 concurrent sessions, p99 <2s voice turn, bank security sign-off'],
        ['V2 — Scale', 'Months 4–6',
         '10 languages, visual rule builder, emergency override UI, push notif, advanced analytics, agent desktop, PrivateLink',
         'Rule builder UAT, push-notif pilot, ISO 27001 gap report'],
        ['V3 — Intelligence', 'Months 7–12',
         'Post-session LLM analysis, native mobile SDK, churn analytics, conversation replay, ISO 27001 cert, multi-region DR',
         'Analysis accuracy >85%, ISO audit passed'],
    ],
    widths=[2.5, 2, 7.5, 5],
)

h2(doc, '6.2 Team Structure')
tbl(doc,
    ['Role', 'FTE', 'V1 responsibility'],
    [
        ['Full-Stack Engineer', '2', 'API platform, bank adapter, rule engine, frontend productionization'],
        ['DevOps / Cloud Engineer', '1', 'VPC, ECS, DocumentDB, Redis, CI/CD, KMS, monitoring'],
        ['AI / LLM Engineer', '1', 'Azure OpenAI, prompts, multilingual STT/TTS tuning'],
        ['QA Engineer', '1', 'Integration tests, load test, security scan, UAT support'],
        ['Solutions Architect', '0.5', 'Bank API contracts, integration design, sprint governance'],
        ['Compliance Consultant', '0.5', 'DPDP, data flow, DPA, audit preparation'],
    ],
    widths=[4.5, 1.5, 11],
)

# Footer
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run(
    'Silver Suits AI  ·  prateek@silversuits.ai  ·  +91 738 173 2333  ·  Mumbai | Bangalore\n'
    'Document Version 3.0  ·  Confidential  ·  May 2026'
)
rf.font.size = Pt(9)
rf.font.color.rgb = GREY

out = '/Users/devanshusaindane/SilverSuits/Voice-to-command/Voice-to-Command/Documentation/AI_Voice_RM_Production_Architecture.docx'
doc.save(out)
print(f'Saved → {out}')
