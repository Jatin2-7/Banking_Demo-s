#!/usr/bin/env python3
"""Generate AI Voice RM Developer Implementation Plan (HLD + LLD + Roadmap)."""
import sys
sys.path.insert(0, '/tmp/docxlib')

from pathlib import Path
import io
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from impl_diagrams import impl_figures

OUT = Path(__file__).parent / 'AI_Voice_RM_Developer_Implementation_Plan.docx'

DARK = RGBColor(0x0A, 0x1F, 0x3D)
ACCENT = RGBColor(0x10, 0x52, 0xAB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x05, 0x96, 0x69)
TEAL = RGBColor(0x04, 0x6B, 0x75)


def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_colour)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def header_row(table, labels, bg='0A1F3D'):
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)


def add_table_row(table, values, shade=False, tone=None):
    row = table.add_row()
    bg_map = {'green': 'E6F9F1', 'amber': 'FFF8E6', 'red': 'FDE8E8', 'blue': 'EBF2FF'}
    for i, val in enumerate(values):
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, 'F5F7FA')
        if tone and bg_map.get(tone):
            set_cell_bg(cell, bg_map[tone])
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(val))
        run.font.size = Pt(8.5)


def make_table(doc, headers, rows, col_widths=None, alt_shade=True, tones=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row(table, headers)
    for idx, r in enumerate(rows):
        tone = tones[idx] if tones and idx < len(tones) else None
        add_table_row(table, r, shade=(alt_shade and idx % 2 == 1), tone=tone)
    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(17)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = ACCENT
    p.runs[0].font.size = Pt(13)


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(11)


def body(doc, text, italic=False):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.size = Pt(10)
    if italic:
        run.italic = True


def callout(doc, text, tone='info'):
    bg = {'info': 'EBF2FF', 'warn': 'FFF8E6', 'ok': 'E6F9F1', 'err': 'FDE8E8'}[tone]
    table = doc.add_table(rows=1, cols=1)
    set_cell_bg(table.rows[0].cells[0], bg)
    p = table.rows[0].cells[0].paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    doc.add_paragraph()


def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)


def add_figure(doc, fig, caption, width=Inches(6.8)):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150, facecolor='#F8FAFC')
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True
    cap.runs[0].font.color.rgb = GREY
    doc.add_paragraph()


def divider(doc):
    doc.add_paragraph('─' * 80)


# ── Build document ────────────────────────────────────────────────────
print('Generating implementation diagrams...')
figs = impl_figures()

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Cover
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('AI Voice RM Platform')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Developer Implementation Plan\nHLD · LLD · Integration Guide · 36-Week Roadmap')
r2.font.size = Pt(14)
r2.font.color.rgb = ACCENT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta.add_run('Indian Bank × Silver Suits AI  |  Version 1.0  |  May 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = GREY

doc.add_page_break()

# TOC placeholder
h1(doc, 'Document Purpose')
body(doc,
     'This document is the engineering blueprint for building the AI Voice RM platform in production. '
     'The Voice-to-Command demo repository and architecture Word docs are reference implementations only. '
     'Use this plan to scope sprints, assign modules, and integrate with an existing banking Android app '
     'or internet-banking website.')
callout(doc,
        'Audience: backend engineers, mobile/web developers, DevOps, QA, and bank integration teams. '
        'Platform scope: in-app voice and screen assistants only (V1–V5). All interaction stays inside '
        'the bank mobile or web application.', tone='info')

h2(doc, 'How to use this document')
make_table(doc,
    headers=['Section', 'Use when you need to…'],
    rows=[
        ['§2 Reference vs Target', 'Understand what the demo already proves vs what must be built'],
        ['§3 AG-UI Integration', 'Wire form assistants into Android or web using the AG-UI protocol'],
        ['§4 Channel Integration', 'Choose WebView vs React Native vs native vs iframe vs NPM SDK'],
        ['§5 HLD', 'Explain system context, deployment, and service boundaries to architects'],
        ['§6 LLD', 'Implement APIs, data models, auth, and SDK modules'],
        ['§7 Security & Ops', 'Configure Azure, bank VPC, logging, and compliance controls'],
        ['§8 Roadmap', 'Plan sprints, owners, acceptance criteria, and dependencies'],
        ['§9 Testing & CI/CD', 'Define test pyramid and release gates'],
    ],
    col_widths=[4.0, 12.4],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '1. Reference Demo vs Production Target')
h2(doc, '1.1 What the demo proves (keep)')
make_table(doc,
    headers=['Capability', 'Demo location', 'Production action'],
    rows=[
        ['Manifest voice engine (UPI, bills, transfer)', 'server/engine/ + server/manifests/', 'Keep; add Redis sessions + bank adapters'],
        ['AG-UI form agents (5 journeys)', 'server/agui/* + client LoanAguiPanel', 'Keep; align events with AG-UI spec; add auth'],
        ['Frustration detection', 'client/src/hooks/useRageDetect.js', 'Extract to SDK; wire to Trigger API'],
        ['STT/TTS proxies', 'server/index.js /api/stt, /api/tts', 'Keep; India residency keys in Key Vault'],
        ['Mock banking REST', 'server/data/backend.js', 'Replace with bank core adapter layer'],
    ],
    col_widths=[4.8, 4.8, 7.0],
)

h2(doc, '1.2 What must be built (not in demo)')
make_table(doc,
    headers=['Gap', 'Priority', 'Owner', 'V1 deliverable'],
    rows=[
        ['Customer JWT auth + session binding', 'P0', 'Backend + Bank', 'Bearer token on every API call'],
        ['Redis session store (replace in-memory Map)', 'P0', 'Backend', 'Horizontal scale + 30 min TTL'],
        ['Bank DocumentDB writes (sessions, turns)', 'P0', 'Backend + Bank DBA', 'Async write pipeline'],
        ['POST /api/trigger + rule engine', 'P0', 'Backend', 'IF/THEN evaluator + audit log'],
        ['Server-side MPIN verify (replace client 1234)', 'P0', 'Bank + Backend', 'Step-up token before CONFIRMATION'],
        ['@silversuits/voice-rm-sdk npm package', 'P0', 'Frontend', 'Publishable embed layer'],
        ['Android WebView shell + JS bridge', 'P0', 'Bank mobile', 'Load hosted bundle with auth injection'],
        ['Analytics dashboard + post-session LLM', 'P1', 'Full stack', 'V3'],
        ['Visual rule-builder UI', 'P1', 'Frontend', 'V2'],
    ],
    col_widths=[5.0, 1.2, 2.4, 8.0],
    tones=[None, 'amber', None, None, 'amber', 'amber', 'amber', None, None],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '2. AG-UI Protocol — Research & Integration Guide')
body(doc,
     'AG-UI (Agent-User Interaction Protocol) is an open, event-based standard maintained by the '
     'ag-ui-protocol community (originally from CopilotKit). It defines how a frontend application '
     'connects to an AI agent backend using HTTP POST + Server-Sent Events (SSE). Official docs: '
     'https://docs.ag-ui.com')

h2(doc, '2.1 Protocol stack position')
make_table(doc,
    headers=['Protocol', 'Layer', 'Role in this platform'],
    rows=[
        ['MCP', 'Agent ↔ Tools', 'Future: connect to bank internal tools via MCP gateway (V3+)'],
        ['A2A', 'Agent ↔ Agent', 'Not required V1–V3; home router uses in-process navigate_to tool'],
        ['AG-UI', 'Agent ↔ User UI', 'Primary: form filling, routing, streaming chat in bank screens'],
    ],
    col_widths=[2.4, 3.2, 10.8],
)

h2(doc, '2.2 How the demo maps to AG-UI')
body(doc, 'The demo implements an AG-UI-compatible subset in client/src/lib/aguiClient.js and server/agui/*. '
          'It is not yet using the official @ag-ui/client npm package but emits compatible event shapes.')

make_table(doc,
    headers=['AG-UI standard event', 'Demo implementation', 'Notes'],
    rows=[
        ['RUN_STARTED', '✓ { thread_id, run_id }', 'Emitted at stream start'],
        ['TEXT_MESSAGE_START/CONTENT/END', 'Partial — TEXT_MESSAGE_CHUNK', 'Migrate to 3-phase in V2 for CopilotKit compat'],
        ['TOOL_CALL_START/ARGS/END', '✓ Full sequence', 'Used for set_field, navigate_to'],
        ['TOOL_CALL_RESULT', '✓ After tool exec', 'Client updates form via onToolCall'],
        ['STATE_DELTA', '✓ JSON Patch ops', 'applyStateDelta() in aguiClient.js'],
        ['STATE_SNAPSHOT', 'Not emitted', 'Add on reconnect / thread resume (V2)'],
        ['RUN_FINISHED / RUN_ERROR', '✓', 'LoanAguiPanel handles both'],
        ['MESSAGES_SNAPSHOT', 'Not emitted', 'Optional; client keeps local message array'],
    ],
    col_widths=[4.0, 4.4, 8.0],
)

h2(doc, '2.3 Standard request contract (RunAgentInput)')
code_block(doc, """POST /api/agui/{agentId}
Authorization: Bearer {customer_jwt}
Accept: text/event-stream
Content-Type: application/json

{
  "thread_id": "uuid-v4",
  "run_id": "uuid-v4",
  "state": { "beneficiary_name": "", "amount": "" },
  "messages": [
    { "id": "msg_1", "role": "user", "content": "Transfer 50000 to Rahul" }
  ],
  "tools": [],
  "context": [{ "type": "user_profile", "data": { "lang": "hi", "tier": "gold" } }],
  "forwarded_props": { "screen_id": "imps_transfer", "journey": "fund_transfer" }
}""")

h2(doc, '2.4 Registered agents (extend by adding runner + config)')
make_table(doc,
    headers=['agentId', 'Screen', 'Tools', 'Config file'],
    rows=[
        ['indian_bank_home_assistant', 'Home / global FAB', 'navigate_to, STATUS_UPDATE', 'homeAguiConfig.js'],
        ['indian_bank_imps_transfer', 'IMPS fund transfer form', 'set_field', 'impsAguiConfig.js'],
        ['indian_bank_loan_los', 'Loan application', 'set_field', 'loanAgentConfig.js'],
        ['indian_bank_deposit', 'FD/RD creation', 'set_field', 'depositAguiConfig.js'],
        ['indian_bank_txn_history', 'Statement / fraud advisory', 'set_field, explain_txn', 'txnHistoryAguiConfig.js'],
    ],
    col_widths=[4.2, 3.6, 3.6, 5.2],
)

h2(doc, '2.5 Migration path to official @ag-ui/client')
body(doc, 'Recommended for V2 after V1 ships with the custom client (lower risk). Steps:')
for step in [
    'Install @ag-ui/client in voice-rm-sdk package.',
    'Replace runAgent() in aguiClient.js with HttpAgent from @ag-ui/client.',
    'Map demo TEXT_MESSAGE_CHUNK → TEXT_MESSAGE_CONTENT in a thin adapter layer.',
    'For React Native: use expo/fetch (supports SSE streaming) — React Native global fetch does NOT stream.',
    'Optional: adopt CopilotKit React components for chat UI; keep LoanAguiPanel for form-specific UX.',
]:
    doc.add_paragraph(step, style='List Bullet')

callout(doc,
        'React Native AG-UI client is community status ("Help Wanted" per ag-ui-protocol repo). '
        'Production pattern: wrap HttpAgent with expo/fetch or react-native-sse. See GitHub issue #510, #1316.',
        tone='warn')

add_figure(doc, figs['agui_sequence'],
           'Figure 3 — AG-UI integration sequence for form assistant screens')

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '3. Channel Integration — Android & Web')
add_figure(doc, figs['integration_options'],
           'Figure 2 — Integration options for Android banking app and internet banking website')

h2(doc, '3.1 Recommended approach (V1 — fastest path)')
body(doc, 'Option A (Android WebView) + Option D (NPM SDK embed for web) minimizes rework of the '
          'existing React demo while meeting bank security requirements.')

h3(doc, '3.1.1 Android WebView integration (step-by-step)')
for step in [
    'Bank mobile team creates VoiceRmActivity (Kotlin) with a hardened WebView (no file access, JS enabled).',
    'Load https://voice-cdn.bank.in/embed/v1/index.html?screen=home from bank CDN (built from voice-rm-sdk).',
    'Inject auth before page load via @JavascriptInterface or evaluateJavascript:',
    '  window.__VOICE_RM_CONFIG__ = { apiBase, jwt, lang, userId, onNavigate: nativeBridge };',
    'Native bridge implements: onNavigate(dest), requestMpin(callback), logEvent(name, payload).',
    'WebView requests mic permission at runtime; pass grant result to JS via postMessage.',
    'On session end, native receives postMessage { type: "SESSION_END", outcome, sessionId } for bank analytics.',
]:
    doc.add_paragraph(step, style='List Number')

code_block(doc, """// Kotlin — minimal bridge
class VoiceRmBridge(private val activity: Activity) {
    @JavascriptInterface
    fun getAuthToken(): String = SessionManager.currentJwt

    @JavascriptInterface
    fun navigateTo(screen: String, contextJson: String) {
        activity.runOnUiThread { Router.open(activity, screen, contextJson) }
    }

    @JavascriptInterface
    fun verifyMpin(onResult: String) {
        MpinActivity.launch(activity) { ok ->
            webView.evaluateJavascript("window.__voiceRm.onMpinResult($ok)", null)
        }
    }
}""")

h3(doc, '3.1.2 Web / Internet banking NPM embed')
code_block(doc, """// Bank React app — UPI route
import { VoiceRmProvider, VoiceModal, useVoiceSession } from '@silversuits/voice-rm-sdk';

function UpiPaymentPage() {
  const { session, sendUtterance } = useVoiceSession({ journey: 'send_money', lang: 'hi' });
  return (
    <VoiceRmProvider apiBase={process.env.VOICE_RM_API} getToken={() => bankAuth.getJwt()}>
      <ExistingUpiForm />
      <VoiceModal session={session} onUtterance={sendUtterance} />
    </VoiceRmProvider>
  );
}""")

h2(doc, '3.2 Integration decision matrix')
make_table(doc,
    headers=['Option', 'Time to V1', 'Reuse demo code', 'Native UX', 'AG-UI SSE', 'Bank effort'],
    rows=[
        ['A · WebView shell', '4–6 wks', '95%', 'Good', 'Via WebView fetch', 'Low — 1 Android dev'],
        ['B · React Native module', '8–12 wks', '80%', 'Excellent', 'expo/fetch required', 'Medium — RN team'],
        ['C · Native Kotlin UI', '12–16 wks', '20%', 'Excellent', 'Java/Dart AG-UI SDK', 'High — full rewrite'],
        ['D · NPM SDK (web)', '3–4 wks', '90%', 'N/A (web)', 'Native fetch SSE', 'Low — frontend embed'],
        ['E · Iframe widget', '2–3 wks', '95%', 'N/A (web)', 'Same as D', 'Lowest — postMessage only'],
        ['F · CopilotKit + @ag-ui/client', '6–8 wks', '70%', 'Good', 'Official client', 'Medium — new deps'],
    ],
    col_widths=[3.2, 2.0, 2.4, 2.0, 2.8, 3.2],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '4. High-Level Design (HLD)')
add_figure(doc, figs['hld_deployment'],
           'Figure 1 — Production deployment topology')

h2(doc, '4.1 System context')
body(doc, 'The platform sits between bank customer channels and bank core systems. Silver Suits hosts '
          'the voice/AG-UI API layer in Azure India. The bank hosts customer-facing apps and owns all '
          'persistent customer data in DocumentDB / on-prem databases.')

h2(doc, '4.2 Logical service boundaries')
make_table(doc,
    headers=['Service', 'Responsibility', 'Scaling', 'State'],
    rows=[
        ['voice-rm-gateway', 'TLS termination, JWT validation, rate limits, routing', 'Azure APIM + App Service', 'Stateless'],
        ['voice-engine', 'Manifest saga, LLM extract, confirm parser', 'Horizontally scaled pods', 'Session in Redis'],
        ['agui-service', 'SSE agent runners, tool loop, STATE_DELTA', 'Same pods as engine', 'Thread state ephemeral per run'],
        ['trigger-service', 'Rule evaluation, actuation dispatch', 'Dedicated worker + queue', 'Rules in config store'],
        ['speech-proxy', 'STT/TTS key proxy, PII redaction in logs', 'Co-located with gateway', 'Stateless'],
        ['bank-adapter', 'Core banking API client, idempotency, circuit breaker', 'Sidecar or module', 'Connection pools'],
        ['analytics-worker', 'Post-session LLM analysis (V3)', 'Queue consumer', 'Writes to bank DB'],
    ],
    col_widths=[3.2, 5.6, 3.6, 3.2],
)

h2(doc, '4.3 Deployment zones')
make_table(doc,
    headers=['Zone', 'Components', 'Connectivity'],
    rows=[
        ['Bank DMZ / Mobile', 'Android app, WebView bundle, IB frontend', 'HTTPS to APIM only'],
        ['Silver Suits Azure (India Central)', 'App Service, Redis, Key Vault, App Insights', 'Private endpoint to bank DB'],
        ['Bank VPC (Mumbai)', 'DocumentDB, core APIs, CRM webhook', 'mTLS from Azure via Private Link'],
        ['AI providers (India residency)', 'ElevenLabs STT, Azure OpenAI, Cartesia TTS', 'Outbound HTTPS from Azure'],
    ],
    col_widths=[3.6, 5.6, 7.4],
)

add_figure(doc, figs['component_map'],
           'Figure 4 — Component / package map for engineering teams')

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '5. Low-Level Design (LLD)')
add_figure(doc, figs['voice_sequence'],
           'Figure 6 — Voice engine turn sequence with bank MPIN gate')

h2(doc, '5.1 Package structure — @silversuits/voice-rm-sdk')
code_block(doc, """@silversuits/voice-rm-sdk/
├── src/
│   ├── index.ts                    # public exports
│   ├── provider/VoiceRmProvider.tsx  # auth + apiBase context
│   ├── voice/
│   │   ├── VoiceModal.tsx          # from demo (refactored)
│   │   ├── useVoiceSession.ts      # wraps engineClient
│   │   └── engineClient.ts
│   ├── agui/
│   │   ├── LoanAguiPanel.tsx
│   │   ├── aguiClient.ts           # → @ag-ui/client wrapper (V2)
│   │   └── applyStateDelta.ts
│   ├── behavior/
│   │   ├── useRageDetect.ts
│   │   └── RMHelpPrompt.tsx
│   ├── speech/
│   │   ├── useElevenSpeech.ts
│   │   └── cartesiaTts.ts
│   └── bridge/
│       └── nativeBridge.ts         # WebView postMessage adapter
├── package.json
└── rollup.config.js                # ESM + CJS; peerDeps: react>=18""")

h2(doc, '5.2 Voice Engine API — POST /api/engine/turn')
make_table(doc,
    headers=['input.type', 'When', 'Required fields', 'Response session.state'],
    rows=[
        ['INIT', 'Session create', 'lang?', 'IDLE → greeting in history'],
        ['START_ACTION', 'Open journey', 'action (manifest id)', 'FILL — first slot prompt'],
        ['TRANSCRIPT', 'User spoke/typed', 'text', 'FILL | DISAMBIGUATE | CHOOSE | CONFIRM'],
        ['SELECTION', 'User picked option', 'optionId', 'Advances saga'],
        ['CONFIRMATION', 'After MPIN OK', 'confirmed: boolean', 'EXECUTING → DONE | FAILED'],
        ['CANCEL', 'User aborts', '—', 'CANCELLED'],
        ['SET_LANG', 'Language switch', 'lang', 'Re-prompts current slot in new lang'],
    ],
    col_widths=[2.4, 2.8, 3.6, 7.8],
)

code_block(doc, """// Request
POST /api/engine/turn
Authorization: Bearer eyJ...
X-Idempotency-Key: {uuid}          // required on CONFIRMATION
X-Bank-Session-Id: {bank_sid}      // bank mobile session correlation

{
  "sessionId": "sess_abc123",        // omit on first call
  "lang": "hi",
  "input": { "type": "TRANSCRIPT", "text": "Rahul ko paanch sau bhejo" }
}

// Response (client renders verbatim)
{
  "ok": true,
  "sessionId": "sess_abc123",
  "session": {
    "state": "FILL",
    "action": "send_money",
    "pending": { "kind": "free_text", "slot": "amount", "prompt": "..." },
    "history": [...],
    "executing": false
  }
}""")

h2(doc, '5.3 Trigger API — POST /api/trigger (production build)')
code_block(doc, """POST /api/trigger
Authorization: Bearer {service_token}   // bank backend calls this
Content-Type: application/json

{
  "event_type": "rage_click",
  "user_id": "CIF123456",
  "screen_id": "imps_transfer",
  "field_context": { "invalid_field": "ifsc_code", "attempts": 2 },
  "lang": "hi",
  "device": { "platform": "android", "app_version": "4.2.1" }
}

// Response
{
  "matched_rule_id": "rule_frustrated_imps",
  "action": "trigger_popup",
  "action_config": {
    "script_id": "imps_assist",
    "agent_id": "indian_bank_imps_transfer",
    "visibility": "modal"
  }
}""")

h2(doc, '5.4 Rule engine actuation handlers (in-app only)')
make_table(doc,
    headers=['action', 'Handler', 'Side effect'],
    rows=[
        ['trigger_popup', 'WebSocket / SSE push to client SDK', 'Opens VoiceModal or AguiPanel with script_id'],
        ['assign_human_rm', 'POST bank CRM webhook', 'Creates ticket; RM notified'],
        ['push_notif', 'FCM/APNs via bank push gateway', 'Re-engagement message with deep link'],
        ['form_prefill', 'Push STATE_DELTA to active AguiPanel thread', 'Warm-starts form with context'],
    ],
    col_widths=[2.8, 4.8, 8.0],
)

h2(doc, '5.5 Data model — DocumentDB collections')
make_table(doc,
    headers=['Collection', 'Key fields', 'Written when', 'Retention'],
    rows=[
        ['sessions', 'session_id, user_id, journey, outcome, started_at, ended_at, lang', 'Session start/end', '7 years (bank policy)'],
        ['turns', 'session_id, turn_idx, role, text, intent, slots, latency_ms', 'Each engine turn', '7 years'],
        ['form_events', 'session_id, field_id, value_hash, valid', 'AG-UI set_field', '7 years'],
        ['rule_audit', 'rule_id, user_id, event_type, action, ts', 'Rule match', '7 years'],
        ['escalations', 'user_id, rm_id, priority, reason, ts', 'assign_human_rm', '7 years'],
        ['analytics', 'session_id, sentiment_arc[], churn_risk, qa_score, intent_category', 'Post-session LLM (V3)', '7 years'],
    ],
    col_widths=[2.4, 5.6, 3.6, 3.0],
)

h2(doc, '5.6 Auth & session binding')
make_table(doc,
    headers=['Token', 'Issuer', 'Claims', 'Used on'],
    rows=[
        ['Customer JWT', 'Bank IAM', 'sub (CIF), sid, lang, tier, exp (15m)', '/engine/turn, /agui/*, /stt, /tts'],
        ['Service token', 'Bank API GW', 'client_id, scope:trigger:write', 'POST /api/trigger only'],
        ['Step-up token', 'Bank MPIN service', 'sub, action, exp (2m), jti', 'CONFIRMATION turns only'],
    ],
    col_widths=[2.4, 2.8, 5.6, 5.8],
)

h2(doc, '5.7 Rage detection — bank embed contract')
body(doc, 'Port useRageDetect.js unchanged. Bank wraps each high-friction screen root view:')
code_block(doc, """const { containerProps, markInvalidField, dismiss } = useRageDetect({
  onFrustrated: () => {
    bankAnalytics.track('rage_detected', { screen: 'imps_transfer' });
    // Option 1: open in-app help immediately
    setHelpOpen(true);
    // Option 2: call bank backend → POST /api/trigger (async rule eval)
    bankApi.postTrigger({ event_type: 'rage_click', screen_id: 'imps_transfer' });
  },
});

return <div {...containerProps}><ImpsForm onInvalid={(f) => markInvalidField(f.id)} /></div>;""")

h2(doc, '5.8 Bank adapter interface (replace mock backend)')
code_block(doc, """// server/adapters/bankCore.js — implement per bank
export const bankAdapter = {
  async searchContacts(userId, query) { /* GET bank core */ },
  async getAccounts(userId) { /* balances */ },
  async sendUpi(userId, payload, idempotencyKey) { /* NPCI / bank UPI API */ },
  async internalTransfer(userId, payload, idempotencyKey) { /* CBS */ },
  async verifyMpin(userId, mpinHash) { /* bank auth service */ },
};""")

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '6. Security, Observability & Compliance')
h2(doc, '6.1 Security controls (V1 minimum)')
for item in [
    'All API endpoints require Bearer JWT; reject missing/expired with 401.',
    'Rate limit: 60 engine turns / user / hour; 20 AG-UI runs / user / hour.',
    'STT/TTS/LLM API keys in Azure Key Vault; never in client bundle.',
    'PII redaction in logs: mask account numbers, VPAs, phone numbers (server/lib/log.js pattern).',
    'CORS: allowlist bank domains only (no wildcard in production).',
    'Consent gate: first session turn blocked until consent_accepted=true in JWT or explicit API call.',
    'Idempotency-Key mandatory on CONFIRMATION and all payment tool executions.',
]:
    doc.add_paragraph(item, style='List Bullet')

h2(doc, '6.2 Observability')
make_table(doc,
    headers=['Signal', 'Tool', 'Key metrics'],
    rows=[
        ['Structured logs', 'Pino → App Insights', 'turn_latency_ms, intent, outcome, error_code'],
        ['Traces', 'OpenTelemetry', 'STT → LLM → TTS waterfall per turn'],
        ['Dashboards', 'Grafana / Azure Monitor', 'sessions/day, escalation rate, STT empty rate'],
        ['Alerts', 'PagerDuty', 'error rate > 2%, p95 latency > 3s, Redis down'],
    ],
    col_widths=[3.0, 4.0, 9.6],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '7. Development Roadmap — 36 Weeks')
add_figure(doc, figs['sprint_roadmap'],
           'Figure 5 — 36-week development roadmap (V1–V3)')

h2(doc, '7.1 V1 MVP — Weeks 1–8 (Go-live criteria)')
make_table(doc,
    headers=['Sprint', 'Deliverable', 'Owner', 'Acceptance criteria'],
    rows=[
        ['1–2', 'Azure infra + CI/CD + APIM + Key Vault', 'DevOps', 'Staging URL live; secrets rotated; TLS 1.2+'],
        ['1–2', 'JWT middleware + CORS allowlist', 'Backend', '401 without token; bank test JWT accepted'],
        ['3–4', 'Redis session store replaces in-memory', 'Backend', '2 pod load test; session survives pod restart'],
        ['3–4', 'voice-rm-sdk v0.1 published (internal npm)', 'Frontend', 'VoiceModal + AguiPanel importable in sample app'],
        ['3–4', 'Android WebView shell (bank team)', 'Bank mobile', 'Demo journey loads; mic works; bridge callbacks fire'],
        ['5–6', 'DocumentDB write pipeline', 'Backend + Bank DBA', 'sessions + turns persisted; bank can query'],
        ['5–6', 'POST /api/trigger + 5 seed rules', 'Backend', 'rage_click → trigger_popup in < 100ms eval'],
        ['5–6', 'Bank adapter v1 (accounts, contacts, UPI mock→real)', 'Backend + Bank', 'send_money executes against bank UAT API'],
        ['7–8', 'Server-side MPIN gate', 'Bank + Backend', 'CONFIRMATION rejected without step-up token'],
        ['7–8', 'UAT + VAPT + consent flow', 'QA + All', 'Bank sign-off; zero P0 bugs; consent logged'],
    ],
    col_widths=[1.4, 4.8, 2.4, 7.8],
)

h2(doc, '7.2 V2 Production — Weeks 9–20')
make_table(doc,
    headers=['Sprint', 'Deliverable', 'Acceptance criteria'],
    rows=[
        ['9–10', 'Push notification actuation (FCM/APNs)', 'Rule match fires push; deep link opens correct screen'],
        ['11–12', 'Visual rule-builder UI', 'Bank ops can edit JSON rules without deploy'],
        ['11–12', '6 additional languages (10 total)', 'STT + manifests + TTS for gu/kn/pa/mr/ml/bn'],
        ['13–14', 'Emergency override dashboard', 'Ops can disable rule in < 60s; audit logged'],
        ['15–16', 'Analytics dashboard v1', 'Volume, outcomes, branch-wise, language charts live'],
        ['17–18', 'Load test 1,000 concurrent sessions', 'p95 turn latency < 1.5s; zero session loss'],
        ['19–20', 'Migrate aguiClient → @ag-ui/client', 'All 5 agents pass AG-UI compatibility test suite'],
    ],
    col_widths=[1.4, 5.6, 9.6],
)

h2(doc, '7.3 V3 Intelligence — Weeks 21–36')
make_table(doc,
    headers=['Sprint', 'Deliverable', 'Acceptance criteria'],
    rows=[
        ['21–24', 'Post-session LLM analysis worker', 'sentiment_arc, intent_category, qa_score per session'],
        ['25–28', 'Churn risk scoring + predictive triggers', 'Rules can match on churn_risk_prev > threshold'],
        ['29–32', 'Full analytics dashboard', 'Conversation replay, multi-journey funnel, cohort views'],
        ['33–36', 'Multi-region DR + ISO 27001 audit prep', 'Failover tested; audit evidence pack complete'],
    ],
    col_widths=[1.4, 5.6, 9.6],
)

h2(doc, '7.4 Team allocation')
make_table(doc,
    headers=['Role', 'FTE', 'V1 focus', 'V2–V3 focus'],
    rows=[
        ['Backend Engineer', '2', 'Engine prod, Redis, trigger API, bank adapter', 'Rule builder API, analytics worker'],
        ['Frontend / SDK Engineer', '1', 'voice-rm-sdk extract, WebView bundle', 'Rule builder UI, dashboard'],
        ['AI / Conversation Engineer', '1', 'Manifest tuning, AG-UI prompts, eval harness', 'Post-session analysis pipeline'],
        ['DevOps', '1', 'Azure, CI/CD, APIM, monitoring', 'DR, load test infra'],
        ['QA Engineer', '1', 'Scripted flows, UAT scripts, regression', 'Load test, security regression'],
        ['Bank Mobile Dev', '1 (bank)', 'WebView shell, bridge, MPIN', 'Push deep links, native mic polish'],
        ['Bank Backend Dev', '1 (bank)', 'JWT issuer, trigger wiring, DB schema', 'CRM webhook, core API UAT'],
        ['Product Manager', '1', 'Sprint planning, bank stakeholder sync', 'Analytics requirements, rule-builder UX'],
    ],
    col_widths=[3.2, 1.0, 5.6, 5.8],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '8. Testing Strategy & CI/CD')
h2(doc, '8.1 Test pyramid')
make_table(doc,
    headers=['Layer', 'What', 'Tooling', 'Gate'],
    rows=[
        ['Unit', 'confirmParser, applyStateDelta, rule eval, expr.js', 'Vitest / Node test', 'Every PR'],
        ['Integration', 'Scripted engine flows (existing tests/)', 'runScriptedFlows.test.js + LLM stub', 'Every PR'],
        ['Contract', 'OpenAPI schema for /engine/turn, /agui/*, /trigger', 'Schemathesis / Dredd', 'Weekly'],
        ['E2E', 'WebView journey: rage → help → UPI payment', 'Playwright + device farm', 'Pre-release'],
        ['Load', '1K concurrent AG-UI + engine sessions', 'k6 / Azure Load Testing', 'V2 week 17–18'],
        ['Security', 'VAPT, dependency scan, secret scan', 'Bank-approved vendor + GitHub Advanced Security', 'Pre go-live'],
    ],
    col_widths=[2.0, 5.6, 4.0, 4.0],
)

h2(doc, '8.2 CI/CD pipeline (Azure DevOps / GitHub Actions)')
for step in [
    'PR → lint + manifest AJV validation + unit tests + scripted flow tests (LLM stubbed).',
    'Merge to main → build Docker image → push to ACR → deploy to staging.',
    'Staging smoke: /api/health, one scripted send_money flow, one AG-UI loan run.',
    'Manual approval → deploy to production (blue/green via App Service slots).',
    'voice-rm-sdk: separate npm publish pipeline on tag v*.*.*.',
]:
    doc.add_paragraph(step, style='List Number')

h2(doc, '8.3 Demo → production migration checklist')
make_table(doc,
    headers=['Demo artifact', 'Production replacement', 'Done?'],
    rows=[
        ['server/engine/session.js in-memory Map', 'Redis session store + serialization', ''],
        ['server/data/backend.js mock', 'bankAdapter → real core APIs', ''],
        ['client MpinSheet hardcoded 1234', 'Bank MPIN service + step-up JWT', ''],
        ['cors() open', 'APIM + allowlist', ''],
        ['VITE_API_BASE inconsistency', 'Single SDK config: apiBase + getToken()', ''],
        ['aguiClient.js custom SSE', '@ag-ui/client HttpAgent (V2)', ''],
        ['No /api/trigger', 'trigger-service + rule config store', ''],
        ['DemoPanel.jsx', 'Remove from production bundle', ''],
    ],
    col_widths=[4.8, 5.6, 2.0],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
h1(doc, '9. Open Decisions Log')
make_table(doc,
    headers=['Decision', 'Options', 'Recommendation', 'Decide by'],
    rows=[
        ['Android integration pattern', 'WebView vs RN vs Native', 'WebView for V1 (Option A)', 'Week 1'],
        ['AG-UI client library', 'Custom vs @ag-ui/client', 'Custom V1; migrate V2', 'Week 9'],
        ['Session store', 'Redis vs Cosmos DB', 'Redis (lower latency for turns)', 'Week 2'],
        ['Rule config storage', 'DocumentDB vs Git-backed JSON', 'DocumentDB with admin UI', 'Week 11'],
        ['WebSocket vs SSE for trigger_popup', 'Both viable', 'SSE (already used for AG-UI)', 'Week 5'],
        ['Bank DB access pattern', 'Private Link vs VPN', 'Private Link (Azure ↔ AWS Mumbai)', 'Week 2'],
    ],
    col_widths=[3.6, 4.0, 4.8, 2.2],
)

# Footer
doc.add_page_break()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run(
    'AI Voice RM Developer Implementation Plan v1.0  ·  Silver Suits AI  ·  Confidential\n'
    'Reference repo: Voice-to-Command (demo)  ·  Architecture docs: Documentation/*.docx'
)
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

doc.save(str(OUT))
print(f'Saved → {OUT}')
