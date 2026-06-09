import sys
sys.path.insert(0, '/tmp/docxlib')

from pathlib import Path
from arch_diagrams import demo_figures, prod_figures

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colours ──────────────────────────────────────────────────────────
DARK   = RGBColor(0x0A, 0x1F, 0x3D)
ACCENT = RGBColor(0x10, 0x52, 0xAB)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x6B, 0x72, 0x80)
GREEN  = RGBColor(0x05, 0x96, 0x69)
AMBER  = RGBColor(0xB4, 0x58, 0x09)
RED    = RGBColor(0xAA, 0x1C, 0x1C)
TEAL   = RGBColor(0x04, 0x6B, 0x75)

# ── helpers ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_colour)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def header_row(table, labels, bg='0A1F3D'):
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)

def add_table_row(table, values, shade=False, tone=None):
    row = table.add_row()
    bg_map = {'green': 'E6F9F1', 'amber': 'FFF8E6', 'red': 'FDE8E8', 'blue': 'EBF2FF'}
    for i, val in enumerate(values):
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, 'F5F7FA')
        if tone and bg_map.get(tone):
            set_cell_bg(cell, bg_map[tone])
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(val))
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)

def make_table(doc, headers, rows, col_widths=None, alt_shade=True, header_bg='0A1F3D', tones=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row(table, headers, bg=header_bg)
    for idx, r in enumerate(rows):
        tone = (tones[idx] if tones and idx < len(tones) else None)
        add_table_row(table, r, shade=(alt_shade and idx % 2 == 1), tone=tone)
    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)
    doc.add_paragraph()
    return table

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(17)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)

def h2(doc, text, colour=None):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = colour or ACCENT
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)

def h3(doc, text, colour=None):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = colour or DARK
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def body(doc, text, italic=False, colour=None, size=10):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.size = Pt(size)
    if italic: run.italic = True
    if colour: run.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(5)
    return p

def label_body(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ': ')
    r1.bold = True; r1.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(3)

def callout(doc, text, tone='info'):
    bg = {'info': 'EBF2FF', 'warn': 'FFF8E6', 'ok': 'E6F9F1', 'err': 'FDE8E8'}[tone]
    fg = {'info': ACCENT, 'warn': AMBER, 'ok': GREEN, 'err': RED}[tone]
    prefix = {'info': 'NOTE', 'warn': 'IMPORTANT', 'ok': 'IN DEMO CODEBASE', 'err': 'NOT IN DEMO'}[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(8)
    r1 = p.add_run(f'  {prefix}  ')
    r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = fg
    r2 = p.add_run('  ' + text)
    r2.font.size = Pt(9); r2.font.color.rgb = GREY

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9.5)

def code_line(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(8.5)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def insert_figure(doc, buf, caption, width_inches=6.4):
    doc.add_picture(buf, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = GREY
    doc.add_paragraph()

print('Generating architecture diagrams...')
(DEMO_FIG1, DEMO_FIG2, DEMO_FIG3, DEMO_FIG4) = demo_figures()
(PROD_FIG1, PROD_FIG2, PROD_FIG3, PROD_FIG4) = prod_figures()
print('Diagrams ready.')

# ═══════════════════════════════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(48)
r = cover.add_run('SILVER SUITS AI')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('AI Coworkers for BFSI Segments to Scale')
r2.font.size = Pt(11); r2.font.color.rgb = GREY; r2.italic = True

doc.add_paragraph()
t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t1.add_run('AI Voice RM Platform')
r3.bold = True; r3.font.size = Pt(22); r3.font.color.rgb = ACCENT

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t2.add_run('Technical Architecture for Productization')
r4.font.size = Pt(15); r4.font.color.rgb = DARK

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = t3.add_run('Integrated: Demo Codebase + Production Target + Roadmap')
r5.font.size = Pt(11); r5.font.color.rgb = GREY; r5.italic = True

doc.add_paragraph()
doc.add_paragraph()

meta_items = [
    ('Client',     'Indian Bank'),
    ('Provider',   'Silver Suits AI'),
    ('Version',    '3.0 Integrated  —  May 2026'),
    ('Repo',       'Voice-to-Command (demo codebase)'),
    ('Status',     'Confidential — Internal Use Only'),
    ('Contact',    'prateek@silversuits.ai  |  +91 738 173 2333'),
]
for k, v in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = p.add_run(k + ':  '); rk.bold = True; rk.font.size = Pt(10)
    rv = p.add_run(v); rv.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  LEGEND — how to read this doc
# ══════════════════════════════════════════════════════════════════════
h1(doc, 'How to Read This Document')
body(doc,
    'This is the single master technical architecture document. It integrates three views that '
    'were previously separate: (A) the verified demo codebase architecture, (B) the production '
    'target architecture with rule engine and bank integration, and (C) the productization '
    'roadmap (resources, timeline, DB, analytics). Sections are labelled to distinguish '
    'what exists today vs what is planned:'
)
make_table(doc,
    headers=['Document section', 'What it covers'],
    rows=[
        ['Part I — §2', 'Code audit: what the Voice-to-Command repo actually implements'],
        ['Part II — §3', 'System architecture diagrams: demo (Fig 1–4) + production (Fig 5–8)'],
        ['Part III — §4–10', 'Data workflows, behavior analysis, rule engine, voice/AG-UI deep dives'],
        ['Part IV — §11–15', 'Resources, timeline, compliance, NFRs'],
    ],
    col_widths=[4.5, 12.1],
)
body(doc,
    'Standalone companion documents (same content, split for presentations) can be regenerated from '
    'generate_arch_diagram_doc.py and generate_production_arch_doc.py.',
    italic=True,
)
make_table(doc,
    headers=['Label', 'Meaning'],
    rows=[
        ['IN DEMO CODEBASE',  'Feature / behaviour confirmed present in the repo code right now'],
        ['PRODUCTIZATION V1', 'Planned for V1 (Weeks 1–8) — not yet built'],
        ['PRODUCTIZATION V2', 'Planned for V2 (Weeks 9–20) — not yet built'],
        ['PRODUCTIZATION V3', 'Planned for V3 (Weeks 21–36) — not yet built'],
        ['NOT IN DEMO',       'Explicitly absent from the current codebase (was wrongly described in v1 of this doc)'],
    ],
    col_widths=[4.0, 12.6],
)
callout(doc,
    'Platform scope: in-app voice and screen assistants only (V1–V5). All user interaction '
    'stays inside the bank mobile or web app — no off-app voice channels.', tone='info')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════
h1(doc, 'Table of Contents')
toc = [
    ('1.',  'Executive Summary & Scope'),
    ('2.',  'Demo Codebase — Verified Code Audit'),
    ('3.',  'System Architecture — Current Demo (Figures 1–4)'),
    ('4.',  'System Architecture — Production Target (Figures 5–8)'),
    ('5.',  'Complete Data Workflow (Demo → Production)'),
    ('6.',  'Behavior Analysis & AI Triggers'),
    ('7.',  'Rule Engine — Demo vs Production'),
    ('8.',  'AG-UI Streaming Agents'),
    ('9.',  'Voice Bot Interaction & Logging'),
    ('10.', 'Bank DB Integration (Productization V1+)'),
    ('11.', 'Analytics Dashboard (Productization V2/V3)'),
    ('12.', 'Resource Requirements'),
    ('13.', 'Timeline & Version Feature List'),
    ('14.', 'Compliance Quick-Reference'),
    ('15.', 'Exception Handling & Fallbacks'),
    ('16.', 'Non-Functional Requirements'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.add_run(f'{num:5}  {title}').font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

h2(doc, 'List of Figures')
make_table(doc,
    headers=['Figure', 'Title', 'Section'],
    rows=[
        ['1', 'Current demo — layered architecture (Voice/UPI + AG-UI)', '§3'],
        ['2', 'Voice turn sequence (demo)', '§3'],
        ['3', 'AG-UI SSE screen agent flow (demo)', '§3'],
        ['4', 'Dialogue state machine + repository map (demo)', '§3'],
        ['5', 'Production — six-layer architecture', '§4'],
        ['6', 'Production runtime — voice + rule engine', '§4'],
        ['7', 'Rule engine pipeline (production)', '§4 / §7'],
        ['8', 'DocumentDB + AWS deployment (production)', '§4'],
    ],
    col_widths=[1.2, 10.5, 2.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════
h1(doc, '1. Executive Summary & Scope')
body(doc,
    'The AI Voice RM Platform is a web-based AI banking assistant embedded inside the Indian Bank '
    'mobile app. When a user shows frustration signals — rage-clicking, repeated invalid field '
    'entries, or extended inactivity — the system automatically presents a contextual AI assistant '
    'popup. The user can speak or type; the AI guides them through the form, executes banking '
    'transactions, and routes them to the right banking journey, all within the mobile app session.'
)
body(doc,
    'The demo codebase (Voice-to-Command repo) is a fully working web simulation of this product. '
    'It runs a React frontend (simulated mobile banking UI) backed by an Express server with '
    'ElevenLabs STT, OpenAI LLM, and Cartesia TTS — all interaction stays inside the bank app. '
    'The productization roadmap adds bank DB integration, analytics, rule engine, and a visual rule-builder.'
)

make_table(doc,
    headers=['Metric', 'Demo (Codebase Today)', 'Full Product (V3)'],
    rows=[
        ['Trigger mechanism', 'Rage-tap detection (JS hook) or DemoPanel button', 'Rage taps + inactivity + bank API events'],
        ['Interaction channel', 'In-app popup (web session only)', 'In-app popup + push notification (V2)'],
        ['STT',   'ElevenLabs Scribe v2 (proxy) or Web Speech API (browser)', 'ElevenLabs Scribe v2 (India data residency)'],
        ['LLM',   'OpenAI GPT-4o-mini', 'Azure OpenAI GPT-4o (India endpoint)'],
        ['TTS',   'Cartesia sonic-3 (proxy)', 'Cartesia sonic-3 (India region)'],
        ['Languages', '4 (en, hi, ta, te)', '10 (en, hi/Hinglish, ta, te, gu, kn, pa, mr, ml, bn)'],
        ['Banking flows', '5 voice flows + 5 AG-UI screen agents', 'Same + extensible via new manifests'],
        ['Data storage', 'In-memory mock only (no real DB writes)', 'Bank-owned DocumentDB (AWS Mumbai / OnPrem)'],
        ['Analytics', 'None', 'Full dashboard (V2 basic, V3 deep analytics)'],
    ],
    col_widths=[3.4, 5.6, 7.6],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  2. CODE AUDIT — WHAT THE DEMO ACTUALLY DOES
# ══════════════════════════════════════════════════════════════════════
h1(doc, '2. Demo Codebase — Verified Code Audit')

h2(doc, '2.1 Repository Structure')
make_table(doc,
    headers=['Path', 'What It Is'],
    rows=[
        ['client/src/App.jsx',                 'Root React app — wires all state, voice hooks, and screens together'],
        ['client/src/components/VoiceModal.jsx','Main AI assistant popup — chat bubbles, mic button, TTS playback, input bar'],
        ['client/src/components/HomeScreen.jsx','Home banking screen — balance card, quick actions, AI concierge button'],
        ['client/src/components/LoanApplicationScreen.jsx','Loan form with live AG-UI form-fill'],
        ['client/src/components/ImpsFundTransferScreen.jsx','IMPS fund transfer form with AI guidance'],
        ['client/src/components/CreateDepositScreen.jsx','FD/RD/MMD deposit creation with AI'],
        ['client/src/components/TransactionHistoryScreen.jsx','Account statement with AI query support'],
        ['client/src/components/RMHelpPrompt.jsx','Bottom-sheet popup shown on frustration detection'],
        ['client/src/components/MpinSheet.jsx','MPIN authentication gate before payment confirmation'],
        ['client/src/hooks/useRageDetect.js',  'Behavior detection: rage taps (5+ in 900ms) and repeated invalid fields'],
        ['client/src/hooks/useElevenSpeech.js','ElevenLabs STT hook — MediaRecorder + VAD silence detection'],
        ['client/src/hooks/useSpeech.js',       'Web Speech API fallback STT hook (browser built-in)'],
        ['client/src/lib/cartesiaTts.js',       'Cartesia TTS client — single audio channel, VAD gate'],
        ['client/src/engine/simEngine.js',      'Thin client wrapper — forwards all turns to POST /api/engine/turn'],
        ['client/src/services/engineClient.js', 'HTTP client for engine API, account statement, and server reset'],
        ['server/index.js',                     'Express app — all API routes: /api/stt, /api/tts, /api/engine/*, /api/agui/*'],
        ['server/engine/engine.js',             'Server-side state machine — processInput(), intent routing, state transitions'],
        ['server/engine/saga.js',               'Generic saga runner — walks manifest steps, handles tool calls, validates'],
        ['server/engine/llm.js',                'LLM call — OpenAI extract() for intent + slots + reply'],
        ['server/engine/tools.js',              'Tool registry: contacts.search, billers.search, balance.check, payment execution'],
        ['server/manifests/*.json',             '5 flow manifests: send_money, check_balance, internal_transfer, pay_bill, book_flight'],
        ['server/manifestRegistry.js',          'Loads and indexes all manifests at startup'],
        ['server/agui/loanAguiRoute.js',        'SSE route dispatcher for all 5 AG-UI agents'],
        ['server/agui/homeAguiRunner.js',       'Home screen agent — navigate_to tool, UPI/IMPS routing rules'],
        ['server/agui/loanAguiRunner.js',       'Loan LOS agent — set_field tool calls, form validation'],
        ['server/agui/impsAguiRunner.js',       'IMPS transfer agent — fill transfer form fields'],
        ['server/agui/depositAguiRunner.js',    'Deposit creation agent'],
        ['server/agui/txnHistoryAguiRunner.js', 'Transaction history agent'],
        ['server/data/backend.js',              'Pure mock banking logic: payments, transfers, bill pay, flight booking'],
        ['server/data/mock.js',                 'Static data: contacts, accounts, billers, per-transaction limits'],
    ],
    col_widths=[5.4, 11.2],
)

h2(doc, '2.2 Confirmed Technical Stack (from code)')
make_table(doc,
    headers=['Component', 'Actual Implementation', 'File'],
    rows=[
        ['STT (primary)',   'ElevenLabs Scribe v2 — MediaRecorder + VAD → POST /api/stt → ElevenLabs API', 'useElevenSpeech.js, server/index.js'],
        ['STT (fallback)',  'Web Speech API (browser built-in) — Chrome/Edge best support',               'useSpeech.js'],
        ['LLM',            'OpenAI GPT-4o-mini — intent + slots + reply extraction per turn',             'server/engine/llm.js'],
        ['TTS',            'Cartesia sonic-3 — POST /api/tts → MP3 bytes → browser Audio()',             'cartesiaTts.js, server/index.js'],
        ['State machine',  'Manifest/saga runner — 5 registered flows, 9 states',                        'server/engine/engine.js, saga.js'],
        ['AG-UI agents',   '5 SSE streaming agents via OpenAI function calling',                          'server/agui/*Runner.js'],
        ['Behavior detect','useRageDetect hook — pointerdown timestamps + invalid field counter',         'useRageDetect.js'],
        ['Auth gate',      'MpinSheet — intercepts voice/tap confirm before executing payment',           'MpinSheet.jsx, App.jsx'],
        ['Languages',      'en, hi, ta, te — 4 languages, BCP47 passed to STT and LLM',                  'client/src/data/languages.js'],
        ['Framework',      'React 18 + Framer Motion (animations) + Tailwind CSS',                       'client/'],
        ['Server',         'Node.js 20 LTS + Express 4 + cors + pino logging',                           'server/'],
    ],
    col_widths=[2.8, 8.0, 5.8],
)

h2(doc, '2.3 What Is NOT in the Demo Codebase')
callout(doc, 'The following items were incorrectly described in the previous document. '
             'They are absent from the codebase and belong to the productization roadmap only.', tone='err')
make_table(doc,
    headers=['Item', 'Status', 'Planned In'],
    rows=[
        ['10-language support',                    'Only 4 languages coded (en/hi/ta/te)',  'V2'],
        ['Bank DB writes (real)',                  'All storage is in-memory mock only',    'V1'],
        ['Analytics dashboard',                    'Not implemented',                       'V2 (basic), V3 (full)'],
        ['Visual rule-builder UI',                 'Not implemented',                       'V2'],
        ['Post-session LLM analysis',              'Not implemented',                       'V3'],
        ['Emergency override dashboard',           'Not implemented',                       'V2'],
        ['Result callback API (real)',             'Not implemented',                       'V2'],
        ['Human RM escalation (real API)',         'Not implemented (mock only)',            'V2'],
        ['Behavior analytics (sentiment/churn)',   'Not implemented',                       'V3'],
    ],
    col_widths=[5.0, 4.2, 3.4],
    tones=['red','red','amber','amber','amber','amber','amber','amber','amber','amber','amber'],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  3. DEMO SYSTEM ARCHITECTURE (FIGURES 1–4)
# ══════════════════════════════════════════════════════════════════════
h1(doc, '3. System Architecture — Current Demo (Voice-to-Command)')
callout(doc,
        'Figures 1–4 describe the verified implementation in the repository. '
        'Web in-app only — mock backend (no bank DB in demo).', tone='ok')

h2(doc, '3.1 Layered Architecture — Demo')
body(doc,
     'Figure 1 shows five layers with two parallel paths: Voice/UPI (left) uses the manifest-driven '
     'state machine; AG-UI screen agents (right) use SSE streaming independently.')

insert_figure(doc, DEMO_FIG1,
    'Figure 1 — Current demo: layered architecture (Voice/UPI + AG-UI, mock backend)')

make_table(doc,
    headers=['Layer', 'Components', 'Repo path'],
    rows=[
        ['User / Browser', 'Indian Bank UI simulation', 'client/ · index.html'],
        ['React Frontend', 'VoiceModal, rage detect, 5 banking screens', 'client/src/components/, hooks/'],
        ['Express Backend', '/api/stt, /api/tts, /api/engine/turn, /api/agui/*', 'server/index.js'],
        ['AI Services', 'ElevenLabs STT, OpenAI GPT-4o-mini, Cartesia TTS', 'Proxied via server'],
        ['Mock & Tools', 'Tool registry, in-memory contacts/billers/payments', 'server/data/, engine/tools.js'],
    ],
    col_widths=[2.8, 5.5, 8.3],
)

h2(doc, '3.2 Voice / UPI Flow — Sequence')
insert_figure(doc, DEMO_FIG2,
    'Figure 2 — Voice turn sequence: speak → STT → engine → tools → TTS → playback')

h2(doc, '3.3 AG-UI Screen Agent Flow — Sequence')
insert_figure(doc, DEMO_FIG3,
    'Figure 3 — AG-UI: screen → SSE agent → set_field / navigate_to → live form update')

h2(doc, '3.4 State Machine & Repository Map')
insert_figure(doc, DEMO_FIG4,
    'Figure 4 — Nine dialogue states + key repository folders')

make_table(doc,
    headers=['Manifest', 'Slots', 'Mock tools'],
    rows=[
        ['send_money', 'recipient, amount, VPA', 'contacts.search → execute_upi_payment'],
        ['check_balance', 'account type', 'balance.check'],
        ['internal_transfer', 'amount, from, to', 'accounts.list → execute_transfer'],
        ['pay_bill', 'biller, account, amount', 'billers.search → execute_bill_payment'],
        ['book_flight', 'route, date, passenger', 'flights.search → execute_flight'],
    ],
    col_widths=[3.2, 5, 8.4],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  4. PRODUCTION TARGET ARCHITECTURE (FIGURES 5–8)
# ══════════════════════════════════════════════════════════════════════
h1(doc, '4. System Architecture — Production Target')
callout(doc,
        'Figures 5–8 describe the V1+ product deployed in bank VPC (AWS Mumbai). '
        'Builds on the demo engine; adds bank integration, rule engine, DocumentDB, analytics.',
        tone='info')

h2(doc, '4.1 Layered Production Architecture')
insert_figure(doc, PROD_FIG1,
    'Figure 5 — Production: six layers (channel → bank → platform → AI → rules → data)')

make_table(doc,
    headers=['Layer', 'Owner', 'Key addition vs demo'],
    rows=[
        ['Channel', 'Bank', 'Native mobile (V3), agent desktop (V2)'],
        ['Bank Integration', 'Bank', 'SSO, CBS events, core APIs, webhooks, mTLS'],
        ['Silver Suits Platform', 'Silver Suits', 'Gateway, Redis sessions, event bus, bank API adapter'],
        ['AI Services', 'Silver Suits', 'Azure OpenAI India; zero retention contracts'],
        ['Rule Engine', 'Silver Suits · bank configures', 'POST /api/trigger, JSON rules, <100ms'],
        ['Data & Analytics', 'Bank', 'DocumentDB, dashboard, 7-year audit'],
    ],
    col_widths=[3.2, 3.5, 9.9],
)

h2(doc, '4.2 Runtime Flows — Voice + Rule Engine')
insert_figure(doc, PROD_FIG2,
    'Figure 6 — Production runtime: voice turn path and rule engine path (parallel)')

h2(doc, '4.3 Rule Engine Architecture')
insert_figure(doc, PROD_FIG3,
    'Figure 7 — Rule engine: trigger → facts → evaluate → dispatch → audit')

h2(doc, '4.4 Data Model & AWS Deployment')
insert_figure(doc, PROD_FIG4,
    'Figure 8 — DocumentDB collections + AWS ap-south-1 deployment topology')

h2(doc, '4.5 Demo vs Production — Component Matrix')
make_table(doc,
    headers=['Component', 'Demo (today)', 'Production (V1+)', 'Host'],
    rows=[
        ['Backend', 'Express :3001', 'ECS Fargate + ALB + rule engine', 'Bank VPC Mumbai'],
        ['STT', 'ElevenLabs proxy (dev keys)', 'ElevenLabs India residency', 'ElevenLabs India'],
        ['LLM', 'OpenAI GPT-4o-mini', 'Azure OpenAI GPT-4o', 'Azure India'],
        ['TTS', 'Cartesia proxy', 'Cartesia India region', 'Cartesia India'],
        ['Banking data', 'mock.js in-memory', 'Bank REST via adapter', 'Bank CBS'],
        ['Sessions', 'React state only', 'DocumentDB + Redis cache', 'Bank-owned'],
        ['Rule engine', 'Client rage detect only', 'Full trigger → actuate pipeline', 'Silver Suits'],
        ['Analytics', 'NOT IN DEMO', 'Dashboard read-only on bank DB', 'Bank ops team'],
    ],
    col_widths=[2.6, 4.2, 4.4, 5.4],
    tones=[None, None, None, None, 'amber', 'amber', 'amber', 'red', 'amber'],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  5. COMPLETE DATA WORKFLOW
# ══════════════════════════════════════════════════════════════════════
h1(doc, '5. Complete Data Workflow')

h2(doc, '5.1 Demo Data Flow (Current)')
make_table(doc,
    headers=['Step', 'Component', 'Data In', 'Data Out', 'Stored?'],
    rows=[
        ['1', 'useRageDetect (browser)', 'pointerdown events, field validation errors', 'onFrustrated() callback fired', 'No — in-memory counters only'],
        ['2', 'RMHelpPrompt',            'onFrustrated event',                           'User taps "Yes, help me"',       'No'],
        ['3', 'VoiceModal opens',        'Session object (blank)',                        'INIT turn sent to server',      'No — session in React state'],
        ['4', 'MediaRecorder + VAD',     'Microphone audio stream',                      'Audio blob (webm/opus)',         'No'],
        ['5', 'POST /api/stt',           'Audio blob + lang param',                      'Transcript text + detected lang','No (ElevenLabs zero retention)'],
        ['6', 'POST /api/engine/turn',   'Transcript + sessionId + lang',                'Updated session (state, history, pending)', 'No — server session map in-memory'],
        ['7', 'LLM (OpenAI)',            'Utterance + state context + history (last 6)',  'intent, slots{}, reply, language', 'No (OpenAI zero retention)'],
        ['8', 'Saga runner + tools',     'Resolved intent + slots',                      'Tool results (contacts, payment)', 'No — mock in-memory'],
        ['9', 'POST /api/tts',           'Bot reply text + lang',                        'MP3 audio bytes (44100Hz, 128kbps)', 'No (Cartesia zero retention)'],
        ['10','Browser Audio',           'MP3 blob URL',                                 'Audio played to user',            'No — blob URL revoked after play'],
        ['11','MpinSheet (CONFIRM only)','User MPIN entry',                              'Success → confirmation sent to server', 'No'],
        ['12','Session result',          'DONE/FAILED/CANCELLED state',                  'Result in session.history[] (React state only)', 'No'],
    ],
    col_widths=[0.6, 3.2, 4.0, 4.4, 4.4],
)

h2(doc, '5.2 Production Data Flow — Pre-Conversation')
make_table(doc,
    headers=['Step', 'Action', 'Data Involved', 'Stored?', 'Version'],
    rows=[
        ['1', 'Bank system emits event (button click, rage-click, inactivity)',
              'event_type, user_id, form_profile{}, frontend_context{}', 'No', 'V1'],
        ['2', 'Bank calls POST /api/trigger on Silver Suits backend',
              'user_profile{}, event_type, form_profile{}, branch_context{}', 'No (rule evaluation only)', 'V1'],
        ['3', 'Rule engine evaluates conditions in-memory (< 100ms)',
              'All trigger payload tokens + loaded rule set', 'No', 'V1'],
        ['4', 'Rule decision: action + script_id + language',
              'action_type, script_id, lang, visibility', 'No', 'V1'],
        ['5', 'Backend creates session record in bank DB',
              'session_id, user_id, journey_type, lang, started_at', 'YES — bank DB', 'V1'],
        ['5a','WebSocket push to browser → VoiceModal popup',
              'script_id, lang, user_name, field_context', 'No (WebSocket ephemeral)', 'V1'],
    ],
    col_widths=[0.6, 4.0, 5.6, 2.4, 2.0],
)

h2(doc, '5.3 Production Data Flow — During Conversation (Per Turn)')
body(doc, 'This is identical to the demo flow (steps 4–9 above) except AI provider keys use '
         'India data-residency configs, and step 7 additionally writes the turn to bank DB.')

h2(doc, '5.4 Production Data Flow — Post-Conversation')
make_table(doc,
    headers=['Step', 'Action', 'Data', 'Destination', 'Version'],
    rows=[
        ['1', 'Session state → DONE/CANCELLED/FAILED',
              'outcome, duration_s, turn_count', 'Bank DB — sessions collection (UPDATE)', 'V1'],
        ['2', 'Optional result callback',
              'session_id, outcome, transcript_summary', 'Bank backend (POST webhook)', 'V2'],
        ['3', 'Post-session LLM analysis',
              'Full transcript text', 'Azure OpenAI (India) — zero retention on LLM', 'V3'],
        ['4', 'Structured analytics JSON stored',
              'sentiment_arc, churn_risk, qa_score, script_compliance, flags[]', 'Bank DB — analytics collection', 'V3'],
        ['5', 'Rule audit log written',
              'rule_id, matched_params{}, action_taken, ts', 'Bank DB — rule_audit collection', 'V1'],
    ],
    col_widths=[0.6, 3.6, 5.0, 4.4, 2.0],
)

h2(doc, '5.5 Data That Is NEVER Stored (Demo and Production)')
never = [
    'Voice audio recordings — discarded immediately after STT processing (zero retention on ElevenLabs)',
    'Voice biometric data — not collected at any stage of the pipeline',
    'Raw audio files — no disk writes on any AI service (STT, LLM, TTS)',
    'LLM conversation context — Azure OpenAI and OpenAI both configured with zero data retention',
    'Cartesia TTS inputs — not stored after audio generation',
    'Customer PII outside bank DB — Silver Suits backend never caches PII on disk (stateless)',
]
for item in never:
    bullet(doc, item)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  5. BEHAVIOR ANALYSIS
# ══════════════════════════════════════════════════════════════════════
h1(doc, '6. Behavior Analysis — How Frustration Is Detected & What Triggers the AI')

h2(doc, '6.1 Current Implementation (IN DEMO CODEBASE — useRageDetect.js)')
callout(doc, 'This is the actual behavior detection implemented in the code. '
             'It runs entirely client-side in the browser.', tone='ok')

body(doc,
    'The useRageDetect hook monitors two independent frustration signals and fires a single '
    "onFrustrated() callback when either threshold is crossed. This opens the 'Need a little help?' "
    'bottom-sheet prompt. The hook auto-resets after 8 seconds so it can fire again if needed.'
)

make_table(doc,
    headers=['Signal', 'Threshold', 'Detection Logic', 'Code Location'],
    rows=[
        ['Rage taps',          '5+ taps in 900ms on same element/screen',
         'pointerdown timestamps pushed to array; array filtered to last 900ms; count >= RAGE_TAP_COUNT fires',
         'useRageDetect.js — handleTap()'],
        ['Repeated invalid field', '2+ invalid submissions of same field',
         'markInvalidField(fieldId) called on validation failure; counter per fieldId; >= INVALID_FIELD_THRESHOLD fires',
         'useRageDetect.js — markInvalidField()'],
    ],
    col_widths=[3.0, 3.4, 5.8, 5.4],
)

h3(doc, 'Constants (from code)')
make_table(doc,
    headers=['Constant', 'Value', 'Purpose'],
    rows=[
        ['RAGE_TAP_COUNT',         '5',     'Number of taps needed within the window to trigger'],
        ['RAGE_TAP_WINDOW',        '900ms', 'Sliding time window for tap counting'],
        ['INVALID_FIELD_THRESHOLD','2',     'Times same field set invalid before triggering'],
        ['Auto-reset delay',       '8000ms','Prevents re-fire for 8 seconds after triggering'],
    ],
    col_widths=[4.0, 2.0, 10.6],
)

h3(doc, 'What happens after frustration is detected (App.jsx)')
steps_after = [
    '1. onFrustrated() callback fires → setRmUpiPromptOpen(true)',
    '2. RMHelpPrompt bottom-sheet appears: "Looks like you hit a snag. Want me to take over?"',
    '3. User taps "Yes, help me" → setOpen(true) → VoiceModal opens → INIT turn sent to server',
    '4. User taps "No, thanks" → dismiss() → rage detection resets → no AI session starts',
]
for s in steps_after:
    body(doc, s, size=9.5)

h2(doc, '6.2 Additional Trigger Points (IN DEMO CODEBASE)')
make_table(doc,
    headers=['Trigger', 'Where', 'How'],
    rows=[
        ['HomeScreen AI button tap',  'HomeScreen.jsx',   'Mic button on home → handleMicTap() → VoiceModal opens with INIT turn'],
        ['DemoPanel scripted prompt', 'DemoPanel.jsx',    'onSpeak(text) → runUtterance(text) → TRANSCRIPT turn with preset text'],
        ['Quick action button',       'HomeScreen.jsx',   'onQuickAction(actionName) → handleOpenWithAction → START_ACTION turn'],
        ['Navigate-to from AG-UI',    'homeAguiRunner.js','Home agent fires navigate_to tool → App.handleNavigate() → opens destination screen'],
    ],
    col_widths=[3.6, 3.2, 9.8],
)

h2(doc, '6.3 Production Behavior Signals (Productization Roadmap)')
body(doc, 'In production, behavior signals come from both the client (same rage detection) '
         'and the bank backend system via the trigger API payload.')

make_table(doc,
    headers=['Signal', 'Source', 'How Sent to Silver Suits', 'Version'],
    rows=[
        ['Rage taps (in-app)',           'Mobile app JS',     'Same useRageDetect hook + POST /api/trigger', 'V1'],
        ['Inactivity duration',          'Mobile app timer',  'inactive_time_s in secondary_signals[]',      'V1'],
        ['Form completion percentage',   'Mobile app',        'form_completion_pct in form_profile{}',       'V1'],
        ['Repeat-click count',           'Mobile app',        'repeat_click_count in secondary_signals[]',   'V1'],
        ['User profile / segment',       'Bank DB',           'user_profile{} in trigger payload',           'V1'],
        ['Previous churn risk score',    'Bank analytics DB', 'user_profile.churn_risk_prev',                'V3'],
        ['Session language preference',  'Bank user profile', 'user_profile.language_pref',                  'V1'],
    ],
    col_widths=[4.0, 2.8, 5.4, 2.4],
)

h2(doc, '6.4 Post-Session Behavior Analytics (V3 — Not Yet Built)')
callout(doc, 'None of the following are in the current codebase. '
             'They require the post-session LLM analysis pipeline planned for V3.', tone='err')
make_table(doc,
    headers=['Datapoint', 'Category', 'Extraction Method'],
    rows=[
        ['Sentiment trajectory (neg/neu/pos per turn)', 'Customer behaviour',     'LLM labels each turn; arc stored as array'],
        ['Churn risk score (0–100)',                    'Customer behaviour',     'LLM predictive scoring on conversation signals'],
        ['Intent categorisation',                       'Customer behaviour',     'LLM assigns primary tag (loan/KYC/UPI/support)'],
        ['Script compliance % (per flow type)',         'Hard skills',            'LLM checks mandatory milestone presence'],
        ['Empathy score',                               'Soft skills',            'LLM detects validating language markers'],
        ['Automated QA score (1–100)',                  'Composite',              'Weighted composite of all parameters'],
        ['Talk-to-listen ratio',                        'Conversational dynamics','Bot token count vs user token count'],
        ['Prohibited content flags',                    'Compliance',             'Regex + LLM scan on transcript'],
        ['Regulatory disclosure check',                 'Compliance',             'LLM verifies mandatory phrases were spoken'],
        ['Up-sell / cross-sell opportunities',          'Revenue',                'LLM extracts product context + intent strength (H/M/L)'],
    ],
    col_widths=[5.4, 3.4, 7.8],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  6. RULE ENGINE
# ══════════════════════════════════════════════════════════════════════
h1(doc, '7. Rule Engine — Demo vs Production')
body(doc,
     'In the demo, only client-side frustration detection opens the AI popup. In production (Figure 7), '
     'the bank sends events to POST /api/trigger and a full rule engine evaluates JSON rules in under 100ms.')

h2(doc, '7.1 What Exists in Demo Today (IN DEMO CODEBASE)')
callout(doc, 'The demo has a routing rule set in homeAguiConfig.js (hardcoded) and a '
             'manifest/saga engine for conversation flows. No visual rule-builder exists yet.', tone='ok')

h3(doc, 'Banking channel routing rules (homeAguiConfig.js — hardcoded)')
make_table(doc,
    headers=['Condition', 'Destination', 'Rationale'],
    rows=[
        ['Amount > ₹1,00,000 OR IMPS/NEFT/bank account/IFSC mentioned', 'fund_transfer',         'Above UPI per-transaction limit of ₹1 lakh'],
        ['Amount ≤ ₹1,00,000 OR UPI/mobile pay/phone mentioned',        'upi_payment',           'Within UPI limit'],
        ['Loan request mentioned',                                        'loan_application',     'Loan LOS journey'],
        ['Deposit/FD/RD/MMD/fixed deposit/recurring mentioned',          'create_deposit',       'Deposit creation journey'],
        ['Transaction history/account statement/suspicious credit',       'transaction_history',  'Statement view journey'],
        ['Completely ambiguous',                                          'Ask 1 clarifying question', 'Avoid wrong routing'],
    ],
    col_widths=[5.6, 3.4, 7.6],
)

h3(doc, 'Conversation state machine rules (server/engine/engine.js)')
make_table(doc,
    headers=['State', 'Allowed Inputs', 'Transition Logic'],
    rows=[
        ['IDLE',        'INIT, START_ACTION, TRANSCRIPT',           'TRANSCRIPT → LLM extracts action → START_ACTION'],
        ['FILL',        'TRANSCRIPT, CANCEL',                       'TRANSCRIPT → extract slots → advance saga'],
        ['ASK',         'TRANSCRIPT, CANCEL',                       'TRANSCRIPT → fill slot → advance'],
        ['DISAMBIGUATE','TRANSCRIPT, SELECTION, CANCEL',            'Pick from options list → advance'],
        ['CHOOSE',      'TRANSCRIPT, SELECTION, CANCEL',            'Pick sub-option → advance'],
        ['CONFIRM',     'CONFIRMATION, TRANSCRIPT (yes/no), CANCEL','yes → execute tool_call; no → CANCELLED'],
        ['DONE',        'RESET, START_ACTION',                      'Terminal — auto-reset on next action'],
        ['FAILED',      'RESET, START_ACTION',                      'Terminal — error surfaced to user'],
        ['CANCELLED',   'RESET, START_ACTION',                      'Terminal — transcript saved as-is'],
    ],
    col_widths=[2.4, 4.2, 10.0],
)

h2(doc, '7.2 Productization Rule Engine (V1+ — New Build)')
callout(doc, 'The production rule engine is a new build on top of the same backend, '
             'evaluating bank-provided trigger payloads rather than in-app routing.', tone='warn')

body(doc, 'Production trigger payload — what the bank sends to POST /api/trigger:')
make_table(doc,
    headers=['Field', 'Type', 'Description', 'Required'],
    rows=[
        ['user_id',           'string',  "Bank's internal user identifier",                                       'Yes'],
        ['user_profile',      'object',  'name, segment, language_pref, account_type, churn_risk_prev',          'Yes'],
        ['event_type',        'enum',    'help_button_clicked / rage_click / inactivity / account_created / custom','Yes'],
        ['form_profile',      'object',  'field_ids[], field_statuses{}, validation_regex{}, form_completion_pct','Conditional'],
        ['frontend_context',  'object',  'current_page, journey_type, button_id, timestamp_ms',                  'Yes'],
        ['branch_context',    'object',  'branch_code, nearest_human_rm_id, geography',                          'Optional'],
        ['secondary_signals', 'array',   'rage_click_count, inactivity_duration_s, repeat_click_count',          'Optional'],
    ],
    col_widths=[3.0, 1.8, 8.0, 2.8],
)

body(doc, 'Example production rules:')
code_line(doc, 'IF [form_completion_pct > 75] AND [repeat_click_count > 5]')
code_line(doc, '   THEN [action=trigger_popup, script=loan_help, lang=en]')
doc.add_paragraph()
code_line(doc, 'IF [inactive_time_s > 20] AND [event_type = inactivity]')
code_line(doc, '   THEN [action=trigger_popup, script=push_user, lang=hi]')
doc.add_paragraph()
code_line(doc, 'IF [churn_risk_prev > 70] AND [event_type = rage_click]')
code_line(doc, '   THEN [action=assign_human_rm, script=retention_assist, priority=high]')

doc.add_paragraph()

h2(doc, '7.3 Actions the Rule Engine Can Actuate (Production)')
make_table(doc,
    headers=['Action', 'How Actuated', 'Parameters', 'Version'],
    rows=[
        ['trigger_popup',       'WebSocket push to browser',              'script_id, lang, user_name, field_context',  'V1'],
        ['assign_human_rm',     "POST bank's CRM webhook",                'user_id, branch_code, rm_id, priority',      'V1'],
        ['send_notification',   'Bank notification API (SMS/push)',        'user_id, template_id, params{}',             'V2'],
        ['fill_form_field',     'WebSocket command to browser',            'field_id, value, validation_regex',          'V1'],
        ['log_event',           'Write to bank DB',                       'session_id, event_type, payload, ts',        'V1'],
        ['abort_session',       'WebSocket close + state = CANCELLED',    'reason_code, user_id',                       'V1'],
    ],
    col_widths=[3.2, 4.0, 5.0, 2.2],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  7. AG-UI STREAMING AGENTS
# ══════════════════════════════════════════════════════════════════════
h1(doc, '8. AG-UI Streaming Agents — Per-Screen AI (Deep Dive)')

h2(doc, '8.1 What AG-UI Is (IN DEMO CODEBASE)')
body(doc,
    'AG-UI (Agent-UI) is a Server-Sent Events (SSE) protocol where the backend streams '
    'AI agent actions to the browser in real time. Each screen has its own dedicated agent '
    'running OpenAI function calling. The agent can call tools (e.g. set_field) that '
    'update the form live while the user is speaking — no waiting for a full response before '
    'the form starts filling.'
)

h2(doc, '8.2 Per-Agent Specification')

agents = [
    {
        'name': 'Home Screen Agent (indian_bank_home_assistant)',
        'file': 'server/agui/homeAguiRunner.js + homeAguiConfig.js',
        'role': 'Acts as a smart banking concierge. Understands what the customer wants (payment, loan, deposit, statement) and routes them to the correct screen using the navigate_to tool.',
        'tools': [
            ('navigate_to', 'destination: upi_payment|fund_transfer|loan_application|create_deposit|transaction_history, context: string, routing_status: string'),
        ],
        'routing': [
            'Amount > ₹1L OR IMPS/NEFT → fund_transfer',
            'Amount ≤ ₹1L OR UPI → upi_payment',
            'Loan request → loan_application',
            'Deposit/FD/RD → create_deposit',
            'Transaction query → transaction_history',
        ],
        'lang': 'Responds in user language (Hindi / Hinglish / English — from system prompt)',
    },
    {
        'name': 'Loan LOS Agent (indian_bank_loan_los)',
        'file': 'server/agui/loanAguiRunner.js + loanAgentConfig.js + loanTools.js',
        'role': 'Fills the loan application form in real time as the user speaks. Can infer multiple field values from one utterance and fill them simultaneously.',
        'tools': [
            ('set_field', 'fieldId: occupation|subProduct|purposeLoan|variant|facility|proposal|interestType|loanAmount|tenureMonths|branchPin, value: string'),
            ('validate_form', 'state: current form state — returns validation snapshot'),
        ],
        'routing': ['10 loan fields', 'Select fields: enum-validated (o0/o1/o2, lp1-lp5, etc.)', 'Numeric: loanAmount, tenureMonths', 'Text: branchPin'],
        'lang': 'Warm RM tone; can ask multiple fields in one question',
    },
    {
        'name': 'IMPS Fund Transfer Agent (indian_bank_imps_transfer)',
        'file': 'server/agui/impsAguiRunner.js + impsAguiConfig.js + impsTools.js',
        'role': 'Guides the user through within-bank or other-bank IMPS transfer form.',
        'tools': [('set_field', 'fieldId: transferType|payeeType|payeeName|payeeAccountNo|ifsc|payeeBank|mobileNo|amount|remarks')],
        'routing': ['Asks ONE question at a time (strict flow)', 'within-bank: name + account number', 'other-bank/account: IFSC + account', 'other-bank/mobile: bank name + mobile number'],
        'lang': 'Friendly Indian Bank RM persona',
    },
    {
        'name': 'Create Deposit Agent (indian_bank_create_deposit)',
        'file': 'server/agui/depositAguiRunner.js + depositAguiConfig.js',
        'role': 'Helps user select deposit type (FD/RD/MMD) and fill amount, tenure, maturity instructions.',
        'tools': [('set_field', 'fieldId: various deposit form fields')],
        'routing': ['FD: lump sum, fixed tenure', 'RD: monthly installment, tenure', 'MMD: money multiplier deposit'],
        'lang': 'Warm, helpful tone',
    },
    {
        'name': 'Transaction History Agent (indian_bank_txn_history)',
        'file': 'server/agui/txnHistoryAguiRunner.js + txnHistoryAguiConfig.js',
        'role': 'Answers user questions about their account statement, suspicious transactions, or expected credits. Can navigate to payment flows for relevant queries.',
        'tools': [('set_field', 'display/filter fields'), ('navigate_to', 'when user wants to initiate a payment based on a transaction')],
        'routing': ['Suspicious credit / fraud alert → explain, log', 'Wanted to pay someone → navigate_to upi_payment'],
        'lang': 'Calm, reassuring tone for fraud/dispute queries',
    },
]

for agent in agents:
    h3(doc, agent['name'])
    label_body(doc, 'File', agent['file'])
    label_body(doc, 'Role', agent['role'])
    h3(doc, '  Tools')
    for tool_name, tool_params in agent['tools']:
        body(doc, f'  {tool_name}({tool_params})', size=9, colour=TEAL)
    h3(doc, '  Logic / Flow')
    for r in agent['routing']:
        bullet(doc, r, level=1)
    doc.add_paragraph()

h2(doc, '8.3 SSE Streaming Protocol')
body(doc,
    'All AG-UI agents stream via Server-Sent Events. Each SSE event is a JSON object '
    'with a type field. The browser processes events in order to update UI state in real time.'
)
make_table(doc,
    headers=['SSE Event Type', 'Payload', 'Browser Effect'],
    rows=[
        ['TEXT_MESSAGE_START',    '{messageId, role}',                       'New message bubble opened'],
        ['TEXT_MESSAGE_CONTENT',  '{messageId, delta}',                      'Text appended to bubble (streaming)'],
        ['TEXT_MESSAGE_END',      '{messageId}',                             'Bubble finalised; TTS triggered on text'],
        ['TOOL_CALL_START',       '{toolCallId, toolCallName}',              'Agent about to call a tool'],
        ['TOOL_CALL_ARGS',        '{toolCallId, delta}',                     'Tool args streamed'],
        ['TOOL_CALL_END',         '{toolCallId}',                            'Tool call complete'],
        ['RUN_FINISHED',          '{threadId, runId}',                       'Agent turn complete'],
        ['RUN_ERROR',             '{message, code}',                         'Error shown to user'],
    ],
    col_widths=[3.6, 4.6, 8.4],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  8. VOICE BOT INTERACTION & LOGGING
# ══════════════════════════════════════════════════════════════════════
h1(doc, '9. Voice Bot Interaction Flow & Per-Turn Logging')

h2(doc, '9.1 Per-Turn Pipeline (IN DEMO CODEBASE)')
make_table(doc,
    headers=['Step', 'Component', 'Action', 'Target Latency'],
    rows=[
        ['1', 'Browser (in-app)',           'User speaks; MediaRecorder starts on mic permission grant',             '—'],
        ['2', 'AnalyserNode (VAD)',         'RMS checked every animation frame; 1.5s silence-stop armed after speech starts', '—'],
        ['3', 'MediaRecorder.stop()',       'chunks[] → Blob (webm/opus); recording ends',                          '—'],
        ['4', 'POST /api/stt?lang=en',      'Audio blob → ElevenLabs Scribe v2 → transcript text + lang code',     '300–500ms'],
        ['5', 'POST /api/engine/turn',      'Transcript + sessionId → server state machine processes turn',         '—'],
        ['6', 'LLM extract() (OpenAI)',     'utterance + state context + last 6 history turns → intent JSON',      '400–700ms'],
        ['7', 'applyIntent() + advance()',  'Slots merged; saga stepped forward; next pending computed',            '< 20ms'],
        ['8', 'Tool calls (if needed)',     'contacts.search / billers.search / execute_payment (mock)',            '50–200ms (mock)'],
        ['9', 'Bot reply text resolved',    'Session updated; history[] and state returned to client',              '—'],
        ['10','POST /api/tts',             'Bot reply text → Cartesia sonic-3 → MP3 bytes returned',              '300–500ms'],
        ['11','Audio.play()',               'MP3 blob URL played in browser; TTS flag set',                         '—'],
        ['12','waitUntilTtsIdle()',          'Auto-mic arms 500ms after TTS ends; loop back to step 1',             '—'],
    ],
    col_widths=[0.6, 3.4, 7.4, 2.2],
)

h2(doc, '9.2 Hands-Free Auto-Mic Logic (App.jsx)')
body(doc,
    'The auto-mic effect re-arms the microphone after each bot reply so the user never '
    'has to tap the mic button between turns. It is gated by multiple conditions to prevent '
    'the mic from capturing TTS audio or opening while the state machine is thinking.'
)
make_table(doc,
    headers=['Condition Checked', 'Why'],
    rows=[
        ['open === true',               'Modal must be visible'],
        ['mpinOpen === false',          'MPIN sheet takes priority — user is entering their PIN'],
        ['speech.supported === true',   'STT backend available'],
        ['session.thinking === false',  'LLM is not currently processing'],
        ['session.executing === false', 'Tool call is not in flight'],
        ['session.state in VOICE_INPUT_STATES', 'Only arm in FILL / DISAMBIGUATE / CHOOSE / CONFIRM states'],
        ['ttsPlaying === false',        'Bot is not currently speaking'],
        ['waitUntilTtsIdle() resolved', 'Belt-and-suspenders TTS idle check'],
        ['500ms pause after TTS ends',  'Let user absorb the bot message before mic opens'],
    ],
    col_widths=[4.4, 12.2],
)

h2(doc, '9.3 MPIN Authentication Gate (MpinSheet.jsx + App.jsx)')
body(doc,
    'Every payment confirmation is gated by MPIN authentication. The voice "yes" interceptor '
    'in App.jsx catches voice confirmations and routes them through the MPIN sheet before '
    'the actual confirmation is sent to the server.'
)
make_table(doc,
    headers=['Scenario', 'Flow'],
    rows=[
        ['User taps "Confirm" button', 'handleConfirm() → setMpinOpen(true) → MpinSheet shown'],
        ['User speaks "yes"/"haan" during CONFIRM state', 'looksLikeYes() → speech.abort() → setMpinOpen(true)'],
        ['User speaks "no"/"cancel"', 'Passes through normally → handleConfirmTap(session, false)'],
        ['MPIN success', 'setMpinOpen(false) → handleConfirmTap(session, true) → server executes payment'],
        ['MPIN cancel', 'setMpinOpen(false) → stay on CONFIRM screen → user can re-tap Confirm'],
    ],
    col_widths=[4.0, 12.6],
)

h2(doc, '9.4 What Is Logged Per Turn (session.history[] — Demo)')
make_table(doc,
    headers=['Field', 'Type', 'Description'],
    rows=[
        ['role',         'enum: user | bot',  'Who said it'],
        ['text',         'string',            'Exact STT transcript or LLM-generated bot reply'],
        ['t',            'unix timestamp ms', 'Date.now() at time of push to history'],
    ],
    col_widths=[2.6, 3.4, 10.6],
)
callout(doc, 'In the demo, session.history[] lives in server-side in-memory Map only. '
             'Nothing is written to a database. In production (V1+), every turn is '
             'INSERTed into bank DB (turns collection) with additional fields: intent, slots, engine_state, lang.', tone='warn')

h2(doc, '9.5 Production Session Lifecycle Events (V1+ — Written to Bank DB)')
make_table(doc,
    headers=['Event', 'Trigger', 'Data Written to Bank DB'],
    rows=[
        ['SESSION_STARTED',    'First turn processed',               'session_id, user_id, journey_type, lang, channel, trigger_event, consent_given, ts'],
        ['TURN_LOGGED',        'After every conversation turn',      'session_id, role, text, intent, slots{}, engine_state, lang, t'],
        ['FORM_FIELD_FILLED',  'fill_form_field tool succeeds',      'session_id, field_id, value, validated, ts'],
        ['CONFIRMATION_GIVEN', 'User confirms (post-MPIN)',          'session_id, action_summary, slots_snapshot, ts'],
        ['FALLBACK_TRIGGERED', 'Exception path entered',             'session_id, fallback_type, reason, ts'],
        ['CONSENT_RECORDED',   'User accepts/declines session start','session_id, consent_given: bool, channel, ts'],
        ['SESSION_ENDED',      'State → DONE/CANCELLED/FAILED',     'session_id, outcome, duration_s, turn_count, final_state, ended_at'],
        ['HUMAN_RM_ASSIGNED',  'Escalation rule fires',              'session_id, user_id, rm_id, branch_code, reason, ts'],
    ],
    col_widths=[3.4, 3.8, 9.4],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  9. BANK DB INTEGRATION
# ══════════════════════════════════════════════════════════════════════
h1(doc, '10. Bank DB Integration — Read / Write Specification (Productization V1+)')
callout(doc, 'None of this exists in the demo. All data in demo is mock/in-memory. '
             'This section describes the V1+ production integration.', tone='warn')

h2(doc, '10.1 Database Schema (DocumentDB / MongoDB)')
make_table(doc,
    headers=['Collection', 'Key Fields', 'Indexes', 'Retention'],
    rows=[
        ['sessions',    'session_id, user_id, journey_type, outcome, started_at, ended_at, lang, turn_count, channel',
                        'user_id + started_at; outcome + started_at', 'Per bank policy (≥7 yrs financial per RBI)'],
        ['turns',       'session_id, role, text, intent, slots{}, engine_state, lang, t',
                        'session_id + t',                              'Same as sessions'],
        ['form_events', 'session_id, field_id, value, validated, ts',
                        'session_id + field_id',                       'Same as sessions'],
        ['analytics',   'session_id, user_id, sentiment_arc[], churn_risk, qa_score, script_compliance, talk_ratio, flags[]',
                        'session_id; user_id + churn_risk',            'Same as sessions'],
        ['rule_audit',  'rule_id, version, matched_params{}, action_taken, ts',
                        'ts; rule_id',                                 '2 years (compliance audit)'],
        ['escalations', 'session_id, user_id, rm_id, branch_code, reason, ts',
                        'user_id + ts',                                '2 years'],
    ],
    col_widths=[2.4, 6.2, 4.0, 4.0],
)

h2(doc, '10.2 Write Operations')
make_table(doc,
    headers=['Operation', 'Trigger', 'Collection', 'Payload Size', 'Version'],
    rows=[
        ['INSERT session',           'Session starts',                'sessions',    '~500 B',    'V1'],
        ['INSERT turn (per turn)',   'After each dialogue turn',      'turns',       '200–800 B', 'V1'],
        ['INSERT form_event',        'Form field filled',             'form_events', '~200 B',    'V1'],
        ['UPDATE session outcome',   'Session ends',                  'sessions',    'Patch only', 'V1'],
        ['INSERT analytics',         'Post-session LLM analysis (V3)',   'analytics',   '2–5 KB',    'V3'],
        ['INSERT rule_audit',        'Rule match fires',              'rule_audit',  '~300 B',    'V1'],
        ['INSERT escalation',        'Human RM assigned',             'escalations', '~200 B',    'V2'],
    ],
    col_widths=[3.6, 3.6, 2.8, 2.2, 2.4],
)

h2(doc, '10.3 Read Operations')
make_table(doc,
    headers=['Operation', 'When', 'Purpose', 'Collection'],
    rows=[
        ['GET user profile',             'Pre-session (rule engine)',       'Contextualise script + language selection',     "Bank's user table (read-only view)"],
        ['GET previous sessions (user)', 'Pre-session (V3 churn)',          'Feed prior outcomes into rule engine',           'sessions'],
        ['GET session + turns',          'Dashboard: timeline',          'Per-user history + conversation replay',         'sessions + turns'],
        ['GET analytics record',         'Dashboard: QA/sentiment',      'Display analytics per session',                  'analytics'],
        ['GET aggregate by date',        'Dashboard: volume charts',     'Session count / outcome rate over time',            'sessions'],
        ['GET escalations by user',      'Rule engine cap check',        'Count recent escalations for abuse guard',       'escalations'],
    ],
    col_widths=[3.8, 3.0, 4.6, 5.2],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  10. ANALYTICS DASHBOARD
# ══════════════════════════════════════════════════════════════════════
h1(doc, '11. Analytics Dashboard — DB Read Patterns (Productization V2/V3)')
callout(doc, 'Analytics dashboard does not exist in the demo. Planned for V2 (basic) and V3 (full).', tone='warn')

make_table(doc,
    headers=['View', 'Query Pattern', 'Collections', 'Version'],
    rows=[
        ['Session volume over time',         'Aggregate sessions by date bucket; filter journey_type',     'sessions',              'V2'],
        ['Outcome breakdown',             'GROUP BY outcome; COUNT per period',                          'sessions',              'V2'],
        ['Branch-wise performance',       'JOIN sessions + user profile → GROUP BY branch_code',         'sessions + user view',  'V2'],
        ['Language distribution',         'GROUP BY lang, COUNT',                                        'sessions',              'V2'],
        ['Escalation rate',               'COUNT escalations / COUNT sessions per period',               'sessions + escalations','V2'],
        ['Conversation timeline (user)',  'FIND sessions WHERE user_id=X; JOIN turns',                  'sessions + turns',      'V3'],
        ['Conversation replay',           'FIND turns WHERE session_id=Y; ORDER BY t',                  'turns',                 'V3'],
        ['QA score distribution',         'Histogram of analytics.qa_score over date range',             'analytics',             'V3'],
        ['Sentiment arc aggregation',     'Avg sentiment_arc[] per journey_type',                        'analytics',             'V3'],
        ['Churn risk cohort',             'Bucket users by churn_risk; show count + avg outcome',        'analytics + sessions',  'V3'],
        ['Multi-journey funnel',          'Count sessions at each milestone per journey_type',            'turns (milestones)',    'V3'],
    ],
    col_widths=[4.0, 5.6, 3.6, 2.4],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  11. RESOURCES
# ══════════════════════════════════════════════════════════════════════
h1(doc, '12. Resource Requirements')

h2(doc, '12.1 Core Team')
make_table(doc,
    headers=['Role', 'FTE', 'Core Responsibilities', 'Active Versions'],
    rows=[
        ['Backend Engineer', '2',
         'Rule engine (production), REST APIs, STT/TTS proxies, DB integration, session management',
         'V1–V3'],
        ['AI / Conversation Engineer', '1',
         'Manifest/saga authoring, LLM prompt design, AG-UI agents, post-session analysis pipeline (V3)',
         'V1–V3'],
        ['Frontend Engineer (React)', '1',
         'Analytics dashboard, rule-builder UI, voice modal improvements, multi-language expansion',
         'V1–V3'],
        ['DevOps / Cloud Engineer', '1',
         'Azure (backend + dashboard), AWS DB access, CI/CD, secrets management, monitoring',
         'V1–V3'],
        ['Product Manager', '1',
         'Sprint planning, bank stakeholder management, compliance sign-offs, requirement grooming',
         'V1–V3'],
        ['QA Engineer', '1',
         'Scripted flow tests, eval harness, edge-case coverage, load testing, regression suite',
         'V1–V3'],
        ['Compliance / Legal', '0.5',
         'DPDP/RBI review, DPA execution, audit-trail sign-offs, grievance redressal',
         'V1–V3'],
    ],
    col_widths=[3.8, 0.8, 9.2, 2.8],
)

h2(doc, '12.2 Bank-Side Integration Effort')
make_table(doc,
    headers=['Task', 'Bank Role', 'Effort', 'When'],
    rows=[
        ['Wire rage-detect + trigger API in mobile app',   'Mobile/Backend engineer', '1–2 weeks', 'V1 — Wk 2'],
        ['Provision DocumentDB schema + access grants',    'DBA',                     '1 week',    'V1 — Wk 2'],
        ['Whitelist Silver Suits IPs on DB + API firewall','Security / DBA',          '1–2 days',  'V1 — Wk 2'],
        ['Implement result callback receiver',              'Backend engineer',        '3–5 days',  'V2 — Wk 10'],
        ['Connect human RM escalation webhook',             'CRM / Backend',          '3–5 days',  'V2 — Wk 11'],
        ['Provide branding guidelines',                     'Marketing / Design',     '1 week',    'V2 — Wk 11'],
        ['Compliance and legal review',                     'Compliance / Legal',     '2–3 weeks', 'V1 — Wk 1–3'],
    ],
    col_widths=[5.0, 3.2, 2.4, 3.0],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  12. TIMELINE & FEATURE LIST
# ══════════════════════════════════════════════════════════════════════
h1(doc, '13. Timeline & Version Feature List')

make_table(doc,
    headers=['Version', 'Weeks', 'Theme', 'Key Milestone'],
    rows=[
        ['V1 — MVP',             '1–8',  'Bank integration + real DB',      'Trigger API live, real sessions persist to bank DB'],
        ['V2 — Production-Ready','9–20', 'Rule builder + analytics v1',     'Rule-builder UI, analytics dashboard, push notifications'],
        ['V3 — Full Intelligence','21–36','Behavior analytics + AI insights','Post-session analysis, conversation timeline, churn prediction'],
    ],
    col_widths=[3.2, 1.6, 4.0, 7.8],
)

h2(doc, '13.1 Sprint Plan')
make_table(doc,
    headers=['Sprint', 'Weeks', 'Ver', 'Deliverables'],
    rows=[
        ['1–2',  '1–2',   'V1', 'Azure backend, bank API contract signed, DB schema v1, CI/CD, TLS/auth setup'],
        ['3–4',  '3–4',   'V1', 'Production rule engine (IF/THEN), POST /api/trigger endpoint, real STT/TTS keys (India residency)'],
        ['5–6',  '5–6',   'V1', 'Bank DB writes (sessions, turns, form_events), consent flow, cap/rate limiting'],
        ['7–8',  '7–8',   'V1', 'Exception handling (all paths), human RM webhook, V1 UAT with bank team'],
        ['9–10', '9–10',  'V2', 'Push notification gateway (FCM/APNs), WebSocket session channel hardening'],
        ['11–12','11–12', 'V2', 'Rule-builder UI (visual editor → JSON config), 6 additional languages'],
        ['13–14','13–14', 'V2', 'Emergency override dashboard, DEV/PROD split, eval harness'],
        ['15–16','15–16', 'V2', 'Analytics dashboard v1 (volume, outcomes, branch-wise, language)'],
        ['17–18','17–18', 'V2', 'Result callback API, branding customisation, V2 load testing (1000 concurrent)'],
        ['19–20','19–20', 'V2', 'Compliance audit, DPA/NDA execution, VAPT, V2 go-live'],
        ['21–24','21–24', 'V3', 'Post-session LLM analysis pipeline — sentiment arc, intent category'],
        ['25–28','25–28', 'V3', 'Churn risk scoring, script compliance, QA score, talk-ratio, compliance flags'],
        ['29–32','29–32', 'V3', 'Full analytics dashboard (multi-journey, timeline, conversation replay)'],
        ['33–36','33–36', 'V3', 'Revenue-opportunity extraction, predictive triggers from churn score, V3 go-live'],
    ],
    col_widths=[1.4, 1.4, 1.0, 12.8],
)

h2(doc, '13.2 V1 Feature List — Weeks 1–8 (Building on Demo)')
make_table(doc,
    headers=['Feature', 'Demo Status', 'V1 Action'],
    rows=[
        ['STT (ElevenLabs Scribe v2)',       'IN DEMO (dev keys)',       'Switch to India data residency + production keys'],
        ['LLM (OpenAI GPT-4o-mini)',         'IN DEMO (dev keys)',       'Switch to Azure OpenAI India endpoint'],
        ['TTS (Cartesia sonic-3)',            'IN DEMO (dev keys)',       'Switch to Cartesia India region + production keys'],
        ['Manifest/saga state machine',       'IN DEMO (complete)',       'No change needed — add new manifests as required'],
        ['AG-UI screen agents (all 5)',       'IN DEMO (complete)',       'No change needed for V1'],
        ['useRageDetect behavior detection',  'IN DEMO (complete)',       'Bank embeds in their mobile app SDK'],
        ['Trigger REST API',                  'NOT IN DEMO',             'Build POST /api/trigger with rule engine evaluation'],
        ['Production rule engine',            'NOT IN DEMO',             'Build in-memory expression evaluator, JSON rule configs'],
        ['Bank DB writes (sessions + turns)', 'NOT IN DEMO (mock only)', 'Wire DocumentDB driver; write on every turn'],
        ['Consent management',                'NOT IN DEMO',             'Add consent gate before first session turn'],
        ['Per-user cap / rate limiting',      'NOT IN DEMO',             'Read escalations count from DB; abort at cap'],
        ['4 → 10 language support',           'NOT IN DEMO (4 only)',    'Add gu/kn/pa/mr/ml/bn to STT lang codes + manifests'],
        ['Result callback API',               'NOT IN DEMO',             'POST outcome + transcript summary to bank after session'],
    ],
    col_widths=[4.6, 3.6, 8.4],
    tones=[None, None, None, 'green', 'green', 'green', 'amber', 'amber', 'amber', 'amber', 'amber', 'amber', 'amber'],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  13. COMPLIANCE
# ══════════════════════════════════════════════════════════════════════
h1(doc, '14. Compliance Quick-Reference')

make_table(doc,
    headers=['Component', 'Host', 'Data Stored?', 'AI Training?', 'India?', 'Cert'],
    rows=[
        ['Backend API',        'Azure India (Central + South)',   'No (stateless)',         'N/A', 'Yes', 'SOC2/ISO27001/PCI DSS'],
        ['STT (ElevenLabs)',   'ElevenLabs India data residency', 'Zero retention',          'No',  'Yes', 'SOC2 Type II, HIPAA'],
        ['LLM (Azure OpenAI)', 'Azure India — Chennai endpoint',  'Zero retention',          'No',  'Yes', 'SOC2 Type II, ISO27001'],
        ['TTS (Cartesia)',      'Cartesia India region',           'Zero retention',          'No',  'Yes', 'SOC2 Type II'],
        ['Bank Database',      'Bank-controlled AWS Mumbai / OnPrem', 'Permanent — bank owns','N/A','Yes', 'AWS SOC2 / bank cert'],
        ['Analytics Dashboard','Azure India',                     'No (reads bank DB)',      'N/A', 'Yes', 'SOC2/ISO27001'],
    ],
    col_widths=[2.8, 3.8, 3.0, 1.6, 1.4, 4.0],
)

body(doc, 'Encryption: AES-256 at rest, TLS 1.2+ in transit — mandatory across all components. '
         'DPDP Act 2023 + RBI guidelines compliance. 7-year minimum retention for financial records (PMLA). '
         'Bank owns all persistent data. Silver Suits never retains voice audio or customer PII on disk.')

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  14. EXCEPTION HANDLING
# ══════════════════════════════════════════════════════════════════════
h1(doc, '15. Exception Handling & Fallbacks')

h2(doc, '15.1 Demo — Currently Handled (IN DEMO CODEBASE)')
make_table(doc,
    headers=['Exception', 'How Handled (code)', 'File'],
    rows=[
        ['STT returns empty transcript',    'engine re-prompts with "didnt_catch"; previousBotMessage re-spoken',          'engine.js — handleSmalltalk()'],
        ['LLM unavailable / API error',     'Speaks "model unavailable" message; does not loop; turn ends gracefully',    'engine.js — intent.error === llm_unavailable'],
        ['Unknown intent (no action match)','handleSmalltalk() — speaks LLM contextual reply OR re-prompts pending slot', 'engine.js — handleSmalltalk()'],
        ['User says cancel/nahi/stop',      'State → CANCELLED; cancellation message spoken',                             'engine.js — doCancel()'],
        ['Bank mock API returns error',     'State → FAILED; BackendError code shown to user; retryable flag',            'backend.js — BackendError class'],
        ['Bank mock: random 5–7% failure',  'Simulates real bank decline; user sees FAILED with retry option',            'backend.js — flaky()'],
        ['Force-fail mode (demo panel)',     'bank_declined / network_error injected via X-Force-Fail header',            'backend.js + DemoPanel.jsx'],
        ['Disambiguation (multiple matches)','State → DISAMBIGUATE; options list shown; user picks by voice or tap',      'saga.js + VoiceModal.jsx'],
        ['TTS 503 (Cartesia not configured)','Console warning only; conversation continues without TTS audio',            'cartesiaTts.js'],
    ],
    col_widths=[4.0, 6.6, 6.0],
)

h2(doc, '15.2 Production — Additional Fallbacks (V2+)')
make_table(doc,
    headers=['Exception', 'Detection', 'Action', 'Logged?'],
    rows=[
        ['Tech glitch / connection failure', 'WebSocket or API error event',  'Log; assign to nearest human RM', 'Yes'],
        ['Session abandoned',                'User closes app mid-session',  'Log; 2 abandons/day → human RM',    'Yes'],
        ['Session dropped midway',           'WebSocket close / network loss','Save transcript; log as incomplete','Yes'],
        ['Insufficient rule parameters',     'Rule eval fails validation',    'Use default action; flag to team', 'Yes'],
        ['User exceeds per-session cap',     'Cap counter in DB',             'Abort; log; assign human RM',     'Yes'],
    ],
    col_widths=[4.0, 3.8, 4.8, 4.0],
)

divider(doc)

# ══════════════════════════════════════════════════════════════════════
#  15. NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════
h1(doc, '16. Non-Functional Requirements')

make_table(doc,
    headers=['Requirement', 'Demo (Current)', 'Production Target', 'Notes'],
    rows=[
        ['E2E turn latency', '~1.0–1.5s (STT+LLM+TTS)', '< 1.2s in-app', 'STT 500ms + LLM 600ms + TTS 400ms (parallel where possible)'],
        ['Rule engine decision', 'N/A (instant in-browser routing)', '< 100ms', 'In-memory eval; no DB calls during evaluation'],
        ['Concurrent sessions', '~10 (dev server)', '1,000+ simultaneous', 'Azure App Service auto-scale; load tested in V2 (week 17–18)'],
        ['Language support', '4 (en/hi/ta/te)', '10 languages', '6 additional languages added in V2'],
        ['System availability', 'N/A (local dev)', '99.9% uptime', 'Azure + ElevenLabs + Cartesia SLAs combined'],
        ['Data freshness (dashboard)', 'N/A', '< 5 min lag', 'Async DB writes; dashboard reads directly from bank DB'],
        ['Rule change cycle', 'Instant (hardcoded)', '14 business days avg (prod rules)', 'Emergency override bypasses for urgent compliance updates'],
        ['VAPT / Security', 'Not applicable (mock)', 'Before each major go-live', 'Penetration testing before V2 and V3 go-live'],
        ['RBI/DPDP audit', 'N/A', 'Continuous readiness', '30-day purge cycle; full audit logs in bank DB'],
    ],
    col_widths=[3.0, 3.4, 3.2, 7.0],
)

# ══════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run(
    "This document is confidential and intended for Indian Bank's technical and product teams only.\n"
    'Silver Suits AI  ·  prateek@silversuits.ai  ·  +91 738 173 2333  ·  Mumbai | Bangalore  ·  May 2026\n'
    'Version 3.0 Integrated — Demo + Production architecture + Productization roadmap — May 2026'
)
r.font.size = Pt(9); r.font.color.rgb = GREY

out = '/Users/devanshusaindane/SilverSuits/Voice-to-command/Voice-to-Command/Documentation/AI_Voice_RM_Platform_Technical_Architecture.docx'
doc.save(out)
print(f'Saved → {out}')
