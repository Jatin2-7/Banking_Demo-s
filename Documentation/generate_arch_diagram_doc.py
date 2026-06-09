"""
System Architecture Document — Demo / Current Codebase (Enterprise edition)
Silver Suits AI — Voice-to-Command Repository
"""
import sys
sys.path.insert(0, '/tmp/docxlib')

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Enterprise palette (matches production doc) ─────────────────────
BG       = '#F8FAFC'
LANE_HDR = '#0F172A'
TEXT     = '#0F172A'
SUB      = '#64748B'
ARROW    = '#334155'

LANES = {
    'user':   ('#E0F2FE', '#0369A1', 'USER / BROWSER'),
    'react':  ('#EEF2FF', '#4338CA', 'REACT FRONTEND  ·  client/'),
    'express':('#F0FDF4', '#15803D', 'EXPRESS BACKEND  ·  server/  :3001'),
    'ai':     ('#F3E8FF', '#7C3AED', 'AI SERVICES  ·  External APIs'),
    'data':   ('#FFEDD5', '#C2410C', 'MOCK DATA & TOOLS  ·  In-memory demo'),
}
VOICE_EC = '#0369A1'
AGUI_EC  = '#7C3AED'


def lane(ax, y, h, W, key):
    fc, ec, label = LANES[key]
    ax.add_patch(Rectangle((0, y), W, h, fc=fc, ec='none', zorder=0))
    ax.add_patch(Rectangle((0, y), 0.55, h, fc=LANE_HDR, ec='none', zorder=1))
    ax.text(0.28, y + h / 2, label, ha='center', va='center',
            fontsize=6.2, fontweight='bold', color='white', rotation=90, zorder=2)


def comp(ax, x, y, w, h, title, sub=None, ec='#4338CA', fc='white', ts=8.5, ss=7):
    p = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0,rounding_size=0.12',
                        fc=fc, ec=ec, lw=1.8, zorder=3)
    ax.add_patch(p)
    cx = x + w / 2
    if sub:
        ax.text(cx, y + h * 0.65, title, ha='center', va='center',
                fontsize=ts, fontweight='bold', color=TEXT, zorder=4)
        ax.text(cx, y + h * 0.30, sub, ha='center', va='center',
                fontsize=ss, color=SUB, zorder=4)
    else:
        ax.text(cx, y + h / 2, title, ha='center', va='center',
                fontsize=ts, fontweight='bold', color=TEXT, zorder=4)


def arr(ax, x1, y1, x2, y2, label=None, color=ARROW, rad=0, dashed=False):
    ls = (0, (4, 3)) if dashed else '-'
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.6,
                                linestyle=ls,
                                connectionstyle=f'arc3,rad={rad}',
                                shrinkA=5, shrinkB=5), zorder=5)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my - 0.12, label, ha='center', va='top',
                fontsize=6.5, color=SUB, zorder=6,
                bbox=dict(fc=BG, ec='none', pad=1, alpha=0.95))


def fig_title(ax, W, title, subtitle, y=0.12):
    ax.text(W / 2, y, title, ha='center', fontsize=14, fontweight='bold', color=TEXT)
    ax.text(W / 2, y + 0.42, subtitle, ha='center', fontsize=9, color=SUB)


def legend(ax, W, y, items):
    for i, (col, lbl) in enumerate(items):
        xo = 0.7 + i * (W - 1.4) / max(len(items), 1)
        ax.plot([xo, xo + 0.5], [y, y], color=col, lw=2.5)
        ax.text(xo + 0.65, y, lbl, fontsize=7, color=SUB, va='center')


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 1 — LAYERED SYSTEM (Voice left | AG-UI right)
# ═══════════════════════════════════════════════════════════════════════
def fig_system_layers():
    W, H = 18, 12.5
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 1 — System Architecture (Current Demo Codebase)',
              'Web-based in-app assistant · Verified Voice-to-Command repo · In-app only')

    lanes_def = [
        (0.9,  1.2,  'user'),
        (2.2,  2.5,  'react'),
        (4.8,  1.65, 'express'),
        (6.55, 2.0,  'ai'),
        (8.65, 1.55, 'data'),
    ]
    for y, h, key in lanes_def:
        lane(ax, y, h, W, key)

    # Flow labels
    ax.text(4.5, 1.05, 'VOICE / UPI FLOW', ha='center', fontsize=9,
            fontweight='bold', color=VOICE_EC)
    ax.text(13.5, 1.05, 'AG-UI SCREEN AGENTS', ha='center', fontsize=9,
            fontweight='bold', color=AGUI_EC)
    ax.plot([9.0, 9.0], [1.2, 8.5], color='#CBD5E1', lw=1.5, linestyle='--', zorder=1)

    ox_v, ox_a = 0.75, 9.2
    cw, ch = 2.45, 0.68

    # User
    comp(ax, 3.0, 1.05, 5.5, 0.75, 'User · Browser · Indian Bank UI simulation', ec=LANES['user'][1])

    # React — voice column
    voice_react = [
        ('useRageDetect', '5 taps / 900ms · invalid ×2'),
        ('RMHelpPrompt', '"Need help?" sheet'),
        ('VoiceModal', 'Chat · mic · TTS play'),
        ('MediaRecorder + VAD', 'useElevenSpeech.js'),
        ('MpinSheet', 'Payment confirm gate'),
    ]
    for i, (t, s) in enumerate(voice_react):
        comp(ax, ox_v + (i % 2) * 2.6, 2.35 + (i // 2) * 0.82, cw, ch, t, s, ec=VOICE_EC)

    # React — AG-UI column
    agui_react = [
        ('HomeScreen', 'Concierge · quick actions'),
        ('LoanApplicationScreen', 'Loan LOS form'),
        ('ImpsFundTransferScreen', 'IMPS form'),
        ('CreateDepositScreen', 'FD / RD / MMD'),
        ('TransactionHistoryScreen', 'Statement + AI query'),
    ]
    for i, (t, s) in enumerate(agui_react):
        comp(ax, ox_a + (i % 2) * 2.6, 2.35 + (i // 2) * 0.82, cw, ch, t, s, ec=AGUI_EC)

    # Express — voice
    comp(ax, ox_v, 4.95, cw, ch, 'POST /api/stt', 'ElevenLabs proxy', ec=LANES['express'][1])
    comp(ax, ox_v + 2.7, 4.95, 2.8, ch, 'POST /api/engine/turn', 'State machine', ec=LANES['express'][1])
    comp(ax, ox_v, 5.75, cw, ch, 'POST /api/tts', 'Cartesia MP3', ec=LANES['express'][1])

    # Express — AG-UI
    comp(ax, ox_a, 4.95, 3.2, ch, 'POST /api/agui/{agentId}', 'SSE streaming', ec=AGUI_EC)
    comp(ax, ox_a + 3.4, 4.95, 2.5, ch, 'GET /api/manifests', 'Flow configs', ec=LANES['express'][1])

    # AI — voice path
    comp(ax, ox_v, 6.75, cw, ch, 'ElevenLabs Scribe v2', 'STT · India', ec=LANES['ai'][1])
    comp(ax, ox_v + 2.7, 6.65, 2.8, 0.88, 'State Machine + Saga', 'engine.js · 9 states', ec=LANES['express'][1], ts=8, ss=6.5)
    comp(ax, ox_v + 2.7, 7.65, 2.8, ch, 'OpenAI GPT-4o-mini', 'Intent + reply', ec=LANES['ai'][1])
    comp(ax, ox_v, 7.65, cw, ch, 'Cartesia sonic-3', 'TTS', ec=LANES['ai'][1])

    # AI — AG-UI
    comp(ax, ox_a, 6.75, 3.2, 0.88, 'OpenAI GPT-4o-mini', 'Function calling · SSE', ec=AGUI_EC, ts=8, ss=6.5)

    # Mock / tools
    comp(ax, ox_v, 8.95, cw, ch, 'Tool Registry', 'tools.js', ec=LANES['data'][1])
    comp(ax, ox_v + 2.7, 8.95, 2.8, ch, 'Mock Backend', 'backend.js · mock.js', ec=LANES['data'][1])
    comp(ax, ox_a, 8.95, 3.2, ch, 'set_field / navigate_to', 'SSE → React state', ec=AGUI_EC)
    comp(ax, ox_a + 3.4, 8.95, 2.5, ch, 'Browser playback', 'Audio.play()', ec=LANES['data'][1])

    # Key flows
    arr(ax, 5.5, 1.8, 2.0, 2.35, color=VOICE_EC)
    arr(ax, 2.0, 5.63, 2.0, 6.75, 'audio', color=LANES['ai'][1])
    arr(ax, 4.1, 5.63, 4.1, 6.65, color=LANES['express'][1])
    arr(ax, 4.1, 7.53, 4.1, 7.65, color=LANES['ai'][1])
    arr(ax, 4.1, 8.33, 4.1, 8.95, 'tools', color=LANES['data'][1])
    arr(ax, 14.0, 3.0, 10.5, 4.95, color=AGUI_EC)
    arr(ax, 10.5, 5.63, 10.5, 6.75, color=AGUI_EC)
    arr(ax, 10.5, 7.63, 10.5, 8.95, color=AGUI_EC, rad=0.15)
    arr(ax, 11.5, 9.29, 14.0, 3.5, 'form update', color=AGUI_EC, dashed=True, rad=0.25)

    legend(ax, W, 10.35, [
        (VOICE_EC, 'Voice / UPI flow'),
        (AGUI_EC, 'AG-UI screen agents'),
        (LANES['ai'][1], 'AI vendor APIs'),
        (LANES['data'][1], 'Mock data (demo only)'),
    ])

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 2 — VOICE FLOW SEQUENCE
# ═══════════════════════════════════════════════════════════════════════
def fig_voice_sequence():
    W, H = 16, 7
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 2 — Voice / UPI Conversation Flow',
              'Manifest-driven state machine · One turn ≈ 1.5–2.5 seconds')

    actors = ['User', 'React App', 'Express API', 'AI Services', 'Mock Backend']
    aw = 2.9
    for i, a in enumerate(actors):
        comp(ax, 0.5 + i * aw, 1.0, aw - 0.2, 0.5, a,
             ec=VOICE_EC if i == 0 else LANES['express'][1] if i == 2 else LANES['ai'][1], ts=8)

    steps = [
        (0, 1, '1. Speak / tap mic'),
        (1, 2, '2. POST audio /api/stt'),
        (2, 3, '3. ElevenLabs transcript'),
        (2, 2, '4. POST /api/engine/turn'),
        (2, 3, '5. GPT-4o-mini intent'),
        (2, 4, '6. Tool call (if needed)'),
        (4, 2, '7. Mock payment'),
        (2, 3, '8. Generate reply'),
        (2, 2, '9. POST /api/tts'),
        (3, 2, '10. Cartesia MP3'),
        (2, 1, '11. Return audio'),
        (1, 0, '12. Play + re-arm mic'),
    ]
    y0 = 1.85
    for j, (fi, ti, lbl) in enumerate(steps):
        y = y0 + j * 0.42
        x1 = 0.5 + fi * aw + aw / 2
        x2 = 0.5 + ti * aw + aw / 2
        col = LANES['data'][1] if fi == 4 or ti == 4 else VOICE_EC if fi <= 1 else LANES['ai'][1] if fi == 3 or ti == 3 else LANES['express'][1]
        arr(ax, x1, y, x2, y, lbl, color=col)

    # MPIN branch note
    comp(ax, 0.5, 6.2, 15, 0.55, 'CONFIRM state: voice "yes" intercepted → MpinSheet → then execute_payment',
         fc='#FFFBEB', ec='#D97706', ts=8, ss=7)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 3 — AG-UI FLOW SEQUENCE
# ═══════════════════════════════════════════════════════════════════════
def fig_agui_sequence():
    W, H = 16, 6.5
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 3 — AG-UI Screen Agent Flow',
              'SSE streaming · Independent of voice state machine · 5 screen-specific agents')

    actors = ['User', 'Bank Screen', 'POST /api/agui', 'OpenAI', 'React Form']
    aw = 3.0
    for i, a in enumerate(actors):
        comp(ax, 0.5 + i * aw, 1.0, aw - 0.2, 0.5, a, ec=AGUI_EC if i in (1, 4) else LANES['ai'][1], ts=8)

    steps = [
        (0, 1, '1. User speaks'),
        (1, 2, '2. Forward messages[]'),
        (2, 3, '3. Stream GPT-4o-mini'),
        (3, 2, '4. SSE TEXT_MESSAGE'),
        (3, 2, '5. SSE TOOL_CALL'),
        (2, 4, '6. set_field / navigate'),
        (4, 1, '7. Form updates'),
        (2, 2, '8. TTS spoken reply'),
    ]
    y0 = 1.75
    for j, (fi, ti, lbl) in enumerate(steps):
        y = y0 + j * 0.48
        x1 = 0.5 + fi * aw + aw / 2
        x2 = 0.5 + ti * aw + aw / 2
        arr(ax, x1, y, x2, y, lbl, color=AGUI_EC)

    agents = 'Agents: home_assistant · loan_los · imps_transfer · create_deposit · txn_history'
    ax.text(W / 2, 5.85, agents, ha='center', fontsize=8, color=SUB,
            bbox=dict(boxstyle='round', facecolor='white', edgecolor='#CBD5E1'))

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


# ═══════════════════════════════════════════════════════════════════════
#  FIGURE 4 — STATE MACHINE + REPO MAP
# ═══════════════════════════════════════════════════════════════════════
def fig_state_and_repo():
    W, H = 16, 8
    fig, ax = plt.subplots(figsize=(W, H), dpi=150)
    ax.set_facecolor(BG); fig.patch.set_facecolor(BG)
    ax.set_xlim(0, W); ax.set_ylim(0, H); ax.invert_yaxis(); ax.axis('off')
    fig_title(ax, W, 'Figure 4 — Dialogue State Machine & Repository Layout',
              'Server-side engine · Manifest-driven flows')

    # State machine (left)
    ax.text(4, 0.95, 'VOICE STATE MACHINE (9 states)', ha='center',
            fontsize=10, fontweight='bold', color=LANES['express'][1])
    states = [
        ('IDLE', 'Start'),
        ('FILL', 'Collect slots'),
        ('ASK', 'Ask question'),
        ('DISAMBIGUATE', 'Pick contact'),
        ('CHOOSE', 'Pick VPA'),
        ('CONFIRM', 'MPIN gate'),
        ('DONE', 'Success'),
        ('FAILED', 'Error'),
        ('CANCELLED', 'User cancel'),
    ]
    for i, (st, desc) in enumerate(states):
        row, col = i // 3, i % 3
        comp(ax, 0.4 + col * 2.6, 1.35 + row * 0.95, 2.4, 0.75, st, desc,
             ec=LANES['express'][1], ts=8, ss=6.5)

    # Repo structure (right)
    ax.text(12, 0.95, 'REPOSITORY STRUCTURE', ha='center',
            fontsize=10, fontweight='bold', color=VOICE_EC)
    repo = [
        ('client/src/', 'React UI · hooks · components'),
        ('client/src/engine/', 'simEngine.js → calls server'),
        ('server/index.js', 'Express routes · API entry'),
        ('server/engine/', 'engine.js · saga.js · llm.js · tools.js'),
        ('server/manifests/', 'send_money.json · 5 flows'),
        ('server/agui/', '5 AG-UI agent configs + runners'),
        ('server/data/', 'mock.js · backend.js (in-memory)'),
    ]
    for i, (path, desc) in enumerate(repo):
        comp(ax, 8.2, 1.35 + i * 0.88, 7.5, 0.72, path, desc, ec=VOICE_EC, ts=8, ss=7)

    # Manifests bar
    comp(ax, 0.4, 5.5, 15.2, 0.7,
         'Manifests: send_money · check_balance · internal_transfer · pay_bill · book_flight',
         'Each JSON defines slots, steps, validations, tool bindings — add new flow = new JSON file',
         ec=LANES['data'][1], fc='#FFFBEB', ts=8.5, ss=7)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=BG, pad_inches=0.15)
    buf.seek(0); plt.close()
    return buf


def get_demo_diagrams():
    """Return diagram buffers for import by unified technical architecture doc."""
    return (
        fig_system_layers(),
        fig_voice_sequence(),
        fig_agui_sequence(),
        fig_state_and_repo(),
    )


_STANDALONE = __name__ == '__main__'
if _STANDALONE:
    print('Generating demo diagrams...')
    buf1, buf2, buf3, buf4 = get_demo_diagrams()

# ═══════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT (standalone only — skipped when imported)
# ═══════════════════════════════════════════════════════════════════════
if _STANDALONE:
    NAVY   = RGBColor(0x0F, 0x17, 0x2A)
ACCENT = RGBColor(0x1E, 0x40, 0xAF)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x64, 0x74, 0x8B)
GREEN  = RGBColor(0x05, 0x76, 0x48)
PURPLE = RGBColor(0x6D, 0x28, 0xD9)
ORANGE = RGBColor(0xC2, 0x41, 0x0C)


def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_colour)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def tbl(doc, headers, rows, widths=None, hdr='0F172A'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hr.cells[i], hdr)
        p = hr.cells[i].paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
    for idx, row in enumerate(rows):
        tr = t.add_row()
        for i, v in enumerate(row):
            if idx % 2 == 1:
                set_cell_bg(tr.cells[i], 'F1F5F9')
            p = tr.cells[i].paragraphs[0]
            p.clear()
            p.add_run(str(v)).font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            set_col_width(t, i, w)
    doc.add_paragraph()
    return t


def h1(doc, t):
    p = doc.add_heading(t, 1)
    if p.runs:
        p.runs[0].font.color.rgb = NAVY
        p.runs[0].font.size = Pt(18)


def h2(doc, t, c=None):
    p = doc.add_heading(t, 2)
    if p.runs:
        p.runs[0].font.color.rgb = c or ACCENT
        p.runs[0].font.size = Pt(13)


def h3(doc, t):
    p = doc.add_heading(t, 3)
    if p.runs:
        p.runs[0].font.color.rgb = NAVY
        p.runs[0].font.size = Pt(11)


def para(doc, t):
    p = doc.add_paragraph(t)
    if p.runs:
        p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def callout(doc, t):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'EEF2FF')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    r = p.add_run(t)
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(8)


def img(doc, buf, cap, w=6.5):
    doc.add_picture(buf, width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        c = doc.add_paragraph(cap)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.runs:
            c.runs[0].font.size = Pt(9)
            c.runs[0].italic = True
            c.runs[0].font.color.rgb = GREY
    doc.add_paragraph()


def pb(doc):
    doc.add_page_break()


# ── Document ────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.top_margin = sec.bottom_margin = Inches(0.7)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# COVER
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SILVER SUITS AI')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = GREY

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('AI Voice RM Platform')
r2.bold = True
r2.font.size = Pt(24)
r2.font.color.rgb = NAVY

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('System Architecture\nCurrent Demo Codebase')
r3.font.size = Pt(16)
r3.font.color.rgb = ACCENT

doc.add_paragraph()
tbl(doc, ['Document', 'Value'], [
    ['Repository', 'Voice-to-Command (verified code audit)'],
    ['Classification', 'Confidential — Internal / Partner Use'],
    ['Version', '3.0'],
    ['Date', 'May 2026'],
    ['Deployment model', 'Local / demo — React + Express :3001'],
    ['Scope note', 'Web in-app only · Mock backend'],
], widths=[4, 12])

pb(doc)

# 1. EXECUTIVE SUMMARY
h1(doc, '1. Executive Summary')
para(doc,
     'This document describes the system architecture of the Voice-to-Command demo as implemented '
     'in source code. It is the reference for what exists today before productization (bank VPC, '
     'rule engine, DocumentDB, production APIs).')

callout(doc,
        'Verified finding: Architecture is web-based only. All voice interaction is in-app '
        'in the repository. Two parallel capabilities: Voice/UPI dialogue engine and AG-UI screen agents.')

h2(doc, 'Solution Overview')
tbl(doc,
    ['Capability', 'Implementation', 'Location in repo'],
    [
        ['In-app voice assistant', 'VoiceModal + server state machine', 'client/components · server/engine/'],
        ['Frustration detection', 'useRageDetect → RMHelpPrompt', 'client/src/hooks/ · components/'],
        ['5 voice banking flows', 'JSON manifests + saga runner', 'server/manifests/'],
        ['5 screen AI agents', 'SSE + OpenAI function calling', 'server/agui/'],
        ['Speech services', 'ElevenLabs STT, Cartesia TTS (proxied)', 'server/index.js'],
        ['Language model', 'OpenAI GPT-4o-mini', 'server/engine/llm.js · server/agui/'],
        ['Banking operations', 'Mock in-memory data', 'server/data/mock.js · backend.js'],
        ['Languages supported', '4: English, Hindi, Tamil, Telugu', 'client/src/data/languages.js'],
    ],
    widths=[4, 5.5, 7.5],
)

h2(doc, 'Demo vs Production (Summary)')
tbl(doc,
    ['Area', 'Current demo', 'Production target (see Production Architecture doc)'],
    [
        ['Hosting', 'Local dev (Vite + Express)', 'Bank VPC · AWS Mumbai'],
        ['Bank APIs', 'Mock backend.js', 'Real CBS / payment APIs'],
        ['Data persistence', 'React state only — no DB', 'DocumentDB sessions + turns'],
        ['Rule engine', 'Client rage detect only', 'Full trigger → evaluate → actuate'],
        ['Analytics', 'None', 'Bank dashboard'],
        ['LLM', 'OpenAI GPT-4o-mini', 'Azure OpenAI GPT-4o (India)'],
    ],
    widths=[3.5, 5, 8.5],
)

pb(doc)

# 2. SYSTEM ARCHITECTURE
h1(doc, '2. System Architecture')
para(doc,
     'Figure 1 shows five architectural layers. The system splits into two runtime paths that share '
     'the same Express server and AI vendors but use different orchestration logic.')

img(doc, buf1, 'Figure 1 — Layered architecture: Voice/UPI flow (left) and AG-UI agents (right)', w=6.6)

h2(doc, '2.1 Layer Specification')
tbl(doc,
    ['Layer', 'Technology', 'Responsibility', 'Key artifacts'],
    [
        ['User / Browser', 'Chrome / mobile browser', 'Indian Bank UI simulation, mic, taps', 'App.jsx, index.html'],
        ['React Frontend', 'React 18 · Vite · Tailwind', 'UI, rage detect, voice modal, banking screens', 'client/src/components/, hooks/'],
        ['Express Backend', 'Node.js · Express :3001', 'API proxies, state machine, AG-UI runners', 'server/index.js, engine/, agui/'],
        ['AI Services', 'ElevenLabs · OpenAI · Cartesia', 'STT, intent, TTS — keys server-side only', 'Env vars · proxy routes'],
        ['Mock & Tools', 'In-memory JS objects', 'Contacts, billers, payments — demo data', 'server/data/, engine/tools.js'],
    ],
    widths=[2.8, 3.2, 5, 6],
)

h2(doc, '2.2 Dual Runtime Paths')
tbl(doc,
    ['Path', 'Entry point', 'Orchestration', 'Output'],
    [
        ['Voice / UPI', 'VoiceModal mic or RMHelpPrompt', 'POST /api/engine/turn → engine.js → saga.js', 'Spoken reply + optional mock payment'],
        ['AG-UI agents', 'Mic on banking screen', 'POST /api/agui/{id} → SSE stream', 'Live form fields + navigation + spoken reply'],
    ],
    widths=[2.5, 4.5, 5.5, 5.5],
)

pb(doc)

# 3. VOICE FLOW
h1(doc, '3. Voice / UPI Conversation Architecture')
img(doc, buf2, 'Figure 2 — Voice turn sequence across components', w=6.6)

h2(doc, '3.1 Turn Processing Pipeline')
tbl(doc,
    ['Step', 'Component', 'Action', 'Typical latency'],
    [
        ['1', 'useRageDetect / VoiceModal', 'Capture intent to speak or open from help prompt', '—'],
        ['2', 'useElevenSpeech (MediaRecorder + VAD)', 'Record audio; stop after 1.5s silence', '—'],
        ['3', 'POST /api/stt', 'Proxy to ElevenLabs Scribe v2', '300–500 ms'],
        ['4', 'POST /api/engine/turn', 'processInput(TRANSCRIPT)', '< 20 ms'],
        ['5', 'llm.js extract()', 'GPT-4o-mini → action, slots, reply', '400–700 ms'],
        ['6', 'saga.js + tools.js', 'Run manifest steps; mock tool calls', '50–200 ms'],
        ['7', 'POST /api/tts', 'Proxy to Cartesia sonic-3', '300–500 ms'],
        ['8', 'VoiceModal', 'Audio.play(); re-arm mic after idle', '—'],
    ],
    widths=[0.7, 3.5, 7.3, 2.5],
)

h2(doc, '3.2 Frustration Detection → AI Entry')
tbl(doc,
    ['Signal', 'Threshold', 'Component', 'Result'],
    [
        ['Rage taps', '5 pointerdown events within 900 ms', 'useRageDetect.js', 'onFrustrated() → RMHelpPrompt'],
        ['Invalid field', 'Same field fails validation 2×', 'useRageDetect.js', 'onFrustrated() → RMHelpPrompt'],
        ['User accepts help', 'Tap "Yes, help me"', 'RMHelpPrompt.jsx', 'VoiceModal opens · INIT turn'],
    ],
    widths=[2.5, 4, 3.5, 7],
)

h2(doc, '3.3 Voice Flow Manifests')
tbl(doc,
    ['Manifest', 'Slots', 'Tools', 'Limits'],
    [
        ['send_money', 'recipient, amount, VPA', 'contacts.search → execute_upi_payment', 'UPI ≤ ₹1L'],
        ['check_balance', 'account type', 'balance.check', 'savings / current / all'],
        ['internal_transfer', 'amount, from, to', 'accounts.list → execute_transfer', 'No same-account transfer'],
        ['pay_bill', 'biller, account, amount', 'billers.search → execute_bill_payment', 'Biller min amount'],
        ['book_flight', 'route, date, passenger, flight', 'flights.search → execute_flight', 'Demo flow'],
    ],
    widths=[3, 4.5, 5, 4.5],
)

h2(doc, '3.4 State Machine (9 States)')
img(doc, buf4, 'Figure 4 — State machine states and repository layout (excerpt)', w=6.6)

tbl(doc,
    ['State', 'Purpose', 'Key transitions'],
    [
        ['IDLE', 'Awaiting start', 'TRANSCRIPT → FILL (LLM picks action)'],
        ['FILL / ASK', 'Collecting slots', 'Slot fill → next saga step'],
        ['DISAMBIGUATE / CHOOSE', 'Multiple matches', 'User SELECTION'],
        ['CONFIRM', 'Payment approval', 'MPIN gate · yes/no parser'],
        ['DONE / FAILED / CANCELLED', 'Terminal', 'RESET or new START_ACTION'],
    ],
    widths=[3, 4, 10],
)

pb(doc)

# 4. AG-UI
h1(doc, '4. AG-UI Screen Agent Architecture')
img(doc, buf3, 'Figure 3 — AG-UI SSE streaming flow', w=6.6)

para(doc,
     'AG-UI agents operate independently of the voice state machine. Each banking screen mounts '
     'its own agent ID and streams Server-Sent Events for text and tool calls.')

h2(doc, '4.1 Agent Catalog')
tbl(doc,
    ['Screen', 'Agent ID', 'Tools', 'Behaviour'],
    [
        ['Home', 'indian_bank_home_assistant', 'navigate_to', 'Route by amount: ≤₹1L UPI, >₹1L IMPS'],
        ['Loan LOS', 'indian_bank_loan_los', 'set_field, validate_form', '10 loan fields · multi-slot inference'],
        ['IMPS', 'indian_bank_imps_transfer', 'set_field', '9 fields · one question at a time'],
        ['Deposit', 'indian_bank_create_deposit', 'set_field', 'FD / RD / MMD'],
        ['Txn History', 'indian_bank_txn_history', 'set_field, navigate_to', 'Query txns · fraud navigation'],
    ],
    widths=[2.2, 4.2, 3.5, 7.1],
    hdr='6D28D9',
)

h2(doc, '4.2 SSE Event Protocol')
tbl(doc,
    ['Event', 'Purpose'],
    [
        ['TEXT_MESSAGE_START / CONTENT / END', 'Streaming chat bubble text'],
        ['TOOL_CALL_START / ARGS / END', 'set_field or navigate_to execution'],
        ['RUN_FINISHED / RUN_ERROR', 'Turn complete or failure'],
    ],
    widths=[5, 12],
    hdr='6D28D9',
)

pb(doc)

# 5. COMPONENTS & APIs
h1(doc, '5. Component Reference')

h2(doc, '5.1 React Frontend Components')
tbl(doc,
    ['Component', 'File', 'Role'],
    [
        ['useRageDetect', 'hooks/useRageDetect.js', 'Frustration signals → onFrustrated callback'],
        ['useElevenSpeech', 'hooks/useElevenSpeech.js', 'STT via MediaRecorder + /api/stt'],
        ['VoiceModal', 'components/VoiceModal.jsx', 'Primary voice UI · chat · mic · TTS'],
        ['RMHelpPrompt', 'components/RMHelpPrompt.jsx', 'Bottom-sheet help entry'],
        ['MpinSheet', 'components/MpinSheet.jsx', 'Blocks voice confirm until MPIN OK'],
        ['simEngine', 'engine/simEngine.js', 'Client wrapper — POST /api/engine/turn'],
        ['HomeScreen + 4 screens', 'components/*Screen.jsx', 'Banking UI + AG-UI integration'],
    ],
    widths=[3, 4.5, 9.5],
)

h2(doc, '5.2 Express API Surface')
tbl(doc,
    ['Route', 'Method', 'Purpose'],
    [
        ['/api/stt', 'POST', 'Audio → transcript (ElevenLabs proxy)'],
        ['/api/tts', 'POST', 'Text → MP3 (Cartesia proxy)'],
        ['/api/engine/turn', 'POST', 'Dialogue state machine step'],
        ['/api/engine/session/:id', 'GET', 'Session state retrieval'],
        ['/api/agui/:agentId', 'POST', 'AG-UI SSE agent stream'],
        ['/api/manifests', 'GET', 'List voice flow manifests'],
        ['/api/account-statement', 'GET', 'Mock balance + transactions'],
        ['/api/health', 'GET', 'Service health + key status'],
        ['/api/mock/*', 'GET/POST', 'Direct mock banking REST'],
    ],
    widths=[4, 1.5, 12.5],
)

h2(doc, '5.3 AI Services & Mock Tools')
tbl(doc,
    ['Service / Tool', 'Provider / File', 'Function'],
    [
        ['STT', 'ElevenLabs Scribe v2', 'Speech → text + language code'],
        ['LLM (voice)', 'OpenAI GPT-4o-mini', 'extract() per turn'],
        ['LLM (AG-UI)', 'OpenAI GPT-4o-mini', 'Function calling + SSE'],
        ['TTS', 'Cartesia sonic-3', 'Text → MP3'],
        ['contacts.search', 'tools.js → backend.js', 'Fuzzy contact lookup'],
        ['execute_upi_payment', 'backend.js', 'Simulated UPI with random failure'],
        ['Mock data store', 'mock.js', '15 contacts · 20 billers · 2 accounts'],
    ],
    widths=[3.5, 4, 9.5],
)

h2(doc, '5.4 Technology Stack')
tbl(doc,
    ['Tier', 'Stack'],
    [
        ['Frontend', 'React 18 · Vite · Tailwind CSS · TypeScript-ready JSX'],
        ['Backend', 'Node.js · Express · ES modules'],
        ['AI', 'ElevenLabs (STT) · OpenAI (LLM) · Cartesia (TTS)'],
        ['Dev', 'concurrently — client :5173 + server :3001'],
    ],
    widths=[3, 14],
)

pb(doc)

# 6. PRODUCTIZATION
h1(doc, '6. Path to Production')
para(doc,
     'The following are not in the demo codebase but are specified in the Production Architecture '
     'document (AI_Voice_RM_Production_Architecture.docx).')

tbl(doc,
    ['Capability', 'Target version', 'Description'],
    [
        ['Bank VPC deployment', 'V1', 'ECS · ALB · Redis · DocumentDB in ap-south-1'],
        ['Real bank APIs', 'V1', 'Replace mock backend with Bank API Adapter'],
        ['Rule engine', 'V1', 'POST /api/trigger · JSON rules · actuation'],
        ['Session persistence', 'V1', 'All turns logged to bank DocumentDB'],
        ['Analytics dashboard', 'V1–V2', 'Volume · outcomes · timelines'],
        ['10 languages', 'V2', 'Extend beyond en/hi/ta/te'],
        ['Push notifications', 'V2', 'FCM / APNs re-engagement'],
        ['Post-session analysis', 'V3', 'Sentiment · churn · QA scoring'],
    ],
    widths=[4, 2, 11],
)

doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run(
    'Silver Suits AI  ·  prateek@silversuits.ai  ·  +91 738 173 2333\n'
    'Document Version 3.0  ·  Code-verified  ·  Confidential  ·  May 2026'
)
rf.font.size = Pt(9)
rf.font.color.rgb = GREY

if __name__ == '__main__':
    print('Generating diagrams...')
    buf1, buf2, buf3, buf4 = get_demo_diagrams()
    # doc built above — save
    out = '/Users/devanshusaindane/SilverSuits/Voice-to-command/Voice-to-Command/Documentation/AI_Voice_RM_System_Architecture.docx'
    doc.save(out)
    print(f'Saved → {out}')
